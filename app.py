"""
LexiAssist v8.0 — Elite AI Legal Engine for Nigerian Lawyers
Single-file deployment with SQLite persistence.
Contract Review · Cost Tracking · User Profiles · Analysis Comparison
Save to Case · Editable References · Custom Templates · Auth Support
"""
from __future__ import annotations

import time
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import hashlib
import html as html_mod
import json
import logging
import os
import re
import psycopg2
import uuid
from datetime import datetime, date
from io import BytesIO
from typing import Any, Optional

from google import genai
from google.genai import types as _genai_types
import pandas as pd
import streamlit as st

try:
    import plotly.express as px
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    import pdfplumber
    HAS_PDF_READ = True
except ImportError:
    HAS_PDF_READ = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _docx_qn
    from docx.oxml import OxmlElement as _OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("LexiAssist")

# ═══════════════════════════════════════════════════════
# PAGE CONFIG
# ═══════════════════════════════════════════════════════
st.set_page_config(
    page_title="LexiAssist v8.0 — Elite Legal AI",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ═══════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════
def _get_db_url() -> str:
    url = ""
    try:
        url = st.secrets["DATABASE_URL"]
    except Exception:
        url = os.getenv("DATABASE_URL", "")
    if not url or not url.strip():
        st.error("❌ DATABASE_URL is not set. Add it to your Streamlit secrets.")
        st.stop()
    # Streamlit Cloud / psycopg2 requires postgresql:// not postgres://
    if url.startswith("postgres://"):
        url = url.replace("postgres://", "postgresql://", 1)
    return url.strip()

# ═══════════════════════════════════════════════════════
# GEMINI MODELS (Best Free Tier – April 2026)
# ═══════════════════════════════════════════════════════
def _parse_models_config():
    models_str = "" 
    try:
        models_str = st.secrets["GEMINI_MODELS"]
    except Exception:
        models_str = os.getenv("GEMINI_MODELS", "")
    if models_str and models_str.strip():
        return [m.strip() for m in models_str.split(",") if m.strip()]
    # Best free models available right now
    return [
        "gemini-2.5-pro",           # ← Highest reasoning quality
        "gemini-2.5-flash",         # ← Best everyday balance (recommended default)
        "gemini-2.5-flash-lite"     # ← Maximum volume when you hit limits
    ]

SUPPORTED_MODELS = _parse_models_config()
DEFAULT_MODEL = "gemini-2.5-flash"   # Change to "gemini-2.5-pro" if you want max quality by default

CASE_STATUSES = ["Active", "Pending", "Completed", "Archived"]
CLIENT_TYPES = ["Individual", "Corporate", "Government", "NGO"]

TASK_TYPES = {
    "general":          {"label": "💬 General Query",            "desc": "Any legal question"},
    "analysis":         {"label": "🔍 Legal Analysis",           "desc": "Issue spotting, CREAC reasoning"},
    "drafting":         {"label": "📄 Document Drafting",        "desc": "Contracts, pleadings, affidavits"},
    "research":         {"label": "📚 Legal Research",           "desc": "Case law, statutes, authorities"},
    "procedure":        {"label": "📋 Procedural Guidance",      "desc": "Filing rules, court practice"},
    "advisory":         {"label": "🎯 Strategic Advisory",       "desc": "Risk mapping, options, strategy"},
    "interpret":        {"label": "⚖️ Statutory Interpretation", "desc": "Legislation analysis"},
    "contract_review":  {"label": "📑 Contract Review",          "desc": "Clause-by-clause risk analysis"},
}

RESPONSE_MODES = {
    "brief":         {"label": "⚡ Brief",          "desc": "Direct answer, 3-5 sentences",        "tokens": 8000,   "temp": 0.1},
    "standard":      {"label": "📝 Standard",       "desc": "Structured analysis, 5-10 paragraphs", "tokens": 32000,  "temp": 0.15},
    "comprehensive": {"label": "🔬 Comprehensive",  "desc": "Full CREAC + Strategy + Risk Ranking",  "tokens": 131072, "temp": 0.2},
}

UPLOAD_TYPES = ["pdf", "docx", "doc", "txt", "xlsx", "xls", "csv", "json", "rtf"]

# Cost per 1M tokens (approx Gemini 2.5 Flash pricing)
COST_PER_1M_INPUT = 0.15
COST_PER_1M_OUTPUT = 0.60

# ═══════════════════════════════════════════════════════
# SYSTEM PROMPTS
# ═══════════════════════════════════════════════════════
IDENTITY_CORE = """You are LexiAssist v8.0 — an elite Senior Partner at a top-tier Nigerian law firm with
35+ years of practice across ALL areas of Nigerian law. You are known for:
- Taking FIRM, CLEAR POSITIONS (never hedging with "may" or "might" when facts permit a conclusion)
- Thinking like a LITIGATOR — always identifying best claim, best defence, weakest party
- Providing ACTIONABLE STRATEGY — not academic theory
- Being BRUTALLY HONEST about risks and exposure

JURISDICTION: Nigeria.
Primary: Constitution of the Federal Republic of Nigeria 1999 (as amended),
Federal Acts, State Laws, Subsidiary Legislation, Rules of Court,
binding case law from the Supreme Court of Nigeria and Court of Appeal.

CITATION INTEGRITY: NEVER fabricate case names or section numbers.
If uncertain, state the legal principle and mark as [CITATION TO BE VERIFIED].
If a case name is well-known and established, cite it confidently.

CRITICAL RULES:
1. TAKE POSITIONS — Say "X IS liable because…" not "X may be liable"
2. ALWAYS identify the WEAKEST PARTY and explain why
3. NEVER end abruptly — always complete your full analysis
4. If the query involves multiple parties, RANK their risk exposure
5. Write to COMPLETION — finish every section you start"""

STRATEGY_BLOCK = """
MANDATORY STRATEGY LAYER (for Standard & Comprehensive modes):
After your legal analysis, you MUST include:

═══ STRATEGIC POSITION ═══
▸ PRIMARY CONCLUSION: State WHO is most exposed and WHY (firm position, no hedging)
▸ RISK RANKING:
  🔴 HIGH RISK → [Party] — [Why]
  🟡 MEDIUM RISK → [Party] — [Why]
  🟢 LOW RISK → [Party] — [Why]

▸ STRATEGY PER PARTY:
  • [Party 1] → [Immediate action recommended]
  • [Party 2] → [Immediate action recommended]
  • [Party 3] → [Immediate action recommended]

▸ LITIGATION ASSESSMENT:
  • Best Claim: [What and by whom]
  • Best Defence: [What and by whom]
  • Weakest Party: [Who and why]
  • Critical Next Step: [Single most important action]
═══════════════════════════
"""

PROMPTS_BY_MODE = {
    "brief": IDENTITY_CORE + """
RESPONSE MODE: BRIEF
- Give the answer in 3-5 clear sentences maximum.
- State your position firmly. No headers, no bullet lists.
- If facts are missing, say: "The outcome turns on X."
- Be direct. Be definitive. No filler.""",

    "standard": IDENTITY_CORE + STRATEGY_BLOCK + """
RESPONSE MODE: STANDARD
- Structure: Issue Identification → Legal Position → Application → Strategy
- Write 5-10 substantial paragraphs of analysis
- Include the MANDATORY STRATEGY LAYER at the end
- COMPLETE your analysis fully — do NOT cut short
- You have ample token space — USE IT to give thorough coverage
- Every paragraph must add value — no repetition""",

    "comprehensive": IDENTITY_CORE + STRATEGY_BLOCK + """
RESPONSE MODE: COMPREHENSIVE (DEEP ANALYSIS)
- This is your MOST THOROUGH mode. Use ALL available space.
- Structure for EACH issue: CONCLUSION → RULE → EXPLANATION → APPLICATION → CONCLUSION (CREAC)
- Identify ALL issues: obvious, hidden, procedural, jurisdictional, limitation
- For EACH issue, cite the governing statute AND at least one leading case
- Include DEVIL'S ADVOCATE section: strongest counter-argument
- Include MANDATORY STRATEGY LAYER (detailed version)
- Include PRACTICAL CHECKLIST of immediate actions
- You have 16,000 tokens available — write a COMPLETE, exhaustive analysis
- NEVER stop mid-analysis — if you identify an issue, ANALYZE it fully
- End with a SUMMARY OF POSITIONS table""",
}

TASK_MODIFIERS = {
    "general": "\nApply the general legal framework. Take a clear position.",
    "analysis": "\nFocus on deep issue-spotting. Apply CREAC to each issue. Distinguish facts carefully.",
    "drafting": "\nDraft a professional Nigerian-standard document. Use [PLACEHOLDER] for missing data. Include all formality requirements (execution, stamping, filing). Do NOT add strategy/risk sections for drafting tasks.",
    "research": "\nWrite a formal Legal Research Memorandum. For each authority: state the principle, quote the ratio (if known), and explain relevance to the query.",
    "procedure": "\nProvide step-by-step procedural guidance. Include: which court, which form/process, filing fees (if known), timelines, and common pitfalls.",
    "advisory": "\nFocus on strategic advisory. Emphasize risk mitigation, commercial impact, and optimal paths. Include risk matrix.",
    "interpret": "\nApply the three rules of statutory interpretation (Literal, Golden, Mischief). State which rule yields the best result and WHY.",
    "contract_review": """
CONTRACT REVIEW MODE — Clause-by-Clause Risk Analysis:
1. For EACH substantive clause, provide:
   • CLAUSE SUMMARY: What it does in plain English
   • RISK LEVEL: 🔴 High / 🟡 Medium / 🟢 Low
   • ISSUES: Legal problems, ambiguities, missing protections
   • RECOMMENDATION: Specific redline or amendment language

2. After clause analysis, include:
═══ RED FLAG MATRIX ═══
| # | Clause | Risk | Issue | Recommended Fix |
|---|--------|------|-------|----------------|
(table of all flagged clauses)

═══ OVERALL ASSESSMENT ═══
▸ Contract Grade: A/B/C/D/F
▸ Signability: Ready / Needs Amendment / Do Not Sign
▸ Top 3 Risks
▸ Missing Clauses (standard protections absent)
═══════════════════════════
""",
}

ISSUE_SPOT_PROMPT = IDENTITY_CORE + """
TASK: Rapid Issue Decomposition (max 250 words)
- CORE ISSUES: List each with area of law and governing principle
- HIDDEN ISSUES: Procedural traps, limitation, standing, regulatory overlap
- MISSING FACTS: Top 3-5 facts that would change the analysis
- COMPLEXITY: Straightforward / Moderate / Complex / Highly Complex
Do NOT provide full analysis. Decomposition ONLY."""

CRITIQUE_PROMPT = IDENTITY_CORE + """
TASK: Quality Assessment of the analysis below (max 150 words).
Score 1-5: Completeness, Legal Accuracy, Strategic Value.
List 1-3 critical gaps. GRADE: A/B/C/D. One sentence overall."""

FOLLOWUP_PROMPT = IDENTITY_CORE + STRATEGY_BLOCK + """
You are continuing a legal conversation.
Context: Original query, previous analysis, and a follow-up question are provided.
- Address the follow-up directly with the same rigor
- Maintain the Litigator/Strategist tone
- Match the specified response mode"""

COMPARISON_PROMPT = IDENTITY_CORE + """
TASK: Compare and contrast the TWO legal analyses provided below.
Structure your comparison as:

═══ ANALYSIS COMPARISON ═══
▸ AREAS OF AGREEMENT: Key points both analyses share
▸ AREAS OF DIVERGENCE: Where they differ and why it matters
▸ THOROUGHNESS: Which is more complete (and what the other missed)
▸ ACCURACY CHECK: Any contradictions or errors in either
▸ VERDICT: Which analysis is BETTER overall and WHY (be specific)
▸ COMBINED RECOMMENDATION: Best position drawing from both
═══════════════════════════

Keep to 300-500 words. Be decisive in your verdict."""

# ═══════════════════════════════════════════════════════
# WITNESS PREPARATION ENGINE — PROMPTS
# ═══════════════════════════════════════════════════════
WITNESS_PREP_SYSTEM = IDENTITY_CORE + """
TASK: Witness Preparation for Nigerian Trial.
You are preparing a witness for court. Your output must be:
1. Examination-in-Chief questions: open-ended, non-leading, logically sequenced, designed to bring out the witness's full account in Nigerian court format.
2. Cross-Examination Risks: realistic, precise attack lines a skilled opponent would deploy — credibility, prior inconsistencies, bias, motive, demeanour weaknesses.
3. Coaching Notes: concise, practical, plain-English instructions the witness can follow.

STRICT RULES:
- Tailor EVERYTHING strictly to the case facts and witness role provided. No generic content.
- Use Nigerian court tone and procedure throughout.
- All questions must be numbered.
- Coaching notes must be actionable, not theoretical.
- Do NOT fabricate facts not given. Flag missing facts with [CLARIFY].
- Keep each section clearly separated with its header.
"""

WITNESS_PREP_PROMPT = """
CASE FACTS:
{case_facts}

WITNESS ROLE: {witness_role}
CASE TYPE: {case_type}

Generate the three sections below. Each section must be clearly labelled.

═══════════════════════════════════
SECTION 1 — EXAMINATION-IN-CHIEF QUESTIONS
═══════════════════════════════════
(Numbered open-ended questions. Non-leading. Structured to build narrative chronologically.)

═══════════════════════════════════
SECTION 2 — CROSS-EXAMINATION RISKS
═══════════════════════════════════
(Bullet-point attack lines. For each: the attack angle, the likely question the opponent asks, and the vulnerability it exploits.)

═══════════════════════════════════
SECTION 3 — COACHING NOTES FOR THE WITNESS
═══════════════════════════════════
(Concise, practical, numbered instructions. What to do, what to avoid, how to behave in the box.)
"""

NEWS_FEED_SUBJECTS = [
    "All Areas",
    "Constitutional Law",
    "Criminal Law & Procedure",
    "Commercial / Contract Law",
    "Company Law",
    "Land / Property Law",
    "Employment & Labour Law",
    "Tax Law",
    "Banking & Finance",
    "Intellectual Property",
    "Family Law",
    "Admiralty / Maritime",
    "Human Rights",
    "Electoral Law",
    "Oil & Gas / Energy",
    "Practice Directions & Court Rules",
    "Legislation Updates",
]

NEWS_FEED_SYSTEM = IDENTITY_CORE + """
TASK: Nigerian Legal News Digest.
You are producing a daily legal intelligence briefing for Nigerian lawyers.
Each item must cover a REAL category of development — new Supreme Court/Court of Appeal decisions,
new legislation, new practice directions, regulatory changes, or notable tribunal rulings.
Do NOT invent specific case names or citation numbers. Describe legal developments at the principle level.
Mark all case references as [CITATION TO BE VERIFIED].
Keep each item tight, practical, and instantly usable by a working lawyer.

STRICT OUTPUT FORMAT — respond ONLY in this exact JSON. Nothing else:
{{
  "generated_date": "DD MMMM YYYY",
  "subject_area": "Subject area covered",
  "items": [
    {{
      "id": 1,
      "title": "Headline title of the development",
      "summary": "2-4 sentence factual summary of what changed or was decided",
      "key_takeaway": "Single sentence — the most critical legal point",
      "practice_impact": "1-2 sentences — what this means for a practising lawyer right now"
    }}
  ]
}}
"""

NEWS_FEED_PROMPT = """
Generate a legal news digest for Nigerian lawyers covering: {subject_area}.
Focus on developments that would have occurred in the last 30-90 days (you may use representative/
typical examples if specific recent cases are uncertain — but mark them [REPRESENTATIVE EXAMPLE]).
Generate exactly 6 news items.
Today's reference date: {today}.
"""

# ═══════════════════════════════════════════════════════
# WITNESS RE-EXAMINATION PROMPT
# ═══════════════════════════════════════════════════════
REEXAM_SYSTEM = IDENTITY_CORE + """
TASK: Generate Re-Examination (Re-Direct) Questions for a Nigerian trial witness.
You are given the cross-examination attack points that the opponent used or will likely use.
Your job is to generate precise, non-leading re-examination questions that REHABILITATE the witness
on each attack point — restoring credibility, clarifying inconsistencies, and neutralising bias allegations.

RULES:
- Only re-examine on matters arising from cross-examination. Do not introduce new matters.
- Questions must be open-ended and non-leading (as required in Nigerian courts under Evidence Act 2011).
- For each attack point addressed, label it clearly.
- End with a brief "Re-examination Strategy Note" on sequencing and emphasis.
- Nigerian court procedure throughout.
"""

REEXAM_PROMPT = """
WITNESS ROLE: {witness_role}
CASE FACTS: {case_facts}

CROSS-EXAMINATION ATTACK POINTS IDENTIFIED:
{cross_exam_risks}

Generate targeted re-examination questions to rehabilitate this witness on each attack point above.
Number each question. Label each attack point being addressed.
End with a Re-examination Strategy Note (3-5 sentences).
"""

# ═══════════════════════════════════════════════════════
# WITNESS CONTRADICTION DETECTOR PROMPT
# ═══════════════════════════════════════════════════════
CONTRADICTION_SYSTEM = IDENTITY_CORE + """
TASK: Multi-Witness Contradiction Analysis for Nigerian trial preparation.
You are given the prepared briefs of two or more witnesses. Your job is to:
1. Identify DIRECT CONTRADICTIONS — where witnesses give conflicting accounts of the same fact
2. Identify GAPS — where one witness's account raises questions the other doesn't address
3. Identify CORROBORATIONS — strong points where accounts align and reinforce each other
4. Provide a Reconciliation Strategy — how counsel can address contradictions before trial

CRITICAL: A contradiction in a prosecution witness and a defence witness may be a strategic advantage.
Distinguish between intra-party contradictions (dangerous) and inter-party ones (expected/exploitable).
Be specific. Quote the conflicting passages directly.
"""

CONTRADICTION_PROMPT = """
Below are the prepared witness briefs for {count} witnesses in this matter.
Analyse for contradictions, gaps, and corroborations.

{witness_summaries}

Structure your output:
1. DIRECT CONTRADICTIONS (each numbered, with both versions quoted)
2. GAPS & UNANSWERED QUESTIONS
3. STRONG CORROBORATIONS
4. RECONCILIATION STRATEGY FOR COUNSEL
"""

# ═══════════════════════════════════════════════════════
# NEWS DEEP DIVE PROMPT
# ═══════════════════════════════════════════════════════
NEWS_DEEPDIVE_SYSTEM = IDENTITY_CORE + STRATEGY_BLOCK + """
TASK: Full legal analysis of a recent Nigerian legal development.
You are given a news item describing a recent case, legislation, or practice direction.
Provide a comprehensive analysis covering: what it means legally, how it changes the law (if at all),
the practical impact on specific practice areas, potential challenges or arguments against it,
and what actions a prudent lawyer should take now.
Use Nigerian law throughout. Mark all case citations as [CITATION TO BE VERIFIED].
"""

NEWS_DEEPDIVE_PROMPT = """
Analyse this Nigerian legal development in full:

TITLE: {title}
SUMMARY: {summary}
KEY TAKEAWAY: {takeaway}
PRACTICE IMPACT: {impact}

Provide a comprehensive Standard-mode legal analysis. Cover:
1. Legal significance and how it fits into existing Nigerian law
2. Which practice areas are affected and how
3. Arguments for and against the position taken
4. Immediate actions a practising lawyer should take
5. Strategic advisory for affected clients
"""

# ═══════════════════════════════════════════════════════
# NEWS RELEVANCE SCAN PROMPT
# ═══════════════════════════════════════════════════════
NEWS_RELEVANCE_SYSTEM = IDENTITY_CORE + """
TASK: Case Relevance Scanner.
You are given a lawyer's case facts and a list of recent Nigerian legal developments.
Score each development for relevance to the case facts on a scale of 0-10.
For each relevant item (score ≥ 5), explain precisely how it affects the case — favourable,
unfavourable, or procedural implications.
Sort output from most relevant to least relevant.
Respond ONLY in this exact JSON format, nothing else:
{
  "scan_summary": "1-2 sentence overview of the most important findings",
  "items": [
    {
      "id": 1,
      "title": "Title of the news item",
      "relevance_score": 8,
      "relevance_label": "HIGH / MEDIUM / LOW / NOT RELEVANT",
      "how_it_affects_case": "Specific explanation of impact on the facts given",
      "favourable_or_unfavourable": "FAVOURABLE / UNFAVOURABLE / NEUTRAL / PROCEDURAL"
    }
  ]
}
"""

NEWS_RELEVANCE_PROMPT = """
CASE FACTS:
{case_facts}

RECENT LEGAL DEVELOPMENTS TO SCAN:
{news_items}

Score each development for relevance to these case facts. Include ALL items in your response,
even those with score 0. Sort by relevance_score descending.
"""

# ═══════════════════════════════════════════════════════
# LEGAL FEE & STAMP DUTY CALCULATOR — DATA
# ═══════════════════════════════════════════════════════
# Nigerian Legal Practitioners (Remuneration for Legal Documentation
# and Other Land Matters) Order — sliding scale
LAND_MATTERS_SCALE = [
    {"band_label": "First ₦5,000",           "up_to": 5_000,         "rate": 0.10},
    {"band_label": "Next ₦10,000",           "up_to": 15_000,        "rate": 0.075},
    {"band_label": "Next ₦15,000",           "up_to": 30_000,        "rate": 0.05},
    {"band_label": "Next ₦70,000",           "up_to": 100_000,       "rate": 0.035},
    {"band_label": "Next ₦400,000",          "up_to": 500_000,       "rate": 0.025},
    {"band_label": "Remainder above ₦500k",  "up_to": float("inf"),  "rate": 0.015},
]
LAND_MATTERS_MIN_FEE = 10_000  # ₦10,000 minimum for any land transaction

# Stamp Duty rates — Stamp Duties Act Cap S8 LFN 2004 (as amended)
STAMP_DUTY_INSTRUMENTS = {
    "deed_of_assignment":       {"label": "Deed of Assignment / Conveyance",    "rate": 0.015,  "basis": "property_value", "note": "1.5% of property value"},
    "tenancy_lt7":              {"label": "Tenancy / Lease (< 7 years)",         "rate": 0.0078, "basis": "annual_rent_x_years", "note": "0.78% × annual rent × years"},
    "tenancy_7to21":            {"label": "Tenancy / Lease (7–21 years)",        "rate": 0.03,   "basis": "annual_rent",    "note": "3% of annual rent"},
    "tenancy_over21":           {"label": "Tenancy / Lease (> 21 years)",        "rate": 0.06,   "basis": "annual_rent",    "note": "6% of annual rent"},
    "mortgage":                 {"label": "Legal Mortgage / Debenture",          "rate": 0.00375,"basis": "loan_amount",    "note": "0.375% of secured amount"},
    "power_of_attorney_gen":    {"label": "General Power of Attorney",           "rate": None,   "basis": "flat",           "flat": 1_000, "note": "₦1,000 flat"},
    "power_of_attorney_spec":   {"label": "Special Power of Attorney",           "rate": None,   "basis": "flat",           "flat": 500,   "note": "₦500 flat"},
    "affidavit":                {"label": "Affidavit",                           "rate": None,   "basis": "flat",           "flat": 200,   "note": "₦200 flat"},
    "memorandum_of_understanding": {"label": "Memorandum of Understanding (MOU)","rate": 0.0075,"basis": "contract_value",  "note": "0.75% of stated value"},
    "share_transfer":           {"label": "Share Transfer / Stock Transfer Form","rate": 0.015,  "basis": "consideration",  "note": "1.5% of consideration"},
    "loan_agreement":           {"label": "Loan / Credit Agreement",             "rate": 0.00375,"basis": "loan_amount",    "note": "0.375% of loan amount"},
    "guarantee":                {"label": "Guarantee / Indemnity",               "rate": 0.0075, "basis": "guaranteed_sum", "note": "0.75% of guaranteed sum"},
    "joint_venture":            {"label": "Joint Venture Agreement",             "rate": 0.0075, "basis": "contract_value",  "note": "0.75% of stated value"},
    "settlement_agreement":     {"label": "Deed of Settlement / Release",        "rate": 0.015,  "basis": "property_value",  "note": "1.5% of settlement amount"},
}

# Court Filing Fees (indicative — varies by state and court rules; verify before filing)
COURT_FILING_FEES = {
    "magistrate_lagos": {
        "label": "Magistrate Court (Lagos State)",
        "bands": [
            {"claim_max": 100_000,     "fee": 2_000,  "label": "Claim up to ₦100,000"},
            {"claim_max": 300_000,     "fee": 5_000,  "label": "Claim ₦100k–₦300k"},
            {"claim_max": 500_000,     "fee": 8_000,  "label": "Claim ₦300k–₦500k"},
            {"claim_max": float("inf"),"fee": 10_000, "label": "Maximum jurisdiction"},
        ],
        "appeal_fee": 15_000,
        "note": "Lagos Magistrate Courts Law 2009 (as amended). Max civil jurisdiction ₦500,000.",
    },
    "high_court_state": {
        "label": "State High Court (e.g. Lagos)",
        "bands": [
            {"claim_max": 500_000,       "fee": 8_000,  "label": "Claim up to ₦500,000"},
            {"claim_max": 1_000_000,     "fee": 15_000, "label": "Claim ₦500k–₦1m"},
            {"claim_max": 5_000_000,     "fee": 25_000, "label": "Claim ₦1m–₦5m"},
            {"claim_max": 20_000_000,    "fee": 40_000, "label": "Claim ₦5m–₦20m"},
            {"claim_max": 100_000_000,   "fee": 75_000, "label": "Claim ₦20m–₦100m"},
            {"claim_max": float("inf"),  "fee": 120_000,"label": "Claim above ₦100m"},
        ],
        "appeal_fee": 50_000,
        "note": "Fees vary by state. Verify with specific court registry before filing.",
    },
    "federal_high_court": {
        "label": "Federal High Court",
        "bands": [
            {"claim_max": 1_000_000,     "fee": 10_000,  "label": "Claim up to ₦1m"},
            {"claim_max": 5_000_000,     "fee": 30_000,  "label": "Claim ₦1m–₦5m"},
            {"claim_max": 20_000_000,    "fee": 60_000,  "label": "Claim ₦5m–₦20m"},
            {"claim_max": 100_000_000,   "fee": 100_000, "label": "Claim ₦20m–₦100m"},
            {"claim_max": float("inf"),  "fee": 150_000, "label": "Claim above ₦100m"},
        ],
        "appeal_fee": 75_000,
        "note": "FHC (Civil Procedure) Rules 2019. Verify current rates with court registry.",
    },
    "national_industrial_court": {
        "label": "National Industrial Court",
        "bands": [
            {"claim_max": 1_000_000,    "fee": 10_000,  "label": "Claim up to ₦1m"},
            {"claim_max": 10_000_000,   "fee": 25_000,  "label": "Claim ₦1m–₦10m"},
            {"claim_max": float("inf"), "fee": 50_000,  "label": "Claim above ₦10m"},
        ],
        "appeal_fee": 50_000,
        "note": "NIC (Civil Procedure) Rules 2017.",
    },
    "court_of_appeal": {
        "label": "Court of Appeal",
        "bands": [
            {"claim_max": float("inf"), "fee": 100_000, "label": "All civil appeals"},
        ],
        "appeal_fee": 0,
        "note": "Court of Appeal Rules 2021. Filing fee covers Notice of Appeal and related documents.",
    },
    "supreme_court": {
        "label": "Supreme Court of Nigeria",
        "bands": [
            {"claim_max": float("inf"), "fee": 200_000, "label": "All matters"},
        ],
        "appeal_fee": 0,
        "note": "Supreme Court Rules (as amended). Verify with Supreme Court Registry, Abuja.",
    },
}


def compute_land_fee(value: float) -> tuple[float, list]:
    """Compute solicitor's fee on a land transaction using the sliding scale.
    Returns (total_fee, breakdown_rows)."""
    fee = 0.0
    breakdown = []
    remaining = value
    prev_band = 0.0
    for band in LAND_MATTERS_SCALE:
        if remaining <= 0:
            break
        cap = band["up_to"]
        taxable = min(remaining, cap - prev_band)
        if taxable <= 0:
            prev_band = cap
            continue
        band_fee = taxable * band["rate"]
        fee += band_fee
        breakdown.append({
            "band": band["band_label"],
            "taxable": taxable,
            "rate": f"{band['rate']*100:.2f}%",
            "fee": band_fee,
        })
        remaining -= taxable
        prev_band = cap
    fee = max(fee, LAND_MATTERS_MIN_FEE)
    return fee, breakdown


def compute_stamp_duty(instrument_key: str, value: float = 0,
                       years: float = 1, annual_rent: float = 0) -> float:
    """Compute stamp duty for an instrument type."""
    inst = STAMP_DUTY_INSTRUMENTS.get(instrument_key)
    if not inst:
        return 0.0
    basis = inst["basis"]
    if basis == "flat":
        return float(inst.get("flat", 0))
    if basis == "property_value":
        return value * inst["rate"]
    if basis == "annual_rent_x_years":
        return (annual_rent or value) * years * inst["rate"]
    if basis == "annual_rent":
        return (annual_rent or value) * inst["rate"]
    if basis in ("loan_amount", "consideration", "guaranteed_sum",
                 "contract_value", "secured_amount"):
        return value * inst["rate"]
    return 0.0


def get_court_filing_fee(court_key: str, claim_value: float) -> tuple[int, str]:
    """Return (fee, note) for filing in a given court with a given claim value."""
    court = COURT_FILING_FEES.get(court_key)
    if not court:
        return 0, ""
    for band in court["bands"]:
        if claim_value <= band["claim_max"]:
            return band["fee"], court["note"]
    return court["bands"][-1]["fee"], court["note"]


# ═══════════════════════════════════════════════════════
# SETTLEMENT & ADR ADVISOR — PROMPTS
# ═══════════════════════════════════════════════════════
SETTLEMENT_SYSTEM = IDENTITY_CORE + STRATEGY_BLOCK + """
TASK: Settlement & ADR Advisory for Nigerian Legal Practice.
You are advising a Nigerian lawyer on settlement strategy and alternative dispute resolution.
Apply your knowledge of: Arbitration and Conciliation Act 2023, Lagos Multi-Door Courthouse (LMDC),
Abuja Multi-Door Courthouse, Rules of Professional Conduct on settlement duties (Rule 17, 22),
Evidence Act 2011 (without prejudice communications), and standard Nigerian litigation practice.

Your output must be structured, firm, and immediately actionable.
Give concrete numbers (settlement ranges, percentages, timelines) — do NOT hedge.
Identify the weaker party, their pressure points, and the optimal strategy for the instructing party.
"""

SETTLEMENT_PROMPT = """
INSTRUCTING PARTY: {instructing_party}
OPPOSING PARTY: {opposing_party}
CASE TYPE: {case_type}
CLAIM AMOUNT / SUBJECT MATTER VALUE: ₦{claim_amount}
COURT / STAGE: {court_stage}
STRENGTH OF CASE (self-assessed): {strength}
URGENCY / TIME PRESSURE: {urgency}

CASE FACTS:
{case_facts}

Generate a full Settlement & ADR Advisory structured as follows:

═══════════════════════════════════════════
SECTION 1 — SETTLEMENT VALUE ANALYSIS
═══════════════════════════════════════════
(Compute: realistic settlement band, ideal settlement amount, walk-away floor/ceiling.
Give actual ₦ figures, not just percentages. Show your reasoning.)

═══════════════════════════════════════════
SECTION 2 — NEGOTIATION STRATEGY
═══════════════════════════════════════════
(Opening position, key concessions to offer, key concessions to demand, sequence of moves.
Identify the opposing party's pressure points and how to exploit them.)

═══════════════════════════════════════════
SECTION 3 — ADR ROUTE RECOMMENDATION
═══════════════════════════════════════════
(Should this go to mediation, arbitration, or direct negotiation? Which ADR centre?
Timeline and cost estimate for ADR vs. continued litigation.)

═══════════════════════════════════════════
SECTION 4 — WITHOUT PREJUDICE OFFER DRAFT
═══════════════════════════════════════════
(Draft a concise "Without Prejudice Save as to Costs" opening offer letter.
Include: offer amount, conditions, deadline, and reservation of rights.)

═══════════════════════════════════════════
SECTION 5 — RISK IF NO SETTLEMENT
═══════════════════════════════════════════
(Litigation risk, cost exposure, likely trial outcome, enforcement risk.
Be brutally honest about the weakest party's position.)
"""

# ═══════════════════════════════════════════════════════
# DUE DILIGENCE ENGINE — PROMPTS & DATA
# ═══════════════════════════════════════════════════════
DD_TRANSACTION_TYPES = {
    "property_purchase":    "🏠 Property / Land Acquisition",
    "company_acquisition":  "🏢 Company / Business Acquisition",
    "loan_security":        "💳 Loan & Security / Debenture",
    "joint_venture":        "🤝 Joint Venture / Partnership",
    "franchise":            "🏪 Franchise Agreement",
    "employment_senior":    "👔 Senior Employment / Directorship",
    "oil_gas_block":        "⛽ Oil & Gas Block / Farm-in",
    "real_estate_dev":      "🏗️ Real Estate Development",
    "ipo_capital_market":   "📈 IPO / Capital Market Transaction",
    "fintech_regulatory":   "📱 Fintech / Payment Service",
}

DD_SYSTEM = IDENTITY_CORE + """
TASK: Generate a comprehensive Nigerian Due Diligence Checklist.
You are preparing a due diligence report framework for a Nigerian transaction.
Apply your knowledge of: Companies and Allied Matters Act (CAMA) 2020, Land Use Act 1978,
Corporate Affairs Commission (CAC) practice, Nigerian Investment Promotion Commission Act,
Securities and Exchange Commission Rules, NDIC regulations, CBN regulations (as applicable),
Federal Inland Revenue Service requirements, and standard Nigerian conveyancing practice.

Structure your output as a complete, actionable checklist that a Nigerian lawyer can take
immediately into the field/office. For each item:
- State what to search/obtain/verify
- State which registry, authority, or source
- Flag the risk if the item is not clear
- Mark priority: 🔴 Critical / 🟡 Important / 🟢 Standard

CRITICAL: Tailor EVERYTHING to the specific transaction type and jurisdiction described.
Do NOT produce a generic checklist. Every item must be specific to Nigerian law and practice.
"""

DD_PROMPT = """
TRANSACTION TYPE: {transaction_type}
TRANSACTION VALUE: ₦{transaction_value}
JURISDICTION: {jurisdiction}
PARTIES: {parties}
BRIEF TRANSACTION DESCRIPTION: {description}
SPECIAL CONCERNS: {special_concerns}

Generate a comprehensive due diligence checklist structured as:

1. CORPORATE / ENTITY SEARCHES
2. TITLE / PROPERTY SEARCHES (if applicable)
3. REGULATORY / LICENSING SEARCHES  
4. FINANCIAL & TAX SEARCHES
5. LITIGATION / ENCUMBRANCE SEARCHES
6. CONTRACTS & COMMERCIAL DOCUMENTS
7. EMPLOYMENT & LABOUR (if applicable)
8. ENVIRONMENTAL / SECTOR-SPECIFIC (if applicable)
9. TRANSACTION STRUCTURE RISK FLAGS
10. RECOMMENDED CONDITIONS PRECEDENT

For each section, provide numbered checklist items with priority flags.
End with a CRITICAL PATH — the 5 searches that must be completed first and why.
"""
DEFAULT_LIMITATION_PERIODS = [
    {"cause": "Simple Contract", "period": "6 years", "authority": "Limitation Act, s. 8(1)(a)"},
    {"cause": "Tort / Negligence", "period": "6 years", "authority": "Limitation Act, s. 8(1)(a)"},
    {"cause": "Personal Injury", "period": "3 years", "authority": "Limitation Act, s. 8(1)(b)"},
    {"cause": "Defamation", "period": "3 years", "authority": "Limitation Act, s. 8(1)(b)"},
    {"cause": "Recovery of Land", "period": "12 years", "authority": "Limitation Act, s. 16"},
    {"cause": "Mortgage Foreclosure", "period": "12 years", "authority": "Limitation Act, s. 18"},
    {"cause": "Recovery of Rent", "period": "6 years", "authority": "Limitation Act, s. 19"},
    {"cause": "Judgment Enforcement", "period": "12 years", "authority": "Limitation Act, s. 8(1)(d)"},
    {"cause": "POPA (Public Officers)", "period": "3 months notice / 12 months suit", "authority": "POPA, s. 2"},
    {"cause": "Fundamental Rights", "period": "12 months", "authority": "FREP Rules, Order II r. 1"},
    {"cause": "Election Petition", "period": "21 days post-declaration", "authority": "Electoral Act 2022, s. 133(1)"},
]

COURT_HIERARCHY = [
    {"level": 1, "name": "Supreme Court of Nigeria", "desc": "Final appellate court", "icon": "🏛️"},
    {"level": 2, "name": "Court of Appeal", "desc": "Intermediate appellate", "icon": "⚖️"},
    {"level": 3, "name": "Federal High Court", "desc": "Federal causes, tax, admiralty", "icon": "🏢"},
    {"level": 3, "name": "State High Courts", "desc": "General civil & criminal", "icon": "🏢"},
    {"level": 3, "name": "National Industrial Court", "desc": "Labour & employment", "icon": "🏢"},
    {"level": 4, "name": "Magistrate / District Courts", "desc": "Summary jurisdiction", "icon": "📋"},
    {"level": 4, "name": "Customary / Sharia Courts", "desc": "Personal law matters", "icon": "📋"},
]

DEFAULT_LEGAL_MAXIMS = [
    {"maxim": "Audi alteram partem", "meaning": "Hear the other side — natural justice"},
    {"maxim": "Nemo judex in causa sua", "meaning": "No one should judge their own cause"},
    {"maxim": "Stare decisis", "meaning": "Stand by decided cases — binding precedent"},
    {"maxim": "Ubi jus ibi remedium", "meaning": "Where there is a right, there is a remedy"},
    {"maxim": "Volenti non fit injuria", "meaning": "No injury to one who consents"},
    {"maxim": "Pacta sunt servanda", "meaning": "Agreements must be honoured"},
    {"maxim": "Nemo dat quod non habet", "meaning": "No one gives what they don't have"},
    {"maxim": "Res judicata", "meaning": "A decided matter cannot be re-litigated"},
    {"maxim": "Actus non facit reum nisi mens sit rea", "meaning": "No guilt without guilty mind"},
    {"maxim": "Ignorantia legis neminem excusat", "meaning": "Ignorance of law excuses no one"},
    {"maxim": "Qui facit per alium facit per se", "meaning": "He who acts through another acts himself"},
    {"maxim": "Generalia specialibus non derogant", "meaning": "General provisions don't override specific ones"},
]

DEFAULT_TEMPLATES = [
    {"id": "builtin_1", "name": "Employment Contract", "cat": "Corporate", "builtin": True,
     "content": "EMPLOYMENT CONTRACT\n\nMade on [DATE] between:\n\n1. [EMPLOYER NAME] (\"Employer\")\n   RC: [NUMBER]\n\n2. [EMPLOYEE NAME] (\"Employee\")\n\nTERMS:\n1. Position: [TITLE]\n2. Start: [DATE]\n3. Probation: [MONTHS]\n4. Salary: N[AMOUNT]/month\n5. Hours: [X] hrs/week\n6. Leave: [X] days/year\n7. Termination: [NOTICE] written notice\n8. Governing Law: Labour Act of Nigeria\n\nSigned:\n_______ (Employer)\n_______ (Employee)"},
    {"id": "builtin_2", "name": "Tenancy Agreement", "cat": "Property", "builtin": True,
     "content": "TENANCY AGREEMENT\n\nMade on [DATE] BETWEEN:\n[LANDLORD] of [ADDRESS] (\"Landlord\")\nAND\n[TENANT] of [ADDRESS] (\"Tenant\")\n\n1. Premises: [ADDRESS]\n2. Term: [DURATION] from [START]\n3. Rent: N[AMOUNT] per [PERIOD]\n4. Deposit: N[AMOUNT]\n5. Use: [Residential/Commercial]\n6. Governing Law: Applicable State Tenancy Law\n\nSigned:\n_______ _______"},
    {"id": "builtin_3", "name": "Power of Attorney", "cat": "Litigation", "builtin": True,
     "content": "GENERAL POWER OF ATTORNEY\n\nI, [GRANTOR], of [ADDRESS], appoint [ATTORNEY] of [ADDRESS] as my Attorney.\n\nPOWERS:\n1. Recover debts and execute settlements\n2. Manage real and personal property\n3. Appear before any court or tribunal\n\nIRREVOCABLE for [PERIOD].\n\nDated: [DATE]\nSigned: _______\nWitness: _______"},
    {"id": "builtin_4", "name": "Written Address (Skeleton)", "cat": "Litigation", "builtin": True,
     "content": "IN THE [COURT NAME]\nSUIT NO: [NUMBER]\n\nBETWEEN:\n[CLAIMANT] ............ Claimant\nAND\n[DEFENDANT] ........... Defendant\n\nWRITTEN ADDRESS OF THE [PARTY]\n\n1.0 INTRODUCTION\n2.0 BRIEF FACTS\n3.0 ISSUES FOR DETERMINATION\n4.0 ARGUMENTS\n   4.1 Issue One\n   4.2 Issue Two\n5.0 CONCLUSION\n\nDated: [DATE]\nCounsel: _______"},
    {"id": "builtin_5", "name": "Demand Letter", "cat": "Commercial", "builtin": True,
     "content": "OUR REF: [REF]\nDATE: [DATE]\n\n[RECIPIENT NAME]\n[ADDRESS]\n\nDear Sir/Madam,\n\nRE: DEMAND FOR PAYMENT OF N[AMOUNT]\n\nWe are Solicitors to [CLIENT NAME] on whose instructions we write.\n\nOur client instructs us that [FACTS].\n\nDEMAND: Pay N[AMOUNT] within [DAYS] days.\n\nFailing which, we have firm instructions to commence legal proceedings without further notice.\n\nYours faithfully,\n[FIRM NAME]"},
]

# ═══════════════════════════════════════════════════════
# THEMES (CSS)
# ═══════════════════════════════════════════════════════
THEMES = {
    "⚖️ Corporate": {
        "display_name": "⚖️ Corporate",
        "description": "Deep navy & gold — professional law firm portal.",
        "bg":               "#f4f6f9",
        "bg_secondary":     "#e8edf4",
        "card_bg":          "#ffffff",
        "border":           "#c5d0e0",
        "text":             "#1a2e4a",
        "text_secondary":   "#4a6080",
        "accent":           "#1a2e4a",
        "accent_secondary": "#c9a84c",
        "positive":         "#059669",
        "negative":         "#dc2626",
        "warning":          "#d97706",
        "sidebar_bg":       "#1a2e4a",
        "input_bg":         "#ffffff",
        "header_gradient":  "linear-gradient(135deg, #1a2e4a, #2d4a6e)",
    },
    "🌿 Emerald": {
        "display_name": "🌿 Emerald",
        "description": "Fresh greens — calm and focused.",
        "bg":               "#f8faf9",
        "bg_secondary":     "#edf7f2",
        "card_bg":          "#ffffff",
        "border":           "#a7d4bc",
        "text":             "#1e293b",
        "text_secondary":   "#3d6b54",
        "accent":           "#047857",
        "accent_secondary": "#0d9488",
        "positive":         "#10b981",
        "negative":         "#dc2626",
        "warning":          "#d97706",
        "sidebar_bg":       "#064e3b",
        "input_bg":         "#ffffff",
        "header_gradient":  "linear-gradient(135deg, #059669, #0d9488)",
    },
    "🌊 Deep Ocean": {
        "display_name": "🌊 Deep Ocean",
        "description": "Calm deep blues — focused and serene.",
        "bg":               "#0B1120",
        "bg_secondary":     "#111B2E",
        "card_bg":          "#14203A",
        "border":           "#1E3A5F",
        "text":             "#E0E7FF",
        "text_secondary":   "#8899BB",
        "accent":           "#64FFDA",
        "accent_secondary": "#7BDFF2",
        "positive":         "#52D68A",
        "negative":         "#FF7675",
        "warning":          "#FFD93D",
        "sidebar_bg":       "#0D1526",
        "input_bg":         "#162040",
        "header_gradient":  "linear-gradient(135deg, #0D1526, #1E3A5F)",
    },
    "🌙 Midnight": {
        "display_name": "🌙 Midnight",
        "description": "Deep purples — contemplative and restful.",
        "bg":               "#0D0B1A",
        "bg_secondary":     "#131024",
        "card_bg":          "#1A1530",
        "border":           "#2E2660",
        "text":             "#E0DCFF",
        "text_secondary":   "#9990CC",
        "accent":           "#A29BFE",
        "accent_secondary": "#C4B5FD",
        "positive":         "#6BCB77",
        "negative":         "#FF6B6B",
        "warning":          "#FFD93D",
        "sidebar_bg":       "#0A0818",
        "input_bg":         "#1E1838",
        "header_gradient":  "linear-gradient(135deg, #0A0818, #2E2660)",
    },
    "🔥 Ember": {
        "display_name": "🔥 Ember",
        "description": "Dark with warm amber — bold and intense.",
        "bg":               "#1A1210",
        "bg_secondary":     "#221812",
        "card_bg":          "#2A1E18",
        "border":           "#5C3D2E",
        "text":             "#FFE4CC",
        "text_secondary":   "#C4977A",
        "accent":           "#FF9800",
        "accent_secondary": "#FFB74D",
        "positive":         "#66BB6A",
        "negative":         "#EF5350",
        "warning":          "#FFD54F",
        "sidebar_bg":       "#16100C",
        "input_bg":         "#30241C",
        "header_gradient":  "linear-gradient(135deg, #16100C, #5C3D2E)",
    },
    "💜 Lavender": {
        "display_name": "💜 Lavender",
        "description": "Soft purples — soothing and creative.",
        "bg":               "#14101A",
        "bg_secondary":     "#1C1626",
        "card_bg":          "#221C30",
        "border":           "#3D3060",
        "text":             "#E8E0F0",
        "text_secondary":   "#A898C8",
        "accent":           "#B388FF",
        "accent_secondary": "#CE93D8",
        "positive":         "#69F0AE",
        "negative":         "#FF8A80",
        "warning":          "#FFE57F",
        "sidebar_bg":       "#110D16",
        "input_bg":         "#281E3A",
        "header_gradient":  "linear-gradient(135deg, #110D16, #3D3060)",
    },
    "☁️ Cloud": {
        "display_name": "☁️ Cloud",
        "description": "Light grays and sky blues — clean and airy.",
        "bg":               "#F5F7FA",
        "bg_secondary":     "#E8ECF1",
        "card_bg":          "#FFFFFF",
        "border":           "#D1D9E6",
        "text":             "#1a202c",
        "text_secondary":   "#4a5568",
        "accent":           "#2563EB",
        "accent_secondary": "#3B82F6",
        "positive":         "#48BB78",
        "negative":         "#FC8181",
        "warning":          "#ECC94B",
        "sidebar_bg":       "#2D3748",
        "input_bg":         "#FFFFFF",
        "header_gradient":  "linear-gradient(135deg, #2D3748, #4299E1)",
    },
    "🌅 Sunset": {
        "display_name": "🌅 Sunset",
        "description": "Warm oranges and ambers — vibrant and expressive.",
        "bg":               "#1A100D",
        "bg_secondary":     "#221610",
        "card_bg":          "#2C1D16",
        "border":           "#6B3A28",
        "text":             "#FFE8D6",
        "text_secondary":   "#C49A7E",
        "accent":           "#FF6B35",
        "accent_secondary": "#FF9F1C",
        "positive":         "#7DCE82",
        "negative":         "#FF4757",
        "warning":          "#FFBE0B",
        "sidebar_bg":       "#15100A",
        "input_bg":         "#30221A",
        "header_gradient":  "linear-gradient(135deg, #15100A, #6B3A28)",
    },
}

THEME_NAMES = list(THEMES.keys())


def get_theme(name: str) -> dict:
    return THEMES.get(name, THEMES["⚖️ Corporate"])


def get_plotly_colors(name: str) -> dict:
    """Return Plotly-compatible colour config matching the active theme."""
    t = get_theme(name)
    return {
        "paper":  t["card_bg"],
        "text":   t["text"],
        "grid":   t["border"],
        "accent": t["accent"],
        "colors": [
            t["accent"], t["accent_secondary"], t["positive"],
            t["negative"], t["warning"], t["text_secondary"],
            "#FF6B6B", "#48DBFB", "#FECA57", "#FF9FF3",
        ],
    }


def get_theme_css(
    theme_name: str,
    font_size_scale: float = 1.0,
    high_contrast: bool = False,
    reduce_motion: bool = False,
) -> str:
    t = get_theme(theme_name)

    text_color = "#FFFFFF" if high_contrast else t["text"]
    text_sec   = "#CCCCCC" if high_contrast else t["text_secondary"]
    bg_color   = "#000000" if (high_contrast and t["bg"][1:3] < "33") else t["bg"]
    base_font  = round(16 * font_size_scale, 1)
    input_font = round(base_font * 0.94, 1)
    mobile_font = round(base_font * 0.92, 1)

    is_light        = int(t["card_bg"].lstrip("#")[0:2], 16) >= 0x77
    is_dark_sidebar = int(t["sidebar_bg"].lstrip("#")[0:2], 16) < 0x44

    if is_light:
        shadow_card  = "0 1px 3px rgba(0,0,0,0.05),0 2px 10px rgba(0,0,0,0.06)"
        shadow_hover = "0 4px 20px rgba(0,0,0,0.09),0 2px 8px rgba(0,0,0,0.05)"
    else:
        shadow_card  = f"0 0 0 1px {t['border']}"
        shadow_hover = f"0 0 0 1px {t['accent']}55"

    sb_text     = "#e6edf8" if is_dark_sidebar else t["text"]
    sb_text_2   = "#8fa5c8" if is_dark_sidebar else t["text_secondary"]
    sb_line     = "rgba(255,255,255,0.08)" if is_dark_sidebar else t["border"]
    sb_input_bg = t["input_bg"]
    sb_input_tx = t["text"]
    sb_hover_bg = "rgba(255,255,255,0.06)" if is_dark_sidebar else t["bg_secondary"]

    badge_ok_bg   = "#dcfce7" if is_light else "#14532d55"
    badge_ok_tx   = "#15803d" if is_light else "#4ade80"
    badge_warn_bg = "#fef9c3" if is_light else "#71360055"
    badge_warn_tx = "#854d0e" if is_light else "#facc15"
    badge_err_bg  = "#fee2e2" if is_light else "#7f1d1d55"
    badge_err_tx  = "#b91c1c" if is_light else "#f87171"
    badge_inf_bg  = f"{t['accent']}18"
    badge_inf_tx  = t["accent"]

    ph_col      = "rgba(30,46,80,0.40)" if is_light else "rgba(200,215,240,0.35)"
    disc_bg     = t["warning"] + ("15" if is_light else "20")
    disc_tx     = "#78350f" if is_light else text_sec

    motion_css = ""
    if reduce_motion:
        motion_css = "*, *::before, *::after { animation-duration:0.01ms!important; transition-duration:0.01ms!important; }"

    acc = t["accent"]
    acc2 = t["accent_secondary"]
    warn = t["warning"]
    pos = t["positive"]
    border = t["border"]
    card = t["card_bg"]
    bg2 = t["bg_secondary"]
    inp = t["input_bg"]
    grad = t["header_gradient"]
    sidebar_bg = t["sidebar_bg"]

    return f"""<style>
/* ── Google Fonts — loaded inside style tag (no external link tags) ── */
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');

/* ── Custom properties ── */
:root {{
  --la-bg:{bg_color};--la-bg2:{bg2};--la-card:{card};--la-border:{border};
  --la-text:{text_color};--la-text2:{text_sec};
  --la-acc:{acc};--la-acc2:{acc2};--la-pos:{pos};--la-neg:{t['negative']};--la-warn:{warn};
  --la-input:{inp};--la-sidebar:{sidebar_bg};
  --r-xs:4px;--r-sm:6px;--r-md:10px;--r-lg:14px;--r-xl:18px;--r-2xl:24px;--r-pill:999px;
  --ease:cubic-bezier(.4,0,.2,1);--ease-out:cubic-bezier(0,0,.2,1);
  --tf:.12s var(--ease);--tb:.18s var(--ease);--ts:.28s var(--ease);
  --font:'Plus Jakarta Sans',-apple-system,BlinkMacSystemFont,'SF Pro Display','Segoe UI Variable Display','Segoe UI',system-ui,sans-serif;
  --mono:'JetBrains Mono','Fira Code','Cascadia Code','SF Mono',Consolas,monospace;
  --sh-card:{shadow_card};--sh-hover:{shadow_hover};
}}

/* ── Reduce motion ── */
{motion_css}

/* ── Global ── */
.stApp{{background:var(--la-bg)!important;color:var(--la-text)!important;
  font-family:var(--font)!important;font-size:{base_font}px!important;
  font-feature-settings:"kern" 1,"liga" 1,"calt" 1!important;
  -webkit-font-smoothing:antialiased!important;-moz-osx-font-smoothing:grayscale!important;
  text-rendering:optimizeLegibility!important;}}

/* ── Universal text sharpening ── */
p,li,td,th,caption,figcaption,span,label,div,
.stMarkdown,.stMarkdown p,.stMarkdown li,.stMarkdown span,
[data-testid="stMarkdownContainer"] *,[data-testid="stText"],[data-testid="stCaptionContainer"],
.stRadio label,.stRadio div,.stCheckbox label,.stCheckbox div,
.stSelectbox label,.stMultiSelect label,.stTextInput label,.stTextArea label,
.stNumberInput label,.stDateInput label,.stSlider label,
.stFileUploader label,.stFileUploader span,.stExpander label,.stExpander p{{
  color:var(--la-text)!important;-webkit-font-smoothing:antialiased!important;
  -moz-osx-font-smoothing:grayscale!important;font-family:var(--font)!important;}}
.stCaption,[data-testid="stCaptionContainer"] p,small,.stHelp{{
  color:var(--la-text2)!important;-webkit-font-smoothing:antialiased!important;}}

/* ── Headings ── */
h1,h2,h3,h4,h5,h6,
.stMarkdown h1,.stMarkdown h2,.stMarkdown h3,.stMarkdown h4,.stMarkdown h5,.stMarkdown h6{{
  color:var(--la-text)!important;font-family:var(--font)!important;
  letter-spacing:-0.025em!important;line-height:1.2!important;
  -webkit-font-smoothing:antialiased!important;}}
.stMarkdown h1{{font-size:1.75rem!important;font-weight:800!important;}}
.stMarkdown h2{{font-size:1.35rem!important;font-weight:700!important;}}
.stMarkdown h3{{font-size:1.1rem!important;font-weight:600!important;}}
.stMarkdown h4{{font-size:0.95rem!important;font-weight:600!important;}}
code,pre,.stMarkdown code,.stMarkdown pre{{
  font-family:var(--mono)!important;font-size:0.88em!important;
  -webkit-font-smoothing:antialiased!important;}}

/* ── Hero ── */
.hero,.hero *,.hero h1,.hero h2,.hero p,.hero span,.hero label{{
  color:#fff!important;-webkit-font-smoothing:antialiased!important;}}
.hero{{background:{grad};padding:2.4rem 2.6rem 2.2rem;border-radius:var(--r-xl);
  margin-bottom:1.8rem;border:1px solid rgba(255,255,255,0.08);
  position:relative;overflow:hidden;}}
.hero::before{{content:'';position:absolute;inset:0;
  background:repeating-linear-gradient(-45deg,transparent,transparent 40px,
  rgba(255,255,255,0.015) 40px,rgba(255,255,255,0.015) 41px);pointer-events:none;}}
.hero h1{{font-size:2.15rem!important;font-weight:800!important;
  letter-spacing:-0.035em!important;margin:0!important;line-height:1.1!important;}}
.hero p{{font-size:1rem!important;opacity:.85;margin:.55rem 0 0 0!important;
  font-weight:400!important;line-height:1.55!important;letter-spacing:.005em!important;}}

/* ── Page sub-header ── */
.page-header,.page-header *,.page-header h2,.page-header p{{
  color:#fff!important;-webkit-font-smoothing:antialiased!important;}}
.page-header{{background:{grad};padding:1.3rem 1.8rem 1.2rem;border-radius:var(--r-lg);
  margin-bottom:1.5rem;border:1px solid rgba(255,255,255,0.07);}}
.page-header h2{{margin:0!important;font-size:1.4rem!important;font-weight:700!important;
  letter-spacing:-0.025em!important;}}
.page-header p{{margin:.3rem 0 0 0!important;opacity:.82;font-size:.9rem!important;
  letter-spacing:.005em!important;font-weight:400!important;}}

/* ── Sidebar ── */
section[data-testid="stSidebar"]{{
  background:{sidebar_bg}!important;border-right:1px solid {sb_line}!important;}}
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] h1,section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,section[data-testid="stSidebar"] h4,
section[data-testid="stSidebar"] li,
section[data-testid="stSidebar"] .stMarkdown,
section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p,
section[data-testid="stSidebar"] [data-testid="stText"],
section[data-testid="stSidebar"] .stRadio label,
section[data-testid="stSidebar"] .stCheckbox label,
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stSlider label{{
  color:{sb_text}!important;-webkit-font-smoothing:antialiased!important;}}
section[data-testid="stSidebar"] .stCaption,
section[data-testid="stSidebar"] small,
section[data-testid="stSidebar"] [data-testid="stCaptionContainer"] p{{
  color:{sb_text_2}!important;font-size:.75rem!important;
  -webkit-font-smoothing:antialiased!important;}}
section[data-testid="stSidebar"] hr{{
  border-color:{sb_line}!important;margin:.5rem 0!important;}}
section[data-testid="stSidebar"] .stTextInput input,
section[data-testid="stSidebar"] .stTextArea textarea,
section[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] *,
section[data-testid="stSidebar"] .stMultiSelect div[data-baseweb="select"] *,
section[data-testid="stSidebar"] .stNumberInput input{{
  color:{sb_input_tx}!important;background-color:{sb_input_bg}!important;
  -webkit-font-smoothing:antialiased!important;}}
section[data-testid="stSidebar"] [data-testid="stFileUploader"] section p,
section[data-testid="stSidebar"] [data-testid="stFileUploader"] section span,
section[data-testid="stSidebar"] [data-testid="stFileDropzoneInstructions"],
section[data-testid="stSidebar"] [data-testid="stFileDropzoneInstructions"] span{{
  color:{sb_input_tx}!important;}}
section[data-testid="stSidebar"] [data-testid="stFileUploader"] small,
section[data-testid="stSidebar"] .stCaption p{{color:{sb_text_2}!important;}}
section[data-testid="stSidebar"] .stButton>button{{
  background:{sb_hover_bg}!important;border:1px solid {sb_line}!important;
  color:{sb_text}!important;border-radius:var(--r-md)!important;
  font-weight:500!important;transition:var(--tb)!important;}}
section[data-testid="stSidebar"] .stButton>button:hover{{
  background:{acc}22!important;border-color:{acc}55!important;}}

/* ── Buttons ── */
.stButton>button{{font-family:var(--font)!important;
  -webkit-font-smoothing:antialiased!important;border-radius:var(--r-md)!important;
  font-weight:500!important;letter-spacing:.005em!important;
  transition:var(--tb)!important;transform:translateZ(0);}}
.stButton>button[kind="primary"],
.stButton>button[data-testid="baseButton-primary"]{{
  background:{acc}!important;color:#fff!important;border:none!important;
  box-shadow:0 1px 3px {acc}55,0 2px 10px {acc}22!important;font-weight:600!important;}}
.stButton>button[kind="primary"]:hover,
.stButton>button[data-testid="baseButton-primary"]:hover{{
  filter:brightness(1.1)!important;
  box-shadow:0 2px 8px {acc}66,0 4px 18px {acc}33!important;
  transform:translateY(-1px)!important;}}
.stButton>button[kind="primary"]:active,
.stButton>button[data-testid="baseButton-primary"]:active{{
  filter:brightness(.97)!important;transform:translateY(0)!important;
  box-shadow:0 1px 3px {acc}44!important;}}
.stButton>button[kind="secondary"],
.stButton>button[data-testid="baseButton-secondary"]{{
  background:transparent!important;color:var(--la-text)!important;
  border:1px solid var(--la-border)!important;font-weight:500!important;}}
.stButton>button[kind="secondary"]:hover,
.stButton>button[data-testid="baseButton-secondary"]:hover{{
  background:var(--la-bg2)!important;border-color:{acc}66!important;
  color:var(--la-acc)!important;}}
.stButton>button:not([kind]){{
  background:transparent!important;color:var(--la-text)!important;
  border:1px solid var(--la-border)!important;}}
.stButton>button:focus:not(:active){{
  box-shadow:0 0 0 3px {acc}33!important;outline:none!important;}}
.stDownloadButton>button{{
  background:var(--la-bg2)!important;color:var(--la-text)!important;
  border:1px solid var(--la-border)!important;border-radius:var(--r-md)!important;
  font-family:var(--font)!important;font-weight:500!important;
  transition:var(--tb)!important;-webkit-font-smoothing:antialiased!important;}}
.stDownloadButton>button:hover{{
  border-color:var(--la-acc)!important;color:var(--la-acc)!important;}}

/* ── Inputs ── */
.stTextInput input,.stTextArea textarea,.stNumberInput input,.stDateInput input{{
  background-color:var(--la-input)!important;color:var(--la-text)!important;
  border:1px solid var(--la-border)!important;border-radius:var(--r-md)!important;
  font-family:var(--font)!important;font-size:{input_font}px!important;
  -webkit-font-smoothing:antialiased!important;
  transition:border-color var(--tf),box-shadow var(--tf)!important;
  padding:.45rem .8rem!important;}}
.stTextInput input:focus,.stTextArea textarea:focus,.stNumberInput input:focus{{
  border-color:var(--la-acc)!important;
  box-shadow:0 0 0 3px {acc}25!important;outline:none!important;}}
.stTextInput>label,.stTextArea>label,.stSelectbox>label,
.stNumberInput>label,.stMultiSelect>label,.stDateInput>label,
.stSlider>label,.stFileUploader>label{{
  color:var(--la-text)!important;font-weight:500!important;
  font-size:.86rem!important;letter-spacing:.01em!important;
  -webkit-font-smoothing:antialiased!important;}}
.stTextInput input::placeholder,.stTextArea textarea::placeholder{{
  color:{ph_col}!important;opacity:1!important;font-style:italic!important;}}
section[data-testid="stSidebar"] .stTextInput input::placeholder,
section[data-testid="stSidebar"] .stTextArea textarea::placeholder{{
  color:{ph_col}!important;opacity:1!important;font-style:italic!important;}}
.stSelectbox div[data-baseweb="select"] *,
.stMultiSelect div[data-baseweb="select"] *{{
  background-color:var(--la-input)!important;color:var(--la-text)!important;
  border-color:var(--la-border)!important;border-radius:var(--r-md)!important;
  font-family:var(--font)!important;-webkit-font-smoothing:antialiased!important;}}

/* ── File uploader ── */
[data-testid="stFileUploader"] section,
[data-testid="stFileUploader"] section p,
[data-testid="stFileUploader"] section span,
[data-testid="stFileDropzoneInstructions"],
[data-testid="stFileDropzoneInstructions"] span,
[data-testid="stFileDropzoneInstructions"] small{{
  color:var(--la-text)!important;-webkit-font-smoothing:antialiased!important;}}
[data-testid="stFileUploader"] small,[data-testid="stFileUploader"] .stCaption{{
  color:var(--la-text2)!important;font-size:.8rem!important;}}

/* ── Stat cards ── */
.stat-card{{background:var(--la-card);border:1px solid var(--la-border);
  border-radius:var(--r-lg);padding:1.2rem 1rem 1.1rem;text-align:center;
  box-shadow:var(--sh-card);
  transition:box-shadow var(--tb),transform var(--tb);
  position:relative;overflow:hidden;}}
.stat-card::after{{content:'';position:absolute;top:0;left:10%;right:10%;height:2px;
  background:linear-gradient(90deg,transparent,{acc}88,transparent);
  border-radius:0 0 2px 2px;}}
.stat-card .stat-value{{font-size:1.95rem!important;font-weight:800!important;
  color:var(--la-acc)!important;line-height:1!important;letter-spacing:-0.03em!important;
  -webkit-font-smoothing:antialiased!important;font-family:var(--font)!important;}}
.stat-card .stat-label{{font-size:.72rem!important;font-weight:600!important;
  color:var(--la-text2)!important;margin-top:.45rem!important;
  text-transform:uppercase!important;letter-spacing:.08em!important;
  -webkit-font-smoothing:antialiased!important;}}

/* ── Custom cards ── */
.custom-card{{background:var(--la-card);border:1px solid var(--la-border);
  border-radius:var(--r-lg);padding:1.1rem 1.35rem;margin-bottom:.7rem;
  box-shadow:var(--sh-card);
  transition:box-shadow var(--tb),border-color var(--tb),transform var(--tb);
  position:relative;}}
.custom-card::before{{content:'';position:absolute;left:0;top:20%;bottom:20%;
  width:3px;border-radius:0 2px 2px 0;background:var(--la-acc);opacity:.8;}}
.custom-card:hover{{box-shadow:var(--sh-hover)!important;
  border-color:{acc}44!important;transform:translateY(-1px);}}
.custom-card h4{{margin:0 0 .3rem 0!important;color:var(--la-text)!important;
  font-size:.93rem!important;font-weight:600!important;letter-spacing:-.01em!important;
  -webkit-font-smoothing:antialiased!important;}}
.custom-card p,.custom-card span{{color:var(--la-text)!important;}}

/* ── History / tool / login cards ── */
.history-item{{background:var(--la-card);border:1px solid var(--la-border);
  border-radius:var(--r-md);padding:.7rem 1rem;margin-bottom:.4rem;cursor:pointer;
  transition:border-color var(--tb),background var(--tb);}}
.history-item:hover{{border-color:{acc}66!important;background:var(--la-bg2)!important;}}
.tool-card{{background:var(--la-card);border:1px solid var(--la-border);
  border-radius:var(--r-md);padding:.9rem 1.1rem;margin-bottom:.5rem;}}
.login-card{{background:var(--la-card);border:1px solid var(--la-border);
  border-top:3px solid var(--la-acc);border-radius:var(--r-xl);
  padding:2.2rem 2.4rem;box-shadow:var(--sh-card);}}

/* ── AI response box ── */
.response-box{{background:var(--la-card);border:1px solid var(--la-border);
  border-radius:var(--r-lg);padding:1.8rem 2rem;
  line-height:1.78!important;font-size:{input_font}px!important;
  font-family:var(--font)!important;white-space:pre-wrap;
  color:var(--la-text)!important;-webkit-font-smoothing:antialiased!important;
  box-shadow:var(--sh-card);position:relative;}}
.response-box::before{{content:'';position:absolute;top:0;left:0;right:0;height:2px;
  background:linear-gradient(90deg,var(--la-acc),var(--la-acc2));
  border-radius:var(--r-lg) var(--r-lg) 0 0;}}

/* ── Disclaimer ── */
.disclaimer{{background:{disc_bg};border-left:3px solid {warn};
  border-radius:0 var(--r-md) var(--r-md) 0;padding:.85rem 1.1rem;margin-top:1rem;
  font-size:.83rem!important;color:{disc_tx}!important;
  line-height:1.55!important;-webkit-font-smoothing:antialiased!important;}}

/* ── Badges ── */
.badge{{display:inline-flex;align-items:center;padding:.2rem .6rem;
  border-radius:var(--r-sm);font-size:.71rem!important;font-weight:600!important;
  letter-spacing:.04em!important;text-transform:uppercase!important;
  -webkit-font-smoothing:antialiased!important;font-family:var(--font)!important;
  white-space:nowrap;}}
.badge-ok{{background:{badge_ok_bg};color:{badge_ok_tx}!important;}}
.badge-warn{{background:{badge_warn_bg};color:{badge_warn_tx}!important;}}
.badge-err{{background:{badge_err_bg};color:{badge_err_tx}!important;}}
.badge-info{{background:{badge_inf_bg};color:{badge_inf_tx}!important;}}

/* ── Tabs ── */
div[data-testid="stTabs"] button{{font-family:var(--font)!important;
  font-weight:500!important;font-size:.86rem!important;
  color:var(--la-text2)!important;background:transparent!important;
  border:none!important;border-bottom:2px solid transparent!important;
  border-radius:0!important;padding:.55rem .9rem!important;
  transition:color var(--tb),border-color var(--tb)!important;
  -webkit-font-smoothing:antialiased!important;letter-spacing:.005em!important;}}
div[data-testid="stTabs"] button:hover{{color:var(--la-acc)!important;
  background:{acc}0d!important;border-radius:var(--r-sm) var(--r-sm) 0 0!important;}}
div[data-testid="stTabs"] button[aria-selected="true"]{{color:var(--la-acc)!important;
  font-weight:600!important;background:transparent!important;
  border-bottom:2px solid var(--la-acc)!important;}}

/* ── Expanders ── */
.streamlit-expanderHeader,[data-testid="stExpander"] summary{{
  background:var(--la-bg2)!important;color:var(--la-text)!important;
  border-radius:var(--r-md)!important;border:1px solid var(--la-border)!important;
  font-weight:500!important;
  transition:background var(--tb),border-color var(--tb)!important;
  -webkit-font-smoothing:antialiased!important;}}
.streamlit-expanderHeader:hover,[data-testid="stExpander"] summary:hover{{
  background:var(--la-card)!important;border-color:{acc}55!important;}}
.streamlit-expanderContent,[data-testid="stExpander"] details{{
  background:var(--la-card)!important;border:1px solid var(--la-border)!important;
  border-top:none!important;border-radius:0 0 var(--r-md) var(--r-md)!important;}}

/* ── Metrics ── */
[data-testid="stMetric"]{{background:var(--la-card)!important;
  border:1px solid var(--la-border)!important;border-radius:var(--r-lg)!important;
  padding:.9rem 1.1rem!important;box-shadow:var(--sh-card)!important;}}
[data-testid="stMetricLabel"] p,div[data-testid="metric-container"] label{{
  color:var(--la-text2)!important;font-size:.74rem!important;font-weight:600!important;
  text-transform:uppercase!important;letter-spacing:.07em!important;
  -webkit-font-smoothing:antialiased!important;}}
[data-testid="stMetricValue"],
div[data-testid="metric-container"] div[data-testid="stMetricValue"]{{
  color:var(--la-acc)!important;font-weight:700!important;
  letter-spacing:-.02em!important;-webkit-font-smoothing:antialiased!important;}}
[data-testid="stMetricDelta"]{{color:var(--la-pos)!important;}}

/* ── Radio & Checkbox ── */
.stRadio>div>label,.stCheckbox>label,
.stRadio [data-testid="stMarkdownContainer"] p,
.stCheckbox [data-testid="stMarkdownContainer"] p{{
  color:var(--la-text)!important;-webkit-font-smoothing:antialiased!important;}}

/* ── Dropdowns ── */
[data-baseweb="menu"],[data-baseweb="menu"] li,[data-baseweb="popover"] li{{
  background-color:var(--la-card)!important;color:var(--la-text)!important;
  font-family:var(--font)!important;-webkit-font-smoothing:antialiased!important;}}
[data-baseweb="menu"] li:hover,[data-baseweb="popover"] li:hover{{
  background-color:{acc}18!important;color:var(--la-acc)!important;}}

/* ── Progress ── */
.stProgress>div>div{{background-color:var(--la-acc)!important;}}
.stProgress{{background-color:var(--la-bg2)!important;border-radius:var(--r-pill)!important;}}

/* ── Chat ── */
[data-testid="stChatMessage"]{{background:var(--la-card)!important;
  border:1px solid var(--la-border)!important;border-radius:var(--r-lg)!important;
  margin-bottom:.6rem!important;}}
[data-testid="stChatMessage"] p{{color:var(--la-text)!important;
  -webkit-font-smoothing:antialiased!important;}}

/* ── Dataframes ── */
.stDataFrame{{border:1px solid var(--la-border)!important;
  border-radius:var(--r-md)!important;overflow:hidden!important;}}
.stDataFrame th{{background:var(--la-bg2)!important;color:var(--la-text)!important;
  font-weight:600!important;font-size:.82rem!important;letter-spacing:.03em!important;
  -webkit-font-smoothing:antialiased!important;}}
.stDataFrame td{{color:var(--la-text)!important;-webkit-font-smoothing:antialiased!important;}}

/* ── Alerts ── */
[data-testid="stAlert"] p,[data-testid="stInfo"] p,[data-testid="stSuccess"] p,
[data-testid="stWarning"] p,[data-testid="stError"] p{{
  color:inherit!important;-webkit-font-smoothing:antialiased!important;}}

/* ── Scrollbar ── */
::-webkit-scrollbar{{width:5px;height:5px;}}
::-webkit-scrollbar-track{{background:transparent;}}
::-webkit-scrollbar-thumb{{background:{border};border-radius:var(--r-pill);
  transition:background var(--tb);}}
::-webkit-scrollbar-thumb:hover{{background:{acc}88;}}
::selection{{background:{acc}30;color:var(--la-text);}}

/* ── Reduce-motion (prefers) ── */
@media (prefers-reduced-motion:reduce){{
  *,*::before,*::after{{animation-duration:.01ms!important;transition-duration:.01ms!important;}}}}

/* ── Mobile ── */
@media (max-width:768px){{
  .stApp{{font-size:{mobile_font}px!important;}}
  .hero{{padding:1.4rem 1.3rem!important;border-radius:var(--r-lg)!important;}}
  .hero h1{{font-size:1.4rem!important;}}
  .hero p{{font-size:.88rem!important;}}
  .page-header{{padding:1rem 1.2rem!important;}}
  .page-header h2{{font-size:1.1rem!important;}}
  .stat-card .stat-value{{font-size:1.45rem!important;}}
  .custom-card{{padding:.75rem .9rem!important;}}
  .response-box{{padding:1rem!important;font-size:.88rem!important;}}
  .stButton>button{{width:100%!important;min-height:2.4rem!important;}}
  div[data-testid="stTabs"] button{{font-size:.72rem!important;padding:.35rem .5rem!important;}}
  .login-card{{padding:1.3rem 1rem!important;}}
  [data-testid="stMetric"]{{padding:.7rem .9rem!important;}}}}
@media (max-width:480px){{
  .hero h1{{font-size:1.15rem!important;}}
  .stat-card .stat-value{{font-size:1.15rem!important;}}
  .badge{{font-size:.64rem!important;}}}}
</style>"""

# ═══════════════════════════════════════════════════════
# EXPORT FUNCTIONS (WITH FIRM BRANDING)
# ═══════════════════════════════════════════════════════
def export_pdf(text: str, title: str = "LexiAssist Analysis") -> bytes:
    if not HAS_FPDF:
        return b"%PDF-1.0\nPDF generation unavailable. Install fpdf2."
    firm = get_firm_name()
    pdf  = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, txt=title, ln=True, align="C")
    pdf.ln(2)
    if firm and firm != "LexiAssist":
        pdf.set_font("Helvetica","B",11); pdf.cell(0,7,txt=firm,ln=True,align="C")
    pdf.set_font("Helvetica","I",9)
    pdf.cell(0,6,txt=f"Generated: {datetime.now():%d %B %Y at %H:%M}",ln=True,align="C")
    pdf.ln(6); pdf.set_draw_color(100,100,100); pdf.line(15,pdf.get_y(),195,pdf.get_y()); pdf.ln(6)
    pdf.set_font("Helvetica",size=10)
    clean = text.encode("latin-1",errors="replace").decode("latin-1")
    for line in clean.split("\n"):
        pdf.multi_cell(0,6,txt=line); pdf.ln(1)
    pdf.ln(8); pdf.set_font("Helvetica","I",8)
    pdf.cell(0,5,txt=f"Generated by {firm} via LexiAssist v8.0 — Verify all citations independently",ln=True,align="C")
    raw = pdf.output(dest="S")
    if isinstance(raw,str):    return raw.encode("latin-1",errors="replace")
    if isinstance(raw,bytearray): return bytes(raw)
    return raw

# ── DOCX helper functions ────────────────────────────────────────────────────
def _docx_shade_paragraph(para, fill_hex: str) -> None:
    pPr = para._p.get_or_add_pPr()
    for x in pPr.findall(_docx_qn("w:shd")): pPr.remove(x)
    shd = _OxmlElement("w:shd")
    shd.set(_docx_qn("w:val"),  "clear"); shd.set(_docx_qn("w:color"), "auto")
    shd.set(_docx_qn("w:fill"), fill_hex.lstrip("#")); pPr.append(shd)

def _docx_set_cell_bg(cell, fill_hex: str) -> None:
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for x in tcPr.findall(_docx_qn("w:shd")): tcPr.remove(x)
    shd = _OxmlElement("w:shd")
    shd.set(_docx_qn("w:val"),  "clear"); shd.set(_docx_qn("w:color"), "auto")
    shd.set(_docx_qn("w:fill"), fill_hex.lstrip("#")); tcPr.append(shd)

def _docx_add_footer(doc, firm: str) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    footer = section.footer
    para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear(); para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr  = para._p.get_or_add_pPr()
    pBdr = _OxmlElement("w:pBdr"); top = _OxmlElement("w:top")
    top.set(_docx_qn("w:val"),  "single"); top.set(_docx_qn("w:sz"),    "4")
    top.set(_docx_qn("w:space"),"1");      top.set(_docx_qn("w:color"), "1a2e4a")
    pBdr.append(top); pPr.append(pBdr)
    tabs_el = _OxmlElement("w:tabs"); tab = _OxmlElement("w:tab")
    tab.set(_docx_qn("w:val"),"right"); tab.set(_docx_qn("w:pos"),"9026")
    tabs_el.append(tab); pPr.append(tabs_el)
    grey = RGBColor(0x6b,0x72,0x80)
    ft   = firm if firm and firm != "LexiAssist" else "LexiAssist v8.0"
    disc = f"{ft}  \u00b7  LexiAssist v8.0  \u00b7  Verify all citations independently"
    def _sr(t, bold=False):
        r=para.add_run(t); r.font.name="Calibri"; r.font.size=Pt(8); r.font.bold=bold; r.font.color.rgb=grey
    def _fr(instr):
        r=para.add_run(); r.font.name="Calibri"; r.font.size=Pt(8); r.font.color.rgb=grey
        b=_OxmlElement("w:fldChar"); b.set(_docx_qn("w:fldCharType"),"begin")
        it=_OxmlElement("w:instrText")
        it.set("{http://www.w3.org/XML/1998/namespace}space","preserve"); it.text=f" {instr} "
        e=_OxmlElement("w:fldChar"); e.set(_docx_qn("w:fldCharType"),"end")
        r._r.extend([b,it,e])
    _sr(disc); _sr("\t"); _sr("Page "); _fr("PAGE"); _sr(" of "); _fr("NUMPAGES")

def _docx_parse_output(text: str) -> list:
    blocks = []
    for line in text.split("\n"):
        s = line.strip()
        if not s: continue
        if "\u2550\u2550\u2550" in s:
            inner = s.replace("\u2550","").strip()
            if inner: blocks.append({"type":"heading2","content":inner})
            continue
        if len(s)>3 and all(c in "\u2550\u2500\u2501\u2014=-*" for c in s): continue
        if s.startswith("\u25b8"):
            blocks.append({"type":"subheading","content":s[1:].strip()}); continue
        if s[:1] in ("\U0001f534","\U0001f7e1","\U0001f7e2"):
            lvl = "HIGH" if "\U0001f534" in s[:4] else ("MEDIUM" if "\U0001f7e1" in s[:4] else "LOW")
            body = s
            for pfx in ("\U0001f534 HIGH RISK \u2192","\U0001f7e1 MEDIUM RISK \u2192","\U0001f7e2 LOW RISK \u2192",
                        "\U0001f534","\U0001f7e1","\U0001f7e2"):
                if body.startswith(pfx): body=body[len(pfx):].strip(); break
            sep = " \u2014 " if " \u2014 " in body else (" - " if " - " in body else None)
            party,reason = body.split(sep,1) if sep else (body,"")
            blocks.append({"type":"risk_row","level":lvl,"party":party.strip(),"reason":reason.strip()})
            continue
        if s.startswith(("\u2022 ","\u25aa ","\u00b7 ")):
            blocks.append({"type":"bullet","content":s[2:].strip()}); continue
        if s.startswith(("  \u2022","  -","\t\u2022","\t-")):
            blocks.append({"type":"bullet","content":s.strip().lstrip("\u2022-").strip()}); continue
        if s.endswith(":") and 5<len(s)<60 and not any(c in s for c in ".?!"):
            blocks.append({"type":"subheading","content":s}); continue
        blocks.append({"type":"body","content":s})
    return blocks

def export_docx(text: str, title: str = "LexiAssist Analysis",
                doc_type: str = "general", meta: dict = None) -> bytes:
    """Professional DOCX: letterhead, risk tables, per-type preambles."""
    if not HAS_DOCX:
        return b"DOCX generation unavailable - install python-docx."
    meta=meta or {}
    NAVY=RGBColor(0x1a,0x2e,0x4a); GOLD=RGBColor(0xc9,0xa8,0x4c)
    DARK_GREY=RGBColor(0x37,0x41,0x51); MID_GREY=RGBColor(0x6b,0x72,0x80)
    LIGHT_BLU=RGBColor(0xa0,0xbc,0xd8)
    firm=get_firm_name() or "LexiAssist"; date_str=datetime.now().strftime("%d %B %Y")
    bio=BytesIO(); doc=DocxDocument()
    sec=doc.sections[0]
    sec.page_width=int(8.27*914400); sec.page_height=int(11.69*914400)
    sec.left_margin=Inches(1.0); sec.right_margin=Inches(1.0)
    sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.9)
    def _sb(style,sz,bold=False,col=None,sb=0,sa=6):
        style.font.name="Calibri"; style.font.size=Pt(sz); style.font.bold=bold
        if col: style.font.color.rgb=col
        style.paragraph_format.space_before=Pt(sb)
        style.paragraph_format.space_after=Pt(sa)
        style.paragraph_format.line_spacing=Pt(sz*1.3)
    _sb(doc.styles["Normal"],11,col=DARK_GREY,sa=6)
    _sb(doc.styles["Heading 1"],18,True,NAVY,8,4)
    _sb(doc.styles["Heading 2"],13,True,NAVY,14,4)
    _sb(doc.styles["Heading 3"],11,True,DARK_GREY,10,2)
    doc.styles["Heading 1"].paragraph_format.keep_with_next=True
    doc.styles["Heading 2"].paragraph_format.keep_with_next=True
    try: _sb(doc.styles["List Bullet"],11,col=DARK_GREY,sa=3)
    except: pass
    hdr=doc.sections[0].header
    hp=hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _docx_shade_paragraph(hp,"1a2e4a")
    hp.paragraph_format.space_before=Pt(5); hp.paragraph_format.space_after=Pt(5)
    rf=hp.add_run(firm.upper() if firm!="LexiAssist" else "LEXIASSIST")
    rf.font.name="Calibri"; rf.font.size=Pt(12); rf.font.bold=True; rf.font.color.rgb=GOLD
    rt=hp.add_run(f"   \u00b7   Legal Analysis   \u00b7   Confidential   \u00b7   {date_str}")
    rt.font.name="Calibri"; rt.font.size=Pt(9); rt.font.color.rgb=LIGHT_BLU
    _docx_add_footer(doc,firm)
    doc.add_paragraph(title,style="Heading 1")
    mp=doc.add_paragraph(); mp.paragraph_format.space_after=Pt(14)
    rm=mp.add_run(f"{firm}   \u00b7   {date_str}   \u00b7   Generated by LexiAssist v8.0")
    rm.font.name="Calibri"; rm.font.size=Pt(9); rm.font.color.rgb=MID_GREY
    pPr=mp._p.get_or_add_pPr(); pBdr=_OxmlElement("w:pBdr"); btm=_OxmlElement("w:bottom")
    btm.set(_docx_qn("w:val"),"single"); btm.set(_docx_qn("w:sz"),"6")
    btm.set(_docx_qn("w:space"),"1"); btm.set(_docx_qn("w:color"),"1a2e4a")
    pBdr.append(btm); pPr.append(pBdr)
    def _pr(label,value):
        if not value: return
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2)
        rl=p.add_run(f"{label}:  ")
        rl.font.name="Calibri"; rl.font.size=Pt(10); rl.font.bold=True; rl.font.color.rgb=NAVY
        rv=p.add_run(value)
        rv.font.name="Calibri"; rv.font.size=Pt(10); rv.font.color.rgb=DARK_GREY
    def _sp(): s=doc.add_paragraph(); s.paragraph_format.space_after=Pt(6)
    if doc_type=="pleading":
        for line,sz,bold in [
            (meta.get("court","IN THE FEDERAL HIGH COURT OF NIGERIA"),11,True),
            (meta.get("division",""),10,False),
            ("HOLDEN AT "+meta.get("location","ABUJA"),10,False),
            ("",9,False),("SUIT NO: "+meta.get("suit_no","_______________"),10,True),
            ("",9,False),("BETWEEN",10,True),
            (meta.get("claimant","_______________"),11,True),
            ("(CLAIMANT / APPELLANT)",9,False),("AND",10,True),
            (meta.get("defendant","_______________"),11,True),
            ("(DEFENDANT / RESPONDENT)",9,False),
        ]:
            if not line: _sp(); continue
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
            r=p.add_run(line); r.font.name="Calibri"; r.font.size=Pt(sz)
            r.font.bold=bold; r.font.color.rgb=NAVY
        _sp(); _sp()
    elif doc_type=="research":
        doc.add_paragraph("LEGAL RESEARCH MEMORANDUM",style="Heading 2")
        _pr("Prepared by",firm); _pr("Prepared for",meta.get("prepared_for",""))
        _pr("Date",date_str); _pr("Subject",title)
        _pr("Area of law",meta.get("area","")); _pr("Jurisdiction",meta.get("jurisdiction","Nigeria")); _sp()
    elif doc_type in ("invoice","fee_note"):
        doc.add_paragraph("PROFESSIONAL FEE NOTE / INVOICE",style="Heading 2")
        _pr("Invoice No.",meta.get("invoice_no","")); _pr("Date",date_str)
        _pr("Client",meta.get("client","")); _pr("Matter",meta.get("matter",""))
        _pr("Amount",meta.get("amount","")); _pr("Due Date",meta.get("due_date","")); _sp()
    elif doc_type=="witness":
        doc.add_paragraph("WITNESS PREPARATION BRIEF",style="Heading 2")
        _pr("Witness",meta.get("witness","")); _pr("Role",meta.get("role",""))
        _pr("Matter",meta.get("matter","")); _pr("Prepared by",firm); _pr("Date",date_str)
        p=doc.add_paragraph()
        r=p.add_run("STRICTLY CONFIDENTIAL \u2014 Attorney-Client Privilege. Not for disclosure to opposing counsel.")
        r.font.name="Calibri"; r.font.size=Pt(9); r.font.bold=True
        r.font.color.rgb=RGBColor(0x99,0x1b,0x1b); _sp()
    elif doc_type=="settlement":
        doc.add_paragraph("SETTLEMENT STRATEGY BRIEF",style="Heading 2")
        _pr("Matter",meta.get("matter","")); _pr("Prepared by",firm); _pr("Date",date_str)
        p=doc.add_paragraph(); r=p.add_run("CONFIDENTIAL \u2014 Without Prejudice")
        r.font.name="Calibri"; r.font.size=Pt(10); r.font.bold=True
        r.font.color.rgb=RGBColor(0x92,0x40,0x0e); _sp()
    elif doc_type=="due_diligence":
        doc.add_paragraph("DUE DILIGENCE REPORT",style="Heading 2")
        _pr("Subject",meta.get("subject","")); _pr("Prepared by",firm)
        _pr("Date",date_str); _pr("Classification","STRICTLY CONFIDENTIAL"); _sp()
    blocks=_docx_parse_output(text); risk_buf=[]
    RCFG={"HIGH":  {"fill":"fef2f2","text":RGBColor(0x7f,0x1d,0x1d)},
          "MEDIUM":{"fill":"fefce8","text":RGBColor(0x71,0x3f,0x12)},
          "LOW":   {"fill":"f0fdf4","text":RGBColor(0x14,0x53,0x2d)}}
    def _flush():
        nonlocal risk_buf
        if not risk_buf: return
        cw=[int(1.7*914400),int(1.1*914400),int(3.47*914400)]
        tbl=doc.add_table(rows=1,cols=3); tbl.style="Table Grid"; tbl.autofit=False
        for i,cell in enumerate(tbl.rows[0].cells): cell.width=cw[i]
        for cell,lbl in zip(tbl.rows[0].cells,["PARTY","RISK","EXPOSURE / REASON"]):
            _docx_set_cell_bg(cell,"1a2e4a"); p=cell.paragraphs[0]; p.clear()
            r=p.add_run(lbl); r.font.name="Calibri"; r.font.size=Pt(9)
            r.font.bold=True; r.font.color.rgb=GOLD
            p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        for rb in risk_buf:
            lvl=rb.get("level","LOW"); cfg=RCFG.get(lvl,RCFG["LOW"]); row=tbl.add_row()
            for i,cell in enumerate(row.cells):
                cell.width=cw[i]; _docx_set_cell_bg(cell,cfg["fill"])
                p=cell.paragraphs[0]; p.clear()
                p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
                r=p.add_run(rb.get("party","") if i==0 else lvl if i==1 else rb.get("reason",""))
                r.font.name="Calibri"; r.font.size=Pt(10 if i!=1 else 9)
                r.font.bold=(i==0 or i==1); r.font.color.rgb=cfg["text"]
        s=doc.add_paragraph(); s.paragraph_format.space_after=Pt(4); risk_buf=[]
    for block in blocks:
        bt=block["type"]
        if bt=="risk_row": risk_buf.append(block); continue
        else: _flush()
        if bt=="heading2":
            doc.add_paragraph(block["content"],style="Heading 2")
        elif bt=="subheading":
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
            r=p.add_run(block["content"])
            r.font.name="Calibri"; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=NAVY
        elif bt=="bullet":
            try:
                p=doc.add_paragraph(block["content"],style="List Bullet")
                p.paragraph_format.space_after=Pt(3)
                for run in p.runs: run.font.name="Calibri"; run.font.size=Pt(11); run.font.color.rgb=DARK_GREY
            except:
                p=doc.add_paragraph(); r=p.add_run(f"\u2022  {block['content']}")
                r.font.name="Calibri"; r.font.size=Pt(11); r.font.color.rgb=DARK_GREY
        elif bt=="body":
            p=doc.add_paragraph(block["content"]); p.style=doc.styles["Normal"]
    _flush()
    doc.add_paragraph()
    disc=doc.add_paragraph()
    pPr2=disc._p.get_or_add_pPr(); pBdr2=_OxmlElement("w:pBdr"); dt=_OxmlElement("w:top")
    dt.set(_docx_qn("w:val"),"single"); dt.set(_docx_qn("w:sz"),"4")
    dt.set(_docx_qn("w:space"),"1"); dt.set(_docx_qn("w:color"),"cccccc")
    pBdr2.append(dt); pPr2.append(pBdr2)
    rd=disc.add_run(f"\u26a0  AI-generated via LexiAssist v8.0. Not legal advice. "
        f"Verify all citations independently. \u00a9 {datetime.now().year} {firm}.")
    rd.font.name="Calibri"; rd.font.size=Pt(8); rd.font.color.rgb=MID_GREY
    doc.save(bio); return bio.getvalue()

def export_txt(text: str, title: str = "LexiAssist Analysis") -> str:
    firm   = get_firm_name()
    header = f"{'='*60}\n{title}\n{firm}\nGenerated: {datetime.now():%d %B %Y at %H:%M}\n{'='*60}\n\n"
    footer = f"\n\n{'='*60}\nGenerated by {firm} via LexiAssist v8.0\n{'='*60}"
    return header + text + footer

def export_html(text: str, title: str = "LexiAssist Analysis") -> str:
    firm = get_firm_name()
    import html as html_mod
    safe = html_mod.escape(text).replace("\n","<br>")
    return (f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
            f"<title>{html_mod.escape(title)}</title>"
            f"<style>body{{font-family:Calibri,Arial,sans-serif;max-width:900px;margin:2rem auto;"
            f"padding:0 1.5rem;color:#1a2e4a;}} h1{{color:#1a2e4a;border-bottom:3px solid #c9a84c;"
            f"padding-bottom:.5rem;}} .meta{{color:#6b7280;font-size:.9rem;margin-bottom:1.5rem;}}"
            f".body{{line-height:1.75;white-space:pre-wrap;}}"
            f".footer{{margin-top:2rem;padding-top:1rem;border-top:1px solid #e2e8f0;"
            f"color:#9ca3af;font-size:.8rem;}}</style></head>"
            f"<body><h1>{html_mod.escape(title)}</h1>"
            f"<div class='meta'>{html_mod.escape(firm)} &middot; "
            f"Generated {datetime.now():%d %B %Y} &middot; LexiAssist v8.0</div>"
            f"<div class='body'>{safe}</div>"
            f"<div class='footer'>\u26a0 AI-generated. Not legal advice. "
            f"Verify all citations independently.</div></body></html>")

def safe_pdf_download(text: str, title: str, fname: str, key: str):
    try:
        pdf_data = export_pdf(text, title)
        st.download_button("📥 PDF", data=pdf_data, file_name=f"{fname}.pdf",
                           mime="application/pdf", key=key, width='stretch')
    except Exception as e:
        st.button("📥 PDF (unavailable)", disabled=True, key=key, width='stretch')
        logger.warning(f"PDF export failed: {e}")

def safe_docx_download(text: str, title: str, fname: str, key: str,
                        doc_type: str = "general", meta: dict = None):
    try:
        docx_data = export_docx(text, title, doc_type=doc_type, meta=meta or {})
        st.download_button("📥 DOCX", data=docx_data, file_name=f"{fname}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           key=key, width='stretch')
    except Exception as e:
        st.button("📥 DOCX (unavailable)", disabled=True, key=key, width='stretch')
        logger.warning(f"DOCX export failed: {e}")


# ═══════════════════════════════════════════════════════
# DATABASE LAYER
# ═══════════════════════════════════════════════════════
class Database:
    """PostgreSQL persistence for all LexiAssist data."""

    def __init__(self):
        self.url = _get_db_url()
        self.conn = self._connect()
        self._init_tables()

    def _connect(self):
        conn = psycopg2.connect(self.url)
        conn.autocommit = False
        return conn

    def _execute(self, sql: str, params=None):
        """Execute with auto-reconnect and transaction-error recovery."""
        try:
            cur = self.conn.cursor()
            cur.execute(sql, params or ())
            return cur
        except (psycopg2.OperationalError, psycopg2.InterfaceError):
            # Stale connection — reconnect and retry
            try:
                self.conn.rollback()
            except Exception:
                pass
            self.conn = self._connect()
            cur = self.conn.cursor()
            cur.execute(sql, params or ())
            return cur
        except psycopg2.Error:
            # Transaction aborted — roll back so the connection is usable again
            try:
                self.conn.rollback()
            except Exception:
                pass
            raise

    def _exec_ddl(self, sql: str):
        """Run a single DDL statement in its own isolated transaction.
        If it fails (e.g. object already exists in a different form), roll back
        cleanly so the connection stays usable for the next statement."""
        try:
            cur = self.conn.cursor()
            cur.execute(sql)
            self.conn.commit()
        except psycopg2.Error as e:
            try:
                self.conn.rollback()
            except Exception:
                pass
            logger.warning(f"DDL skipped (non-fatal): {e!s:.120}")

    def _init_tables(self):
        # Each statement runs in its own transaction so one failure never
        # poisons subsequent DDL (PostgreSQL aborts the whole txn on error).
        ddl_statements = [
            """CREATE TABLE IF NOT EXISTS kv_store (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL DEFAULT '[]'
            )""",
            """CREATE TABLE IF NOT EXISTS users (
                user_id TEXT PRIMARY KEY,
                username TEXT UNIQUE NOT NULL,
                email TEXT DEFAULT '',
                password_hash TEXT NOT NULL,
                firm_name TEXT DEFAULT '',
                lawyer_name TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                address TEXT DEFAULT '',
                role TEXT DEFAULT 'user',
                created_at TEXT DEFAULT '',
                last_login TEXT DEFAULT ''
            )""",
            """CREATE TABLE IF NOT EXISTS user_profile (
                id INTEGER PRIMARY KEY CHECK (id = 1),
                firm_name TEXT DEFAULT '',
                lawyer_name TEXT DEFAULT '',
                email TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                address TEXT DEFAULT '',
                password_hash TEXT DEFAULT ''
            )""",
            """CREATE TABLE IF NOT EXISTS cost_logs (
                id TEXT PRIMARY KEY,
                timestamp TEXT,
                model TEXT,
                task TEXT,
                mode TEXT,
                input_chars INTEGER DEFAULT 0,
                output_chars INTEGER DEFAULT 0,
                estimated_cost REAL DEFAULT 0,
                query_preview TEXT DEFAULT '',
                user_id TEXT DEFAULT 'legacy'
            )""",
            """CREATE TABLE IF NOT EXISTS case_analyses (
                id TEXT PRIMARY KEY,
                case_id TEXT NOT NULL,
                query TEXT,
                response TEXT,
                task TEXT,
                mode TEXT,
                timestamp TEXT,
                user_id TEXT DEFAULT 'legacy'
            )""",
            """CREATE TABLE IF NOT EXISTS user_sessions (
                token TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                created_at TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                last_used TEXT DEFAULT '',
                device_hint TEXT DEFAULT ''
            )""",
        ]
        for stmt in ddl_statements:
            self._exec_ddl(stmt)

        # Safely add columns to existing tables — each in its own transaction
        for tbl in ("cost_logs", "case_analyses"):
            self._exec_ddl(
                f"ALTER TABLE {tbl} ADD COLUMN IF NOT EXISTS user_id TEXT DEFAULT 'legacy'"
            )

        # Ensure legacy profile row exists
        self._exec_ddl(
            "INSERT INTO user_profile (id) VALUES (1) ON CONFLICT DO NOTHING"
        )

    def _uid(self) -> str:
        """Return current user_id from Streamlit session, fallback to 'legacy'."""
        try:
            uid = st.session_state.get("current_user_id", "")
            return uid if uid else "legacy"
        except Exception:
            return "legacy"

    # ── KV Store — raw (keep for internal use) ──
    def _save_list_raw(self, key: str, data: list):
        self._execute(
            "INSERT INTO kv_store (key, value) VALUES (%s, %s) "
            "ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value",
            (key, json.dumps(data, default=str)),
        )
        self.conn.commit()

    def _load_list_raw(self, key: str) -> list:
        cur = self._execute("SELECT value FROM kv_store WHERE key = %s", (key,))
        row = cur.fetchone()
        if row:
            try:
                return json.loads(row[0])
            except Exception:
                return []
        return []

    # ── KV Store — user-namespaced (primary API) ──
    def save_list(self, key: str, data: list):
        """Save data namespaced to the current user."""
        uid = self._uid()
        self._save_list_raw(f"u:{uid}:{key}", data)

    def load_list(self, key: str) -> list:
        """Load data namespaced to the current user."""
        uid = self._uid()
        return self._load_list_raw(f"u:{uid}:{key}")

    # ── User Profile ──
    def get_profile(self) -> dict:
        """Load current user's profile from users table + extended kv fields."""
        uid = self._uid()
        if uid and uid != "legacy":
            return self.get_user_profile(uid)
        # Fallback for legacy / unauthenticated
        cur = self._execute(
            "SELECT firm_name, lawyer_name, email, phone, address, password_hash "
            "FROM user_profile WHERE id = 1"
        )
        row = cur.fetchone()
        if row:
            return {
                "firm_name": row[0] or "", "lawyer_name": row[1] or "",
                "email": row[2] or "", "phone": row[3] or "",
                "address": row[4] or "", "password_hash": row[5] or "",
            }
        return {"firm_name": "", "lawyer_name": "", "email": "", "phone": "", "address": "", "password_hash": ""}

    def save_profile(self, profile: dict):
        """Save current user's profile."""
        uid = self._uid()
        if uid and uid != "legacy":
            self.save_user_profile(uid, profile)
        else:
            self._execute(
                "UPDATE user_profile SET firm_name=%s, lawyer_name=%s, email=%s, "
                "phone=%s, address=%s, password_hash=%s WHERE id=1",
                (
                    profile.get("firm_name", ""), profile.get("lawyer_name", ""),
                    profile.get("email", ""), profile.get("phone", ""),
                    profile.get("address", ""), profile.get("password_hash", ""),
                ),
            )
            self.conn.commit()

    # ── Users table CRUD ──
    def has_any_users(self) -> bool:
        cur = self._execute("SELECT COUNT(*) FROM users")
        return cur.fetchone()[0] > 0

    def create_user(self, data: dict) -> bool:
        try:
            self._execute(
                "INSERT INTO users (user_id, username, email, password_hash, firm_name, "
                "lawyer_name, phone, address, role, created_at) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                (
                    data.get("user_id", uuid.uuid4().hex[:12]),
                    data.get("username", "").lower().strip(),
                    data.get("email", ""),
                    data.get("password_hash", ""),
                    data.get("firm_name", ""),
                    data.get("lawyer_name", ""),
                    data.get("phone", ""),
                    data.get("address", ""),
                    data.get("role", "user"),
                    datetime.now().isoformat(),
                ),
            )
            self.conn.commit()
            return True
        except Exception as e:
            logger.error(f"create_user failed: {e}")
            try:
                self.conn.rollback()
            except Exception:
                pass
            return False

    def get_user_by_username(self, username: str) -> Optional[dict]:
        cur = self._execute(
            "SELECT user_id, username, email, password_hash, firm_name, lawyer_name, "
            "phone, address, role, created_at, last_login FROM users WHERE username = %s",
            (username.lower().strip(),),
        )
        row = cur.fetchone()
        if row:
            return {
                "user_id": row[0], "username": row[1], "email": row[2],
                "password_hash": row[3], "firm_name": row[4], "lawyer_name": row[5],
                "phone": row[6], "address": row[7], "role": row[8],
                "created_at": row[9], "last_login": row[10],
            }
        return None

    def get_user_by_id(self, user_id: str) -> Optional[dict]:
        cur = self._execute(
            "SELECT user_id, username, email, password_hash, firm_name, lawyer_name, "
            "phone, address, role, created_at, last_login FROM users WHERE user_id = %s",
            (user_id,),
        )
        row = cur.fetchone()
        if row:
            return {
                "user_id": row[0], "username": row[1], "email": row[2],
                "password_hash": row[3], "firm_name": row[4], "lawyer_name": row[5],
                "phone": row[6], "address": row[7], "role": row[8],
                "created_at": row[9], "last_login": row[10],
            }
        return None

    def list_users(self) -> list:
        cur = self._execute(
            "SELECT user_id, username, email, firm_name, lawyer_name, role, created_at, last_login "
            "FROM users ORDER BY created_at ASC"
        )
        rows = cur.fetchall()
        return [
            {
                "user_id": r[0], "username": r[1], "email": r[2],
                "firm_name": r[3], "lawyer_name": r[4], "role": r[5],
                "created_at": r[6], "last_login": r[7],
            }
            for r in rows
        ]

    def update_user(self, user_id: str, updates: dict):
        allowed = ("email", "password_hash", "firm_name", "lawyer_name",
                   "phone", "address", "role", "last_login")
        fields = [f"{k} = %s" for k in updates if k in allowed]
        values = [v for k, v in updates.items() if k in allowed]
        if not fields:
            return
        values.append(user_id)
        self._execute(f"UPDATE users SET {', '.join(fields)} WHERE user_id = %s", values)
        self.conn.commit()

    def delete_user(self, user_id: str):
        self._execute("DELETE FROM users WHERE user_id = %s", (user_id,))
        self._execute("DELETE FROM case_analyses WHERE user_id = %s", (user_id,))
        self._execute("DELETE FROM cost_logs WHERE user_id = %s", (user_id,))
        self._execute("DELETE FROM kv_store WHERE key LIKE %s", (f"u:{user_id}:%",))
        self.conn.commit()

    def update_user_last_login(self, user_id: str):
        self.update_user(user_id, {"last_login": datetime.now().isoformat()})

    def get_user_profile(self, user_id: str) -> dict:
        user = self.get_user_by_id(user_id)
        base = {
            "firm_name": "", "lawyer_name": "", "email": "",
            "phone": "", "address": "", "password_hash": "",
        }
        if user:
            base.update({
                "firm_name": user.get("firm_name", ""),
                "lawyer_name": user.get("lawyer_name", ""),
                "email": user.get("email", ""),
                "phone": user.get("phone", ""),
                "address": user.get("address", ""),
                "password_hash": user.get("password_hash", ""),
            })
        # Merge extended profile fields (notification settings etc.)
        ext_data = self._load_list_raw(f"u:{user_id}:profile_extended")
        if ext_data and isinstance(ext_data, list) and ext_data:
            base.update(ext_data[0])
        return base

    def save_user_profile(self, user_id: str, profile: dict):
        core_fields = ("firm_name", "lawyer_name", "email", "phone", "address", "password_hash")
        core = {k: profile.get(k, "") for k in core_fields}
        self.update_user(user_id, core)
        # Save extended fields (notifications etc.) separately
        extended = {k: v for k, v in profile.items() if k not in core_fields}
        if extended:
            self._save_list_raw(f"u:{user_id}:profile_extended", [extended])

    # ── Cost Logs (user-scoped) ──
    def add_cost_log(self, entry: dict):
        uid = self._uid()
        self._execute(
            "INSERT INTO cost_logs "
            "(id, timestamp, model, task, mode, input_chars, output_chars, "
            "estimated_cost, query_preview, user_id) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING",
            (
                entry.get("id", uuid.uuid4().hex[:8]),
                entry.get("timestamp", datetime.now().isoformat()),
                entry.get("model", ""), entry.get("task", ""), entry.get("mode", ""),
                entry.get("input_chars", 0), entry.get("output_chars", 0),
                entry.get("estimated_cost", 0.0), entry.get("query_preview", ""), uid,
            ),
        )
        self.conn.commit()

    def get_cost_logs(self, limit: int = 200) -> list:
        uid = self._uid()
        cur = self._execute(
            "SELECT id, timestamp, model, task, mode, input_chars, output_chars, "
            "estimated_cost, query_preview FROM cost_logs "
            "WHERE user_id = %s ORDER BY timestamp DESC LIMIT %s",
            (uid, limit),
        )
        rows = cur.fetchall()
        return [
            {
                "id": r[0], "timestamp": r[1], "model": r[2], "task": r[3],
                "mode": r[4], "input_chars": r[5], "output_chars": r[6],
                "estimated_cost": r[7], "query_preview": r[8],
            }
            for r in rows
        ]

    def get_cost_summary(self) -> dict:
        uid = self._uid()
        today = date.today().isoformat()
        month_start = date.today().replace(day=1).isoformat()
        total = self._execute(
            "SELECT COALESCE(SUM(estimated_cost),0), COUNT(*) FROM cost_logs WHERE user_id = %s",
            (uid,)
        ).fetchone()
        daily = self._execute(
            "SELECT COALESCE(SUM(estimated_cost),0), COUNT(*) FROM cost_logs "
            "WHERE user_id = %s AND timestamp >= %s", (uid, today)
        ).fetchone()
        monthly = self._execute(
            "SELECT COALESCE(SUM(estimated_cost),0), COUNT(*) FROM cost_logs "
            "WHERE user_id = %s AND timestamp >= %s", (uid, month_start)
        ).fetchone()
        return {
            "total_cost": total[0], "total_calls": total[1],
            "daily_cost": daily[0], "daily_calls": daily[1],
            "monthly_cost": monthly[0], "monthly_calls": monthly[1],
        }

    # ── Case Analyses (user-scoped) ──
    def add_case_analysis(self, case_id: str, data: dict):
        uid = self._uid()
        self._execute(
            "INSERT INTO case_analyses (id, case_id, query, response, task, mode, timestamp, user_id) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING",
            (
                data.get("id", uuid.uuid4().hex[:8]), case_id,
                data.get("query", ""), data.get("response", ""),
                data.get("task", ""), data.get("mode", ""),
                data.get("timestamp", datetime.now().isoformat()), uid,
            ),
        )
        self.conn.commit()

    def get_case_analyses(self, case_id: str) -> list:
        uid = self._uid()
        cur = self._execute(
            "SELECT id, query, response, task, mode, timestamp FROM case_analyses "
            "WHERE case_id = %s AND user_id = %s ORDER BY timestamp DESC",
            (case_id, uid),
        )
        rows = cur.fetchall()
        return [
            {
                "id": r[0], "query": r[1], "response": r[2],
                "task": r[3], "mode": r[4], "timestamp": r[5],
            }
            for r in rows
        ]

    def delete_case_analysis(self, analysis_id: str):
        self._execute("DELETE FROM case_analyses WHERE id = %s", (analysis_id,))
        self.conn.commit()

    def delete_case_analyses_for_case(self, case_id: str):
        uid = self._uid()
        self._execute(
            "DELETE FROM case_analyses WHERE case_id = %s AND user_id = %s",
            (case_id, uid)
        )
        self.conn.commit()

    # ── Lifecycle (user-scoped via namespaced kv) ──
    def save_lifecycle(self, case_id: str, data: dict):
        self.save_list(f"lifecycle_{case_id}", [data])

    def load_lifecycle(self, case_id: str) -> dict:
        result = self.load_list(f"lifecycle_{case_id}")
        if result and isinstance(result, list) and len(result) > 0:
            return result[0]
        return {}

    def save_lifecycle_progress(self, case_id: str, progress: dict):
        self.save_list(f"lifecycle_progress_{case_id}", [progress])

    def load_lifecycle_progress(self, case_id: str) -> dict:
        result = self.load_list(f"lifecycle_progress_{case_id}")
        if result and isinstance(result, list) and len(result) > 0:
            return result[0]
        return {}

    # ── Migration: copy legacy un-namespaced data to a new user account ──
    def has_legacy_data(self) -> bool:
        for key in ("cases", "clients", "time_entries", "invoices", "chat_history"):
            cur = self._execute("SELECT value FROM kv_store WHERE key = %s", (key,))
            row = cur.fetchone()
            if row and row[0] and row[0] not in ("[]", "{}", ""):
                try:
                    if json.loads(row[0]):
                        return True
                except Exception:
                    pass
        return False

    def migrate_legacy_data_to_user(self, user_id: str) -> int:
        migrated = 0
        legacy_keys = ["cases", "clients", "time_entries", "invoices", "chat_history",
                       "custom_templates", "custom_limitation_periods", "custom_maxims"]
        for key in legacy_keys:
            cur = self._execute("SELECT value FROM kv_store WHERE key = %s", (key,))
            row = cur.fetchone()
            if row and row[0] and row[0] != "[]":
                namespaced = f"u:{user_id}:{key}"
                self._execute(
                    "INSERT INTO kv_store (key, value) VALUES (%s, %s) "
                    "ON CONFLICT (key) DO NOTHING",
                    (namespaced, row[0]),
                )
                migrated += 1
        # Migrate lifecycle keys
        cur2 = self._execute(
            "SELECT key, value FROM kv_store WHERE key LIKE 'lifecycle_%'"
        )
        for lkey, lval in (cur2.fetchall() or []):
            nkey = f"u:{user_id}:{lkey}"
            self._execute(
                "INSERT INTO kv_store (key, value) VALUES (%s, %s) ON CONFLICT DO NOTHING",
                (nkey, lval),
            )
            migrated += 1
        # Migrate case analyses
        self._execute(
            "UPDATE case_analyses SET user_id = %s WHERE user_id IN ('legacy', '') OR user_id IS NULL",
            (user_id,)
        )
        # Migrate cost logs
        self._execute(
            "UPDATE cost_logs SET user_id = %s WHERE user_id IN ('legacy', '') OR user_id IS NULL",
            (user_id,)
        )
        self.conn.commit()
        return migrated

    # ── Session Tokens ──
    def create_session_token(self, user_id: str, days: int = 30, device_hint: str = "") -> str:
        """Create a persistent session token valid for `days` days. Returns the token."""
        import datetime as _dt
        token = uuid.uuid4().hex + uuid.uuid4().hex  # 64-char random, unguessable
        now = datetime.now()
        expires = now + _dt.timedelta(days=days)
        try:
            self._execute(
                "INSERT INTO user_sessions "
                "(token, user_id, created_at, expires_at, last_used, device_hint) "
                "VALUES (%s, %s, %s, %s, %s, %s)",
                (token, user_id, now.isoformat(), expires.isoformat(),
                 now.isoformat(), device_hint),
            )
            self.conn.commit()
        except Exception as e:
            try:
                self.conn.rollback()
            except Exception:
                pass
            logger.error(f"create_session_token failed: {e}")
        return token

    def validate_session_token(self, token: str) -> Optional[dict]:
        """Validate a session token. Returns the user dict if valid, else None."""
        if not token or len(token) < 32:
            return None
        try:
            cur = self._execute(
                "SELECT user_id, expires_at FROM user_sessions WHERE token = %s", (token,)
            )
            row = cur.fetchone()
            if not row:
                return None
            user_id, expires_at = row
            # Check expiry
            try:
                exp = datetime.fromisoformat(expires_at)
                if datetime.now() > exp:
                    self.revoke_session_token(token)
                    return None
            except Exception:
                pass
            # Touch last_used
            try:
                self._execute(
                    "UPDATE user_sessions SET last_used = %s WHERE token = %s",
                    (datetime.now().isoformat(), token),
                )
                self.conn.commit()
            except Exception:
                try:
                    self.conn.rollback()
                except Exception:
                    pass
            return self.get_user_by_id(user_id)
        except Exception:
            return None

    def revoke_session_token(self, token: str):
        """Delete a single session token."""
        try:
            self._execute("DELETE FROM user_sessions WHERE token = %s", (token,))
            self.conn.commit()
        except Exception:
            try:
                self.conn.rollback()
            except Exception:
                pass

    def revoke_all_user_sessions(self, user_id: str):
        """Delete all session tokens for a user (sign out all devices)."""
        try:
            self._execute("DELETE FROM user_sessions WHERE user_id = %s", (user_id,))
            self.conn.commit()
        except Exception:
            try:
                self.conn.rollback()
            except Exception:
                pass

    def get_user_sessions(self, user_id: str) -> list:
        """List all active (non-expired) sessions for a user."""
        try:
            now = datetime.now().isoformat()
            cur = self._execute(
                "SELECT token, created_at, expires_at, last_used, device_hint "
                "FROM user_sessions WHERE user_id = %s AND expires_at > %s "
                "ORDER BY last_used DESC",
                (user_id, now),
            )
            rows = cur.fetchall()
            return [
                {
                    "token": r[0], "created_at": r[1], "expires_at": r[2],
                    "last_used": r[3], "device_hint": r[4],
                }
                for r in rows
            ]
        except Exception:
            return []

    def cleanup_expired_sessions(self):
        """Remove expired tokens (call periodically)."""
        try:
            self._execute(
                "DELETE FROM user_sessions WHERE expires_at < %s",
                (datetime.now().isoformat(),),
            )
            self.conn.commit()
        except Exception:
            try:
                self.conn.rollback()
            except Exception:
                pass

    def close(self):
        self.conn.close()

    def ensure_connected(self):
        """Ping the connection; reconnect + re-init tables if dead."""
        try:
            self.conn.cursor().execute("SELECT 1")
        except Exception:
            try:
                self.conn.rollback()
            except Exception:
                pass
            try:
                self.conn = self._connect()
                self._init_tables()
            except Exception as e:
                logger.error(f"DB reconnect failed: {e}")


@st.cache_resource
def get_db() -> Database:
    """Singleton DB connection per Streamlit server process."""
    return Database()

def persist(key: str):
    """Save a session_state list to DB under the current user's namespace."""
    get_db().save_list(key, st.session_state.get(key, []))


def persist_profile():
    """Save current user's full profile to DB."""
    get_db().save_profile(st.session_state.get("profile", {}))


def load_user_data():
    """Load all user-specific data from DB into session state. Called once after login."""
    if not st.session_state.get("current_user_id"):
        return
    db = get_db()
    st.session_state.cases = db.load_list("cases") or []
    st.session_state.clients = db.load_list("clients") or []
    st.session_state.time_entries = db.load_list("time_entries") or []
    st.session_state.invoices = db.load_list("invoices") or []
    st.session_state.chat_history = db.load_list("chat_history") or []
    st.session_state.custom_templates = db.load_list("custom_templates") or []
    st.session_state.custom_limitation_periods = db.load_list("custom_limitation_periods") or []
    st.session_state.custom_maxims = db.load_list("custom_maxims") or []
    st.session_state.profile = db.get_profile()


# ═══════════════════════════════════════════════════════
# MULTI-USER AUTH
# ═══════════════════════════════════════════════════════
def hash_password(password: str) -> str:
    """PBKDF2-HMAC-SHA256 with random salt. Format: pbkdf2$<salt>$<dk>"""
    import secrets as _sec
    salt = _sec.token_hex(16)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 260_000)
    return f"pbkdf2${salt}${dk.hex()}"


def verify_password(password: str, stored: str) -> bool:
    """Verify against PBKDF2 hash, or legacy plain SHA-256 (auto-upgrades on login)."""
    if stored.startswith("pbkdf2$"):
        try:
            _, salt, dk_hex = stored.split("$")
            dk = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 260_000)
            return dk.hex() == dk_hex
        except Exception:
            return False
    return hashlib.sha256(password.encode()).hexdigest() == stored


def is_allow_registration() -> bool:
    try:
        return str(st.secrets.get("ALLOW_REGISTRATION", "false")).lower() == "true"
    except Exception:
        return os.getenv("ALLOW_REGISTRATION", "false").lower() == "true"


def do_login(username: str, password: str, remember_me: bool = True) -> bool:
    """Authenticate user, load their data into session. Returns True on success."""
    db = get_db()
    user = db.get_user_by_username(username.strip())
    if not user:
        return False
    if not verify_password(password, user["password_hash"]):
        return False
    # Auto-upgrade legacy SHA-256 → PBKDF2 on next login
    if not user["password_hash"].startswith("pbkdf2$"):
        get_db().update_user(user["user_id"], {"password_hash": hash_password(password)})
    uid = user["user_id"]
    st.session_state.authenticated = True
    st.session_state.current_user_id = uid
    st.session_state.current_username = user["username"]
    st.session_state.current_user_role = user["role"]
    db.update_user_last_login(uid)
    load_user_data()
    st.session_state.user_data_loaded = True
    # ── Persistent session token ──
    if remember_me:
        token = db.create_session_token(uid, days=30)
        st.session_state["_session_token"] = token
        pass  # token stored in session_state only — never exposed in URL
    return True


def do_auto_login_from_token(token: str) -> bool:
    """Silently restore a session from a persistent token. Returns True if valid."""
    if not token:
        return False
    db = get_db()
    db.cleanup_expired_sessions()
    user = db.validate_session_token(token)
    if not user:
        return False  # Token invalid/expired
    uid = user["user_id"]
    st.session_state.authenticated = True
    st.session_state.current_user_id = uid
    st.session_state.current_username = user["username"]
    st.session_state.current_user_role = user["role"]
    st.session_state["_session_token"] = token
    load_user_data()
    st.session_state.user_data_loaded = True
    return True


def do_logout():
    """Revoke session token, clear query params, and wipe session state."""
    db = get_db()
    token = st.session_state.get("_session_token", "")
    if token:
        db.revoke_session_token(token)
    try:
        st.query_params.clear()
    except Exception:
        pass
    clear_keys = [
        "authenticated", "current_user_id", "current_username", "current_user_role",
        "user_data_loaded", "_session_token",
        "cases", "clients", "time_entries", "invoices",
        "chat_history", "custom_templates", "custom_limitation_periods", "custom_maxims",
        "profile", "last_response", "original_query", "research_results",
        "loaded_template", "imported_doc",
        "wp_result", "wp_role_label", "wp_facts_saved", "wp_reexam_result",
        "wp_witness_log", "wp_contra_result",
        "nf_feed_data", "nf_subject_loaded", "nf_deepdive", "nf_bookmarks", "nf_scan_result",
    ]
    for k in clear_keys:
        st.session_state.pop(k, None)
    st.rerun()


def render_login_screen():
    st.markdown(get_theme_css(st.session_state.get("theme", "⚖️ Corporate")), unsafe_allow_html=True)
    st.markdown("""
<div class="hero">
  <h1>⚖️ LexiAssist v8.0</h1>
  <p>Elite AI Legal Engine &nbsp;&middot;&nbsp; Nigerian Law &nbsp;&middot;&nbsp; Built for Practitioners</p>
</div>""", unsafe_allow_html=True)
    _, col, _ = st.columns([1, 2, 1])
    with col:
        st.markdown('<div class="login-card">', unsafe_allow_html=True)
        st.markdown("#### 🔒 Sign In to Your Workspace")
        st.markdown("<hr style='margin:0.6rem 0 1rem 0;border-color:#1a2e4a18'>", unsafe_allow_html=True)
        login_tabs = ["🔒 Login", "📝 Register"] if is_allow_registration() else ["🔒 Login"]
        if len(login_tabs) > 1:
            tab_login, tab_reg = st.tabs(login_tabs)
        else:
            tab_login = st.container(); tab_reg = None
        with tab_login:
            with st.form("login_form", clear_on_submit=False):
                username_inp = st.text_input("Username", placeholder="your.username", key="login_username_inp")
                password_inp = st.text_input("Password", type="password", key="login_password_inp")
                remember_me  = st.checkbox("Stay signed in for 30 days", value=True, key="login_remember_me")
                if st.form_submit_button("🔒 Sign In", type="primary", width='stretch'):
                    if not username_inp.strip() or not password_inp:
                        st.error("❌ Enter both username and password.")
                    elif do_login(username_inp.strip(), password_inp, remember_me=remember_me):
                        st.success(f"✅ Welcome back, @{st.session_state.current_username}!")
                        time.sleep(0.3); st.rerun()
                    else:
                        st.error("❌ Invalid username or password.")
        if tab_reg is not None:
            with tab_reg: render_register_form("reg_self")
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown(
            "<div style='text-align:center;margin-top:1.2rem;color:#64748b;font-size:0.82rem;'>"
            "Contact your firm administrator to create an account.</div>",
            unsafe_allow_html=True)


def render_register_form(key_prefix: str, admin_mode: bool = False):
    """Reusable registration / account-creation form."""
    db = get_db()
    is_first_user = not db.has_any_users()

    with st.form(f"{key_prefix}_form", clear_on_submit=True):
        r1, r2 = st.columns(2)
        with r1:
            reg_username = st.text_input("Username *", placeholder="e.g. amaka.obi", key=f"{key_prefix}_uname")
            reg_pw = st.text_input("Password *", type="password", key=f"{key_prefix}_pw")
            reg_confirm = st.text_input("Confirm Password *", type="password", key=f"{key_prefix}_confirm")
        with r2:
            reg_lawyer = st.text_input("Full Name *", placeholder="Barr. Amaka Obi", key=f"{key_prefix}_lname")
            reg_firm = st.text_input("Firm Name", placeholder="Obi & Associates", key=f"{key_prefix}_firm")
            reg_email = st.text_input("Email", placeholder="amaka@obilaw.com", key=f"{key_prefix}_email")

        role_options = ["user", "admin"] if admin_mode else ["user"]
        reg_role = st.selectbox("Role", role_options, key=f"{key_prefix}_role") if admin_mode else "user"

        btn_label = "🛡️ Create Admin Account" if is_first_user else "✅ Create Account"
        if st.form_submit_button(btn_label, type="primary", width='stretch'):
            uname = reg_username.strip().lower()
            if not uname or not reg_pw or not reg_lawyer.strip():
                st.error("❌ Username, password, and full name are required.")
                return False
            if len(uname) < 3:
                st.error("❌ Username must be at least 3 characters.")
                return False
            if reg_pw != reg_confirm:
                st.error("❌ Passwords do not match.")
                return False
            if len(reg_pw) < 6:
                st.error("❌ Password must be at least 6 characters.")
                return False
            if db.get_user_by_username(uname):
                st.error(f"❌ Username '{uname}' is already taken.")
                return False

            role = "admin" if (is_first_user or reg_role == "admin") else "user"
            user_id = uuid.uuid4().hex[:12]
            ok = db.create_user({
                "user_id": user_id,
                "username": uname,
                "password_hash": hash_password(reg_pw),
                "firm_name": reg_firm.strip(),
                "lawyer_name": reg_lawyer.strip(),
                "email": reg_email.strip(),
                "role": role,
            })
            if ok:
                if is_first_user:
                    # Migrate any legacy data to this admin account
                    migrated = db.migrate_legacy_data_to_user(user_id)
                    if migrated > 0:
                        st.info(f"ℹ️ {migrated} legacy data item(s) migrated to your account.")
                if not admin_mode:
                    # Auto-login after self-registration
                    do_login(uname, reg_pw)
                    st.success(f"✅ Account created! Welcome, {reg_lawyer.strip()}.")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.success(f"✅ Account created for {reg_lawyer.strip()} (@{uname}) [{role}].")
                return True
            else:
                st.error("❌ Account creation failed. Try a different username.")
                return False
    return False


def render_create_admin_screen():
    """First-run screen: no users exist yet."""
    st.markdown(get_theme_css(st.session_state.get("theme", "⚖️ Corporate")), unsafe_allow_html=True)
    st.markdown("""
<div class="hero">
  <h1>⚖️ LexiAssist v8.0</h1>
  <p>First-time setup &nbsp;&middot;&nbsp; Create your administrator account below</p>
</div>""", unsafe_allow_html=True)
    _, col, _ = st.columns([1, 2, 1])
    with col:
        st.markdown('<div class="login-card">', unsafe_allow_html=True)
        st.markdown("#### 🛡️ Create Administrator Account")
        st.info("ℹ️ No accounts exist yet. The first account you create becomes the admin.")
        st.markdown("<hr style='margin:0.6rem 0 1rem 0;border-color:#1a2e4a18'>", unsafe_allow_html=True)
        render_register_form("first_admin")
        st.markdown("</div>", unsafe_allow_html=True)


def render_setup_screen():
    st.markdown("""
    <div class="hero">
        <h1>⚖️ LexiAssist v8.0</h1>
        <p>Elite AI Legal Engine for Nigerian Lawyers</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### 🔐 Secure API Configuration")
    st.markdown("""
    Your API key is required to power the AI engine. It is **never displayed**
    in the sidebar or stored outside this session.

    **Recommended:** Add your key to Streamlit Secrets (`.streamlit/secrets.toml`)
    or set the `GEMINI_API_KEY` environment variable so this screen never appears.
    """)

    with st.form("api_setup_form"):
        key_input = st.text_input(
            "Google Gemini API Key",
            type="password",
            placeholder="Paste your API key here…",
            help="Get a key at https://aistudio.google.com/app/apikey",
        )
        model_sel = st.selectbox("AI Model", SUPPORTED_MODELS, index=0)
        submitted = st.form_submit_button("🔗 Connect", type="primary", width='stretch')

        if submitted:
            if key_input and len(key_input.strip()) >= 10:
                st.session_state.gemini_model = model_sel
                with st.spinner("🔗 Connecting to Gemini…"):
                    if manual_connect(key_input.strip()):
                        st.success("✅ Connected! Redirecting…")
                        time.sleep(1)
                        st.rerun()
            else:
                st.error("❌ Please enter a valid API key.")

    st.divider()
    st.caption("💡 **Tip:** To skip this screen permanently, add to `.streamlit/secrets.toml`:")
    st.code('GEMINI_API_KEY = "your-key-here"\nGEMINI_MODEL = "gemini-2.5-flash"\n# ALLOW_REGISTRATION = "true"  # let users self-register', language="toml")


# ═══════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════
def render_sidebar():
    with st.sidebar:
        firm = get_firm_name()
        corp = (st.session_state.get("theme", "⚖️ Corporate") == "⚖️ Corporate")
        name_display = firm if (firm and firm != "LexiAssist") else "LexiAssist v8.0"
        tag_display  = "Powered by LexiAssist v8.0" if (firm and firm != "LexiAssist") else "Elite AI Legal Engine"
        hdr_col = "#c9a84c" if corp else "#1a2e4a"
        cap_col = "#a0bcd8" if corp else "#64748b"
        div_col = "#2d4a6e" if corp else "#e2e8f0"
        st.markdown(f"""
<div style="padding:1rem 0.4rem 0.8rem 0.4rem;border-bottom:1px solid {div_col};">
  <div style="font-size:1.05rem;font-weight:800;color:{hdr_col};letter-spacing:-0.01em;">⚖️ {esc(name_display)}</div>
  <div style="font-size:0.74rem;margin-top:0.15rem;color:{cap_col};">{esc(tag_display)}</div>
</div>""", unsafe_allow_html=True)
        uname = st.session_state.get("current_username","")
        urole = st.session_state.get("current_user_role","")
        if uname:
            role_icon = "🛡️ Admin" if urole=="admin" else "👤 User"
            bg_c = "#ffffff10" if corp else "#22c55e18"
            bd_c = "#ffffff25" if corp else "#22c55e55"
            tx_c = "#c9a84c"   if corp else "#059669"
            st.markdown(f"""
<div style="margin:0.8rem 0 0.4rem 0;padding:0.6rem 0.8rem;background:{bg_c};border:1px solid {bd_c};border-radius:8px;">
  <div style="font-weight:700;font-size:0.9rem;color:{tx_c};">@{esc(uname)}</div>
  <div style="font-size:0.75rem;opacity:0.75;margin-top:0.1rem;">{role_icon}</div>
</div>""", unsafe_allow_html=True)
            if st.button("🚪 Sign Out", key="sidebar_logout_btn", width='stretch'):
                do_logout()
        st.divider()
        c1,c2 = st.columns(2)
        with c1: st.metric("Cases",    len(get_active_cases()))
        with c2: st.metric("Sessions", len(st.session_state.chat_history))
        st.divider()
        st.markdown("**🧠 Response Mode**")
        modes = list(RESPONSE_MODES.keys())
        cur_m = modes.index(st.session_state.response_mode) if st.session_state.response_mode in modes else 1
        mode = st.radio("Depth", modes, index=cur_m,
            format_func=lambda x: RESPONSE_MODES[x]["label"],
            key="sidebar_mode_radio", label_visibility="collapsed")
        if mode != st.session_state.response_mode:
            st.session_state.response_mode = mode; st.rerun()
        sel = RESPONSE_MODES[st.session_state.response_mode]
        st.caption(sel["desc"]); st.caption(f"Token limit: {sel['tokens']:,}")
        st.divider()
        st.markdown("**🎨 Theme**")
        theme_names = list(THEMES.keys())
        cur_t = theme_names.index(st.session_state.theme) if st.session_state.theme in theme_names else 0
        theme = st.selectbox(
            "Theme", theme_names, index=cur_t,
            key="sidebar_theme_sel", label_visibility="collapsed",
            help=THEMES[theme_names[cur_t]]["description"])
        if theme != st.session_state.theme:
            st.session_state.theme = theme; st.rerun()
        # ── Accessibility controls ────────────────────────────────
        with st.expander("♿ Accessibility", expanded=False):
            fss = st.slider(
                "Font size", 0.85, 1.4, st.session_state.font_size_scale, 0.05,
                key="sidebar_font_scale",
                help="Scale all text up or down")
            if fss != st.session_state.font_size_scale:
                st.session_state.font_size_scale = fss; st.rerun()
            hc = st.checkbox(
                "High contrast", value=st.session_state.high_contrast,
                key="sidebar_high_contrast",
                help="Maximise text/background contrast")
            if hc != st.session_state.high_contrast:
                st.session_state.high_contrast = hc; st.rerun()
            rm = st.checkbox(
                "Reduce motion", value=st.session_state.reduce_motion,
                key="sidebar_reduce_motion",
                help="Disable all CSS animations and transitions")
            if rm != st.session_state.reduce_motion:
                st.session_state.reduce_motion = rm; st.rerun()
        st.divider()
        st.markdown("**🤖 AI Engine**")
        if st.session_state.api_configured:
            st.success(f"✅ Connected · `{st.session_state.gemini_model}`")
            ms = st.selectbox("Model", SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.gemini_model)
                      if st.session_state.gemini_model in SUPPORTED_MODELS else 0,
                key="sidebar_model_sel", label_visibility="collapsed")
            if ms != st.session_state.gemini_model:
                st.session_state.gemini_model = ms; st.rerun()
            summary = get_db().get_cost_summary()
            if summary["total_calls"] > 0:
                st.caption(f"💰 Today: ${summary['daily_cost']:.4f} ({summary['daily_calls']} calls)")
                st.caption(f"📅 Month: ${summary['monthly_cost']:.4f} ({summary['monthly_calls']} calls)")
        else:
            st.error("🔴 Not connected")
        st.divider()
        st.markdown("**💾 Data**")
        if st.button("📥 Export All Data (JSON)", width='stretch', key="sidebar_export_btn"):
            export_data = {
                "export_date": datetime.now().isoformat(), "version": "8.0",
                "cases": st.session_state.cases,
                "clients": st.session_state.clients,
                "time_entries": st.session_state.time_entries,
                "invoices": st.session_state.invoices,
                "chat_history": st.session_state.chat_history,
                "custom_templates": st.session_state.custom_templates,
                "custom_limitation_periods": st.session_state.custom_limitation_periods,
                "custom_maxims": st.session_state.custom_maxims,
                "profile": st.session_state.profile,
                "cost_logs": get_db().get_cost_logs(500),
            }
            st.download_button("⬇️ Download JSON",
                json.dumps(export_data, indent=2, default=str),
                f"lexiassist_backup_{datetime.now():%Y%m%d_%H%M}.json",
                "application/json", key="sidebar_dl_json", width='stretch')
        st.markdown("**📤 Import Files**")
        uploaded = st.file_uploader("Upload", type=UPLOAD_TYPES, accept_multiple_files=False,
            key="sidebar_file_upload", label_visibility="collapsed",
            help="Supports: PDF, DOCX, TXT, XLSX, CSV, JSON, RTF")
        if uploaded:
            try:
                ext = uploaded.name.split(".")[-1].lower()
                if ext == "json":
                    raw = json.loads(uploaded.getvalue().decode("utf-8", errors="ignore"))
                    if isinstance(raw,dict) and any(k in raw for k in ["cases","clients"]):
                        for k in ["cases","clients","time_entries","invoices","chat_history",
                                  "custom_templates","custom_limitation_periods","custom_maxims"]:
                            if k in raw: st.session_state[k]=raw[k]; persist(k)
                        if "profile" in raw and isinstance(raw["profile"],dict):
                            st.session_state.profile.update(raw["profile"]); persist_profile()
                        st.success("✅ Data imported!"); st.rerun()
                    else:
                        text=json.dumps(raw,indent=2)
                        st.session_state.imported_doc={"name":uploaded.name,"type":ext,
                            "size":len(uploaded.getvalue()),"full_text":text,"preview":text[:600]}
                        st.success(f"✅ {uploaded.name} loaded → AI Assistant"); st.rerun()
                else:
                    text=extract_file_text(uploaded)
                    st.session_state.imported_doc={"name":uploaded.name,"type":ext,
                        "size":len(uploaded.getvalue()),"full_text":text,
                        "preview":text[:600]+("…" if len(text)>600 else "")}
                    st.success(f"✅ {uploaded.name} loaded → AI Assistant"); st.rerun()
            except Exception as e:
                st.error(f"❌ Import error: {e}")
        st.divider()
        st.caption("⚖️ LexiAssist v8.0 © 2026")
        st.caption("🇳🇬 Nigerian Law · 🤖 AI-Powered")


def render_home():
    firm = get_firm_name()
    subtitle = f"{esc(firm)} · " if firm and firm != "LexiAssist" else ""
    st.markdown(f"""
    <div class="hero">
        <h1>⚖️ LexiAssist v8.0</h1>
        <p>{subtitle}Elite AI Legal Engine for Nigerian Lawyers<br>
        Position-taking · Strategy-driven · Risk-ranked · Litigator-minded</p>
    </div>
    """, unsafe_allow_html=True)

    # Stats row
    cost_summary = get_db().get_cost_summary()
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.markdown(f'<div class="stat-card"><div class="stat-value">{len(st.session_state.cases)}</div><div class="stat-label">Total Cases</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="stat-card"><div class="stat-value">{len(get_active_cases())}</div><div class="stat-label">Active Cases</div></div>', unsafe_allow_html=True)
    with c3:
        st.markdown(f'<div class="stat-card"><div class="stat-value">{len(st.session_state.clients)}</div><div class="stat-label">Clients</div></div>', unsafe_allow_html=True)
    with c4:
        st.markdown(f'<div class="stat-card"><div class="stat-value">{total_hours():.1f}h</div><div class="stat-label">Billable Hours</div></div>', unsafe_allow_html=True)
    with c5:
        st.markdown(f'<div class="stat-card"><div class="stat-value">{len(st.session_state.chat_history)}</div><div class="stat-label">AI Sessions</div></div>', unsafe_allow_html=True)

    st.markdown("")

    # ── Getting Started card — shown only to brand-new users ──
    is_new_user = (
        len(st.session_state.cases) == 0
        and len(st.session_state.chat_history) == 0
        and len(st.session_state.clients) == 0
    )
    if is_new_user:
        st.markdown("""
<div class="custom-card" style="border-left-color:var(--la-accent-2); margin-bottom:1.2rem;">
<h4 style="font-size:1.05rem; margin-bottom:0.6rem;">👋 Welcome to LexiAssist — here's how to get started</h4>
<p style="margin:0 0 0.8rem 0; font-size:0.93rem;">
Your workspace is empty. Pick any of the three actions below to begin:
</p>
<table style="width:100%; border-collapse:separate; border-spacing:0 6px; font-size:0.9rem;">
<tr>
  <td style="width:32px; font-size:1.25rem; vertical-align:top;">🧠</td>
  <td><strong>Ask your first legal question</strong><br>
  <span style="color:var(--la-text-secondary);">Open the <em>AI Assistant</em> tab, type your query, choose a response mode and hit Generate.</span></td>
</tr>
<tr>
  <td style="font-size:1.25rem; vertical-align:top;">📁</td>
  <td><strong>Create your first case</strong><br>
  <span style="color:var(--la-text-secondary);">Go to the <em>Cases</em> tab, click ➕ Add Case and fill in the matter details — all analysis you generate can be saved to it.</span></td>
</tr>
<tr>
  <td style="font-size:1.25rem; vertical-align:top;">📑</td>
  <td><strong>Review a contract</strong><br>
  <span style="color:var(--la-text-secondary);">On the <em>AI Assistant</em> tab select <em>Contract Review</em> mode, paste or upload your contract, and receive a full clause-by-clause risk matrix.</span></td>
</tr>
</table>
</div>""", unsafe_allow_html=True)

    col_left, col_right = st.columns([3, 2])
    with col_left:
        st.markdown("### 📅 Upcoming Hearings")
        hearings = get_hearings()
        if hearings:
            for h in hearings[:8]:
                d = days_until(h["date"])
                badge = "badge-err" if d <= 3 else ("badge-warn" if d <= 7 else "badge-ok")
                st.markdown(f"""<div class="custom-card">
                    <h4>{esc(h['title'])}</h4>
                    Suit: {esc(h['suit'])} · Court: {esc(h['court'])}<br>
                    📅 {esc(fmt_date(h['date']))}
                    <span class="badge {badge}">{esc(relative_date(h['date']))}</span>
                </div>""", unsafe_allow_html=True)
        else:
            st.info("No upcoming hearings. Add cases with hearing dates.")

    with col_right:
        st.markdown("### 🧠 Recent AI Sessions")
        history = st.session_state.chat_history
        if history:
            for entry in reversed(history[-6:]):
                mode_lbl = RESPONSE_MODES.get(entry.get("mode", ""), {}).get("label", "")
                st.markdown(f"""<div class="history-item">
                    <strong>{esc(entry.get('query', '')[:80])}{'…' if len(entry.get('query', '')) > 80 else ''}</strong><br>
                    <small>{esc(entry.get('timestamp', ''))} · {esc(mode_lbl)} · {entry.get('word_count', 0)} words</small>
                </div>""", unsafe_allow_html=True)
        else:
            st.info("No AI sessions yet. Go to AI Assistant to start.")

        # Cost summary on home
        if cost_summary["total_calls"] > 0:
            st.markdown("### 💰 AI Costs")
            kc1, kc2 = st.columns(2)
            with kc1:
                st.metric("Today", f"${cost_summary['daily_cost']:.4f}")
            with kc2:
                st.metric("This Month", f"${cost_summary['monthly_cost']:.4f}")

    st.markdown("---")
    st.markdown("### 🏆 Elite Features")
    f1, f2, f3, f4 = st.columns(4)
    with f1:
        st.markdown("""<div class="custom-card">
            <h4>🎯 Position-Taking</h4>
            <p>No more "may be liable" — firm conclusions backed by authority</p>
        </div>""", unsafe_allow_html=True)
    with f2:
        st.markdown("""<div class="custom-card">
            <h4>📑 Contract Review</h4>
            <p>Clause-by-clause risk analysis with red flag matrix</p>
        </div>""", unsafe_allow_html=True)
    with f3:
        st.markdown("""<div class="custom-card">
            <h4>⚔️ Strategy Layer</h4>
            <p>Actionable next steps per party — litigator-grade advice</p>
        </div>""", unsafe_allow_html=True)
    with f4:
        st.markdown("""<div class="custom-card">
            <h4>💾 Cloud Persistence</h4>
            <p>PostgreSQL-backed — all data synced across sessions and devices</p>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: AI ASSISTANT (FULL-FEATURED)
# ═══════════════════════════════════════════════════════
def render_ai():
    st.markdown("""<div class="page-header">
        <h2>🧠 AI Legal Assistant</h2>
        <p>Position-taking · Strategy-driven · Risk-ranked · Contract Review</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ AI not connected. Configure your API key on the setup screen.")
        return

    mode = st.session_state.response_mode
    mode_info = RESPONSE_MODES[mode]
    st.info(f"**Mode: {mode_info['label']}** — {mode_info['desc']} (up to {mode_info['tokens']:,} tokens)")

    # ── Imported Document Context ──
    doc_context = ""
    if st.session_state.imported_doc:
        with st.expander(f"📎 Imported: {st.session_state.imported_doc['name']}", expanded=False):
            doc = st.session_state.imported_doc
            st.caption(f"Type: {doc['type'].upper()} · Size: {doc['size']:,} bytes")
            st.text_area("Preview", doc["preview"], height=120, disabled=True, key="doc_preview_ta")
            dc1, dc2 = st.columns(2)
            with dc1:
                if st.button("📋 Use as Context", key="use_doc_ctx_btn", width='stretch'):
                    doc_context = doc["full_text"]
                    st.success("✅ Document loaded as context for your query.")
            with dc2:
                if st.button("🗑️ Clear Document", key="clear_doc_btn", width='stretch'):
                    st.session_state.imported_doc = None
                    st.rerun()
        if not doc_context and st.session_state.imported_doc:
            doc_context = st.session_state.imported_doc.get("full_text", "")

    # ── Session History with Compare Selection ──
    if st.session_state.chat_history:
        with st.expander(f"📚 Session History ({len(st.session_state.chat_history)} entries) — select 2 to compare", expanded=False):
            # Compare selections
            compare_sels = st.session_state.get("compare_selections", [])

            for i, entry in enumerate(reversed(st.session_state.chat_history[-20:])):
                real_idx = len(st.session_state.chat_history) - 1 - i
                mode_lbl = RESPONSE_MODES.get(entry.get("mode", ""), {}).get("label", "")
                task_lbl = TASK_TYPES.get(entry.get("task", ""), {}).get("label", "")

                hc1, hc2, hc3 = st.columns([0.5, 4.5, 1])
                with hc1:
                    is_checked = real_idx in compare_sels
                    checked = st.checkbox(
                        "Sel", value=is_checked, key=f"cmp_chk_{real_idx}",
                        label_visibility="collapsed",
                    )
                    if checked and real_idx not in compare_sels:
                        compare_sels.append(real_idx)
                        if len(compare_sels) > 2:
                            compare_sels.pop(0)
                        st.session_state.compare_selections = compare_sels
                    elif not checked and real_idx in compare_sels:
                        compare_sels.remove(real_idx)
                        st.session_state.compare_selections = compare_sels

                with hc2:
                    st.markdown(f"""<div class="history-item">
                        <strong>{esc(entry.get('query', '')[:100])}</strong><br>
                        <small>{esc(entry.get('timestamp', ''))} · {esc(task_lbl)} · {esc(mode_lbl)} · {entry.get('word_count', 0)} words</small>
                    </div>""", unsafe_allow_html=True)
                with hc3:
                    if st.button("📖", key=f"load_hist_{real_idx}", width='stretch', help="Load this session"):
                        st.session_state.selected_history_idx = real_idx
                        st.session_state.last_response = entry["response"]
                        st.session_state.original_query = entry["query"]
                        st.session_state.last_task = entry.get("task", "general")
                        st.session_state.last_mode = entry.get("mode", "standard")
                        st.rerun()

            # Compare button
            compare_sels = st.session_state.get("compare_selections", [])
            if len(compare_sels) == 2:
                st.markdown("---")
                st.markdown(f"**📊 Compare:** Session {compare_sels[0]+1} vs Session {compare_sels[1]+1}")
                if st.button("🔬 Run Analysis Comparison", type="primary", key="run_compare_btn", width='stretch'):
                    entry_a = st.session_state.chat_history[compare_sels[0]]
                    entry_b = st.session_state.chat_history[compare_sels[1]]
                    with st.spinner("🔬 Comparing analyses…"):
                        verdict = run_comparison(entry_a, entry_b)
                    st.session_state["comparison_result"] = verdict
                    st.rerun()
            elif len(compare_sels) == 1:
                st.caption("☑️ Select one more session to enable comparison.")

    # ── Show comparison result ──
    if st.session_state.get("comparison_result"):
        st.markdown("---")
        st.markdown("### 📊 Analysis Comparison Verdict")
        verdict = st.session_state["comparison_result"]
        st.markdown(f'<div class="response-box">{esc(verdict)}</div>', unsafe_allow_html=True)

        fname = f"LexiAssist_Comparison_{datetime.now():%Y%m%d_%H%M}"
        vc1, vc2, vc3, vc4 = st.columns(4)
        with vc1:
            st.download_button("📥 TXT", export_txt(verdict, "Analysis Comparison"), f"{fname}.txt", "text/plain", key="cmp_dl_txt", width='stretch')
        with vc2:
            st.download_button("📥 HTML", export_html(verdict, "Analysis Comparison"), f"{fname}.html", "text/html", key="cmp_dl_html", width='stretch')
        with vc3:
            safe_pdf_download(verdict, "Analysis Comparison", fname, "cmp_dl_pdf")
        with vc4:
            safe_docx_download(verdict, "Analysis Comparison", fname, "cmp_dl_docx")

        if st.button("✖️ Close Comparison", key="close_cmp_btn"):
            st.session_state["comparison_result"] = ""
            st.session_state.compare_selections = []
            st.rerun()
        st.markdown("---")

    # ── Show selected history entry ──
    if st.session_state.selected_history_idx is not None:
        idx = st.session_state.selected_history_idx
        if 0 <= idx < len(st.session_state.chat_history):
            entry = st.session_state.chat_history[idx]
            st.markdown("---")
            st.markdown(f"### 📖 Viewing: Session from {entry.get('timestamp', '')}")
            task_lbl = TASK_TYPES.get(entry.get("task", ""), {}).get("label", "")
            mode_lbl = RESPONSE_MODES.get(entry.get("mode", ""), {}).get("label", "")
            st.caption(f"{task_lbl} · {mode_lbl} · {entry.get('word_count', 0)} words")
            st.markdown(f"**Query:** {esc(entry['query'])}")
            st.markdown(f'<div class="response-box">{esc(entry["response"])}</div>', unsafe_allow_html=True)

            fname = f"LexiAssist_{entry.get('timestamp', '').replace(' ', '_').replace(':', '')}"
            hx1, hx2, hx3, hx4 = st.columns(4)
            with hx1:
                st.download_button("📥 TXT", export_txt(entry["response"]), f"{fname}.txt", "text/plain", key=f"hist_dl_txt_{idx}", width='stretch')
            with hx2:
                st.download_button("📥 HTML", export_html(entry["response"]), f"{fname}.html", "text/html", key=f"hist_dl_html_{idx}", width='stretch')
            with hx3:
                safe_pdf_download(entry["response"], "Legal Analysis", fname, f"hist_dl_pdf_{idx}")
            with hx4:
                safe_docx_download(entry["response"], "Legal Analysis", fname, f"hist_dl_docx_{idx}")

            if st.button("✖️ Close", key="close_hist_view"):
                st.session_state.selected_history_idx = None
                st.rerun()
            st.markdown("---")

    # ── Main Query Input ──
    st.markdown("### 💬 New Query")
    tc1, tc2 = st.columns([2, 1])
    with tc1:
        task_keys = list(TASK_TYPES.keys())
        task = st.selectbox(
            "Task Type", task_keys,
            format_func=lambda x: f"{TASK_TYPES[x]['label']} — {TASK_TYPES[x]['desc']}",
            key="ai_task_sel",
        )
    with tc2:
        st.markdown("")
        st.markdown(f"**Mode:** {mode_info['label']}")
        st.caption(f"Max output: {mode_info['tokens']:,} tokens")

    # Special hint for contract review
    if task == "contract_review":
        st.info("📑 **Contract Review Mode:** Paste or upload a contract. The AI will analyse each clause for risk, flag issues, and provide a red flag matrix with an overall signability grade.")

    prefill = st.session_state.pop("loaded_template", "") if "loaded_template" in st.session_state and st.session_state.get("loaded_template") else ""
    query = st.text_area(
        "Your Legal Query",
        value=prefill,
        height=200,
        placeholder="Describe your legal question in detail…\n\nFor Contract Review: paste the full contract text here, or upload the document via the sidebar.",
        key="ai_query_ta",
    )

    # ── Action Buttons ──
    bc1, bc2, bc3 = st.columns(3)
    with bc1:
        generate_btn = st.button(
            f"🧠 Generate ({mode_info['label']})",
            type="primary", width='stretch',
            disabled=not query.strip(), key="ai_generate_btn",
        )
    with bc2:
        issue_btn = st.button(
            "🔍 Issue Spot", width='stretch',
            disabled=not query.strip(), key="ai_issue_btn",
        )
    with bc3:
        clear_btn = st.button(
            "🗑️ Clear", width='stretch', key="ai_clear_btn",
        )

    if clear_btn:
        st.session_state.last_response = ""
        st.session_state.original_query = ""
        st.session_state.selected_history_idx = None
        st.session_state["comparison_result"] = ""
        st.session_state.compare_selections = []
        st.rerun()

    # ── Empty-query guard ──
    if (generate_btn or issue_btn) and not query.strip():
        st.warning("⚠️ Please enter your legal query before generating a response.")

    # ── Issue Spotting ──
    if issue_btn and query.strip():
        with st.spinner("🔍 Decomposing issues…"):
            result = run_issue_spot(query.strip())
        st.markdown("### 🔍 Issue Decomposition")
        st.markdown(f'<div class="response-box">{esc(result)}</div>', unsafe_allow_html=True)

    # ── Main Generation ──
    if generate_btn and query.strip():
        with st.spinner(f"🧠 Generating {mode_info['label']} analysis…"):
            start_t = time.time()
            result = run_ai_query(query.strip(), task, mode, doc_context)
            elapsed = time.time() - start_t

        st.session_state.last_response = result
        st.session_state.original_query = query.strip()
        st.session_state.last_task = task
        st.session_state.last_mode = mode
        st.session_state.selected_history_idx = None
        add_to_history(query.strip(), result, task, mode)
        st.caption(f"⏱️ Generated in {elapsed:.1f}s · {len(result.split()):,} words")

    # ── Display Response ──
    if st.session_state.last_response and st.session_state.selected_history_idx is None:
        response = st.session_state.last_response
        st.markdown("---")
        task_lbl = TASK_TYPES.get(st.session_state.get("last_task", "general"), {}).get("label", "Analysis")
        st.markdown(f"### 📋 {task_lbl} Result")

        # Export row
        fname = f"LexiAssist_Analysis_{datetime.now():%Y%m%d_%H%M}"
        ex1, ex2, ex3, ex4 = st.columns(4)
        with ex1:
            st.download_button("📥 TXT", export_txt(response), f"{fname}.txt", "text/plain", key="resp_dl_txt", width='stretch')
        with ex2:
            st.download_button("📥 HTML", export_html(response), f"{fname}.html", "text/html", key="resp_dl_html", width='stretch')
        with ex3:
            safe_pdf_download(response, "Legal Analysis", fname, "resp_dl_pdf")
        with ex4:
            safe_docx_download(response, "Legal Analysis", fname, "resp_dl_docx")

        st.markdown(f'<div class="response-box">{esc(response)}</div>', unsafe_allow_html=True)

        # ── Copy to clipboard ──
        _copy_html = f"""
<style>
#la-copy-btn {{
    display:inline-flex; align-items:center; gap:6px;
    padding:5px 14px; border-radius:6px; border:1px solid rgba(128,128,128,0.35);
    background:transparent; cursor:pointer; font-size:13px;
    color:inherit; font-family:inherit; margin-bottom:4px;
    transition:background 0.15s;
}}
#la-copy-btn:hover {{ background:rgba(128,128,128,0.12); }}
</style>
<div style="text-align:right; padding:4px 0 0 0;">
  <button id="la-copy-btn" onclick="(function(){{
    navigator.clipboard.writeText({json.dumps(response)}).then(function(){{
      var b=document.getElementById('la-copy-btn');
      b.innerHTML='&#10003;&nbsp;Copied!';
      b.style.color='#16a34a';
      setTimeout(function(){{b.innerHTML='&#128203;&nbsp;Copy response';b.style.color=''}},2200);
    }}).catch(function(){{
      var b=document.getElementById('la-copy-btn');
      b.innerHTML='Copy not supported in this browser';
    }});
  }})()">&#128203;&nbsp;Copy response</button>
</div>"""
        st.html(_copy_html)

        # ── CASE STRENGTH METER ──
        if st.session_state.get("last_task") in ("analysis", "advisory", "contract_review"):
            with st.expander("📊 Case Strength Meter", expanded=True):
                st.caption("AI-assessed win probability per party based on the analysis above.")
                if st.button("⚡ Generate Strength Assessment", key="strength_meter_btn", type="primary"):
                    strength_prompt = f"""
Based on this legal analysis, extract ALL parties mentioned and estimate each party's
litigation strength as a percentage.
Respond ONLY in this exact JSON format, nothing else:
{{
  "parties": [
    {{"name": "Party Name", "role": "Claimant/Defendant/Third Party", "strength": 75, "reason": "One sentence why"}},
    {{"name": "Party Name", "role": "Defendant", "strength": 35, "reason": "One sentence why"}}
  ],
  "overall_complexity": "Low/Medium/High/Extreme",
  "recommended_action": "One sentence immediate action"
}}
ANALYSIS:
{response[:6000]}
"""
                    with st.spinner("Calculating case strength..."):
                        raw = generate(strength_prompt, IDENTITY_CORE, "brief", "analysis")
                    try:
                        clean = raw.strip().replace("```json","").replace("```","").strip()
                        data = json.loads(clean)
                        for p in data.get("parties", []):
                            strength = int(p.get("strength", 50))
                            color = "#dc2626" if strength < 40 else ("#f59e0b" if strength < 65 else "#059669")
                            bar_html = f"""
<div style="margin-bottom:1rem;">
  <div style="display:flex;justify-content:space-between;margin-bottom:4px;">
    <strong>{esc(p['name'])}</strong>
    <span class="badge badge-info">{esc(p['role'])}</span>
    <strong style="color:{color};">{strength}%</strong>
  </div>
  <div style="background:#e5e7eb;border-radius:999px;height:14px;">
    <div style="width:{strength}%;background:{color};height:14px;border-radius:999px;"></div>
  </div>
  <small style="color:#6b7280;">{esc(p.get('reason',''))}</small>
</div>"""
                            st.markdown(bar_html, unsafe_allow_html=True)
                        st.markdown(f"**Complexity:** `{data.get('overall_complexity','—')}`")
                        st.markdown(f"**Immediate Action:** {esc(data.get('recommended_action','—'))}")
                    except Exception:
                        st.markdown(raw)

                # ── STRATEGY SIMULATOR (inside same expander) ──
                st.markdown("---")
                st.markdown("#### 🎯 Strategy Simulator — *What If We Do X?*")
                st.caption("Simulate any litigation move and get AI probability, risks, and opponent counter-strategy.")

                sim_cols = st.columns([3, 1])
                with sim_cols[0]:
                    sim_action = st.text_input(
                        "Proposed Action",
                        placeholder="e.g. File a preliminary objection challenging jurisdiction",
                        key="sim_action_inp",
                        label_visibility="collapsed",
                    )
                with sim_cols[1]:
                    sim_btn = st.button(
                        "🎯 Simulate",
                        key="sim_run_btn",
                        type="primary",
                        width='stretch',
                        disabled=not sim_action.strip(),
                    )

                # Quick action buttons
                st.caption("Quick simulations:")
                qa1, qa2, qa3, qa4 = st.columns(4)
                with qa1:
                    if st.button("Preliminary Objection", key="qa1_btn", width='stretch'):
                        st.session_state["sim_prefill"] = "File a preliminary objection challenging the court's jurisdiction"
                        st.rerun()
                with qa2:
                    if st.button("Strike Out Application", key="qa2_btn", width='stretch'):
                        st.session_state["sim_prefill"] = "File an application to strike out the suit for want of locus standi"
                        st.rerun()
                with qa3:
                    if st.button("Interlocutory Injunction", key="qa3_btn", width='stretch'):
                        st.session_state["sim_prefill"] = "Apply for an interlocutory injunction to preserve the subject matter"
                        st.rerun()
                with qa4:
                    if st.button("Settlement Offer", key="qa4_btn", width='stretch'):
                        st.session_state["sim_prefill"] = "Make a without-prejudice settlement offer to the opposing party"
                        st.rerun()

                # Apply prefill if set
                if st.session_state.get("sim_prefill"):
                    sim_action = st.session_state.pop("sim_prefill")

                if sim_btn and sim_action.strip():
                    sim_prompt = f"""
You are a senior Nigerian litigation strategist. A lawyer is considering the following
litigation action in the case described below. Analyse it fully.

Respond ONLY in this exact JSON format, nothing else:
{{
  "action": "The proposed action",
  "probability_of_success": 72,
  "verdict": "RECOMMENDED/RISKY/DO NOT PROCEED",
  "reasoning": "2-3 sentences explaining the probability",
  "risks": [
    "Risk 1",
    "Risk 2",
    "Risk 3"
  ],
  "opponent_counter_strategy": [
    "What opponent will likely do in response 1",
    "What opponent will likely do in response 2"
  ],
  "our_counter_to_counter": [
    "How we neutralise opponent response 1",
    "How we neutralise opponent response 2"
  ],
  "better_alternative": "A better action to consider, or empty string if this is already optimal",
  "nigerian_authority": "The most relevant Nigerian case or statute supporting or opposing this action"
}}

CASE ANALYSIS CONTEXT:
{response[:5000]}

PROPOSED ACTION: {sim_action}
"""
                    with st.spinner("🎯 Simulating strategy..."):
                        sim_raw = generate(sim_prompt, IDENTITY_CORE, "brief", "advisory")
                    try:
                        sim_clean = sim_raw.strip().replace("```json","").replace("```","").strip()
                        sim_data = json.loads(sim_clean)

                        prob = int(sim_data.get("probability_of_success", 50))
                        verdict = sim_data.get("verdict", "RISKY")

                        if verdict == "RECOMMENDED":
                            verdict_color = "#059669"
                            verdict_bg = "#f0fdf4"
                            verdict_icon = "✅"
                        elif verdict == "DO NOT PROCEED":
                            verdict_color = "#dc2626"
                            verdict_bg = "#fef2f2"
                            verdict_icon = "🚫"
                        else:
                            verdict_color = "#d97706"
                            verdict_bg = "#fffbeb"
                            verdict_icon = "⚠️"

                        prob_color = "#dc2626" if prob < 40 else ("#f59e0b" if prob < 65 else "#059669")

                        st.markdown(f"""
<div style="background:{verdict_bg};border:2px solid {verdict_color};
border-radius:0.75rem;padding:1.2rem;margin-top:1rem;">
  <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.8rem;">
    <h4 style="margin:0;color:{verdict_color};">
      {verdict_icon} {esc(sim_data.get('action',''))}
    </h4>
    <span style="font-size:1.6rem;font-weight:800;color:{prob_color};">{prob}%</span>
  </div>
  <div style="background:#e5e7eb;border-radius:999px;height:12px;margin-bottom:0.8rem;">
    <div style="width:{prob}%;background:{prob_color};height:12px;border-radius:999px;"></div>
  </div>
  <p style="margin:0;">{esc(sim_data.get('reasoning',''))}</p>
</div>""", unsafe_allow_html=True)

                        sr1, sr2 = st.columns(2)
                        with sr1:
                            st.markdown("**🔴 Risks:**")
                            for r in sim_data.get("risks", []):
                                st.markdown(f"- {esc(r)}")
                            st.markdown("**⚔️ Opponent Will:**")
                            for c in sim_data.get("opponent_counter_strategy", []):
                                st.markdown(f"- {esc(c)}")
                        with sr2:
                            st.markdown("**🛡️ Our Counter:**")
                            for cc in sim_data.get("our_counter_to_counter", []):
                                st.markdown(f"- {esc(cc)}")
                            if sim_data.get("nigerian_authority"):
                                st.markdown(f"**📖 Authority:** {esc(sim_data['nigerian_authority'])}")

                        if sim_data.get("better_alternative"):
                            st.info(f"💡 **Better Alternative:** {sim_data['better_alternative']}")

                        # Save simulation to case history
                        if st.session_state.cases:
                            sim_text = (
                                f"STRATEGY SIMULATION\n"
                                f"Action: {sim_data.get('action','')}\n"
                                f"Probability: {prob}%\n"
                                f"Verdict: {verdict}\n"
                                f"Reasoning: {sim_data.get('reasoning','')}\n"
                            )
                            add_to_history(
                                f"[Strategy Sim] {sim_action[:80]}",
                                sim_text, "advisory", "brief",
                            )

                    except Exception:
                        st.markdown(sim_raw)

        # ── SAVE TO CASE ──
        cases = st.session_state.cases
        if cases:
            st.markdown("### 💾 Save to Case")
            stc1, stc2 = st.columns([3, 1])
            with stc1:
                case_names = [f"{c.get('title', 'Untitled')} ({c.get('suit_no', '—')})" for c in cases]
                selected_case = st.selectbox(
                    "Select case to attach this analysis:",
                    case_names, key="save_to_case_sel", label_visibility="collapsed",
                )
            with stc2:
                if st.button("💾 Save", key="save_to_case_btn", type="primary", width='stretch'):
                    case_idx = case_names.index(selected_case)
                    target_case = cases[case_idx]
                    save_analysis_to_case(
                        target_case["id"],
                        st.session_state.original_query,
                        response,
                        st.session_state.get("last_task", "general"),
                        st.session_state.get("last_mode", "standard"),
                    )
                    st.success(f"✅ Analysis saved to case: {target_case.get('title', '')}")

        # Quality critique
        if mode in ("standard", "comprehensive"):
            with st.expander("🔎 Quality Assessment", expanded=False):
                if st.button("Run Critique", key="run_critique_btn"):
                    with st.spinner("Assessing quality…"):
                        critique = run_critique(st.session_state.original_query, response)
                    st.markdown(f'<div class="response-box">{esc(critique)}</div>', unsafe_allow_html=True)

        # Follow-up
        st.markdown("### 🔄 Follow-Up Question")
        followup = st.text_input(
            "Ask a follow-up based on the analysis above:",
            placeholder="E.g.: 'What if the contract had an arbitration clause?'",
            key="followup_input",
        )
        if st.button("🔄 Follow Up", disabled=not followup.strip(), key="followup_btn"):
            with st.spinner("🔄 Processing follow-up…"):
                fu_result = run_followup(
                    st.session_state.original_query,
                    response, followup.strip(), mode,
                )
            st.session_state.last_response = fu_result
            add_to_history(f"[Follow-up] {followup.strip()}", fu_result, "general", mode)
            st.rerun()

        st.markdown('<div class="disclaimer"><strong>⚖️ Disclaimer:</strong> AI-generated legal analysis. This does not constitute legal advice. Verify all citations and authorities independently before reliance.</div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: LEGAL RESEARCH
# ═══════════════════════════════════════════════════════
def render_research():
    st.markdown("""<div class="page-header">
        <h2>📚 Legal Research</h2>
        <p>Case law · Statutes · Authorities · Research Memoranda</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    mode = st.session_state.response_mode
    mode_info = RESPONSE_MODES[mode]
    st.info(f"**Research Mode: {mode_info['label']}** — {mode_info['desc']}")

    query = st.text_area(
        "🔍 Research Query", height=140,
        placeholder="E.g.: 'What are the grounds for setting aside an arbitral award under the Arbitration and Mediation Act 2023?'",
        key="research_query_ta",
    )

    # ── Quick Precedent Finder ──
    with st.expander("🔖 Quick Precedent Finder", expanded=False):
        st.caption("Get the top 5 most relevant Nigerian cases on any legal point — instantly.")
        prec_cols = st.columns([3, 1])
        with prec_cols[0]:
            prec_query = st.text_input(
                "Legal Issue",
                placeholder="e.g. unlawful termination of employment, right of pre-emption in land law",
                key="prec_query_inp",
                label_visibility="collapsed",
            )
        with prec_cols[1]:
            prec_btn = st.button(
                "🔖 Find Cases",
                key="prec_btn",
                disabled=not prec_query.strip(),
                width='stretch',
                type="primary",
            )
        if prec_btn and prec_query.strip():
            prec_prompt = f"""
You are a Nigerian law librarian. For the legal issue below, provide the TOP 5 most
authoritative Nigerian cases. Respond ONLY in this exact JSON format, nothing else:
{{
  "cases": [
    {{
      "name": "Full case name",
      "citation": "[(year)] volume report page",
      "court": "Supreme Court/Court of Appeal/Federal High Court",
      "year": "1995",
      "ratio": "One sentence — the exact legal principle established",
      "relevance": "One sentence — why this case applies to the issue"
    }}
  ]
}}
LEGAL ISSUE: {prec_query}
"""
            with st.spinner("🔖 Searching Nigerian precedents..."):
                raw = generate(prec_prompt, IDENTITY_CORE, "brief", "research")
            try:
                clean = raw.strip().replace("```json", "").replace("```", "").strip()
                data = json.loads(clean)
                for i, case in enumerate(data.get("cases", []), 1):
                    court = case.get("court", "")
                    if "Supreme" in court:
                        court_badge = "badge-err"
                    elif "Appeal" in court:
                        court_badge = "badge-warn"
                    else:
                        court_badge = "badge-ok"
                    st.markdown(f"""
<div class="custom-card">
  <div style="display:flex;justify-content:space-between;align-items:flex-start;">
    <h4 style="margin:0;">#{i} · {esc(case.get('name',''))}</h4>
    <span class="badge {court_badge}">{esc(court)}</span>
  </div>
  <div style="margin:0.4rem 0;">
    📖 <code>{esc(case.get('citation',''))}</code> · 📅 {esc(case.get('year',''))}
  </div>
  <div><strong>Ratio:</strong> {esc(case.get('ratio',''))}</div>
  <div style="color:#6b7280;">
    <strong>Why relevant:</strong> {esc(case.get('relevance',''))}
  </div>
</div>""", unsafe_allow_html=True)
            except Exception:
                st.markdown(raw)
    st.markdown("---")
    rc1, rc2 = st.columns([1, 1])
    with rc1:
        research_btn = st.button(
            f"📚 Research ({mode_info['label']})",
            type="primary", width='stretch',
            disabled=not query.strip(), key="research_go_btn",
        )
    with rc2:
        clear_btn = st.button("🗑️ Clear Results", width='stretch', key="research_clear_btn")

    if clear_btn:
        st.session_state.research_results = ""
        st.rerun()

    if research_btn and query.strip():
        with st.spinner("📚 Researching…"):
            start_t = time.time()
            result = run_research(query.strip(), mode)
            elapsed = time.time() - start_t
        st.session_state.research_results = result
        add_to_history(f"[Research] {query.strip()}", result, "research", mode)
        st.caption(f"⏱️ {elapsed:.1f}s · {len(result.split()):,} words")

    result = st.session_state.research_results
    if result:
        st.markdown("---")
        fname = f"LexiAssist_Research_{datetime.now():%Y%m%d_%H%M}"
        ex1, ex2, ex3, ex4 = st.columns(4)
        with ex1:
            st.download_button("📥 TXT", export_txt(result, "Legal Research"), f"{fname}.txt", "text/plain", key="res_dl_txt", width='stretch')
        with ex2:
            st.download_button("📥 HTML", export_html(result, "Legal Research"), f"{fname}.html", "text/html", key="res_dl_html", width='stretch')
        with ex3:
            safe_pdf_download(result, "Legal Research", fname, "res_dl_pdf")
        with ex4:
            safe_docx_download(result, "Legal Research", fname, "res_dl_docx", doc_type="research")

        st.markdown(f'<div class="response-box">{esc(result)}</div>', unsafe_allow_html=True)

        # Save research to case
        cases = st.session_state.cases
        if cases:
            st.markdown("### 💾 Save to Case")
            stc1, stc2 = st.columns([3, 1])
            with stc1:
                case_names_r = [f"{c.get('title', 'Untitled')} ({c.get('suit_no', '—')})" for c in cases]
                sel_case_r = st.selectbox("Select case:", case_names_r, key="res_save_case_sel", label_visibility="collapsed")
            with stc2:
                if st.button("💾 Save", key="res_save_case_btn", type="primary", width='stretch'):
                    cidx = case_names_r.index(sel_case_r)
                    target = cases[cidx]
                    save_analysis_to_case(target["id"], f"[Research] {query.strip()}", result, "research", mode)
                    st.success(f"✅ Research saved to case: {target.get('title', '')}")

        st.markdown('<div class="disclaimer"><strong>⚖️ Disclaimer:</strong> AI-generated research. Verify all citations independently.</div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# END OF PART 2 — Continue with Part 3 below this line
# ═══════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════
# PART 3: Cases, Calendar, Templates (CRUD), Clients,
#          Billing (+ Cost Tracker), Tools (editable),
#          Profile, and main() entry point
# ═══════════════════════════════════════════════════════


# ═══════════════════════════════════════════════════════
# PAGE: CASES (WITH SAVED ANALYSES)
# ═══════════════════════════════════════════════════════
def render_cases():
    st.markdown("""<div class="page-header">
        <h2>📁 Case Manager</h2>
        <p>Track cases, hearings, deadlines, suit numbers, and saved analyses</p>
    </div>""", unsafe_allow_html=True)

    tab_list, tab_add = st.tabs(["📋 All Cases", "➕ Add Case"])

    with tab_add:
        with st.form("add_case_form", clear_on_submit=True):
            st.markdown("#### ➕ New Case")
            ac1, ac2 = st.columns(2)
            with ac1:
                title = st.text_input("Case Title *", key="case_title_inp")
                suit_no = st.text_input("Suit Number", key="case_suit_inp")
                court = st.text_input("Court", key="case_court_inp")
            with ac2:
                status = st.selectbox("Status", CASE_STATUSES, key="case_status_inp")
                client_opts = ["— None —"] + [c.get("name", "?") for c in st.session_state.clients]
                client_sel = st.selectbox("Client", client_opts, key="case_client_inp")
                next_hearing = st.date_input("Next Hearing", value=None, key="case_hearing_inp")
            notes = st.text_area("Notes", height=80, key="case_notes_inp")

            if st.form_submit_button("➕ Add Case", type="primary"):
                if title.strip():
                    client_id = ""
                    if client_sel != "— None —":
                        cidx = client_opts.index(client_sel) - 1
                        if 0 <= cidx < len(st.session_state.clients):
                            client_id = st.session_state.clients[cidx]["id"]
                    add_case({
                        "title": title.strip(), "suit_no": suit_no.strip(),
                        "court": court.strip(), "status": status,
                        "client_id": client_id,
                        "next_hearing": str(next_hearing) if next_hearing else "",
                        "notes": notes.strip(),
                    })
                    st.success(f"✅ Case '{title}' added!")
                    st.rerun()
                else:
                    st.error("❌ Case title is required.")

    with tab_list:
        cases = st.session_state.cases
        if not cases:
            st.info("No cases yet. Add one in the ➕ Add Case tab.")
            return

        fc1, fc2 = st.columns([1, 2])
        with fc1:
            filt_status = st.selectbox("Filter by Status", ["All"] + CASE_STATUSES, key="case_filter_sel")
        with fc2:
            filt_search = st.text_input("🔍 Search cases", key="case_search_inp", placeholder="Title, suit number, court…")

        filtered = cases
        if filt_status != "All":
            filtered = [c for c in filtered if c.get("status") == filt_status]
        if filt_search.strip():
            s = filt_search.strip().lower()
            filtered = [c for c in filtered if s in c.get("title", "").lower() or s in c.get("suit_no", "").lower() or s in c.get("court", "").lower()]

        st.caption(f"Showing {len(filtered)} of {len(cases)} cases")

        for c in filtered:
            d = days_until(c.get("next_hearing", ""))
            badge = "badge-err" if d <= 3 else ("badge-warn" if d <= 7 else "badge-ok")
            hearing_txt = fmt_date(c.get("next_hearing", ""))
            cname = get_client_name(c.get("client_id", ""))

            st.markdown(f"""<div class="custom-card">
                <h4>{esc(c.get('title', 'Untitled'))}</h4>
                <span class="badge badge-info">{esc(c.get('status', ''))}</span>
                Suit: <strong>{esc(c.get('suit_no', '—'))}</strong> ·
                Court: {esc(c.get('court', '—'))} ·
                Client: {esc(cname)} ·
                Hearing: {esc(hearing_txt)}
                <span class="badge {badge}">{esc(relative_date(c.get('next_hearing', '')))}</span>
            </div>""", unsafe_allow_html=True)

            with st.expander(f"✏️ Manage: {c.get('title', '')[:50]}", expanded=False):
                manage_tab, analyses_tab = st.tabs(["⚙️ Details", "📎 Saved Analyses"])

                with manage_tab:
                    mc1, mc2 = st.columns(2)
                    with mc1:
                        new_status = st.selectbox(
                            "Status", CASE_STATUSES,
                            index=CASE_STATUSES.index(c["status"]) if c.get("status") in CASE_STATUSES else 0,
                            key=f"cs_{c['id']}",
                        )
                        new_hearing = st.date_input("Hearing", value=None, key=f"ch_{c['id']}")
                        new_notes = st.text_area("Notes", value=c.get("notes", ""), height=60, key=f"cn_{c['id']}")
                        if st.button("💾 Save Changes", key=f"save_{c['id']}", width='stretch'):
                            upd = {"status": new_status, "notes": new_notes}
                            if new_hearing:
                                upd["next_hearing"] = str(new_hearing)
                            update_case(c["id"], upd)
                            st.success("✅ Updated!")
                            st.rerun()
                    with mc2:
                        st.markdown(f"**Created:** {esc(fmt_date(c.get('created_at', '')))}")
                        if c.get("updated_at"):
                            st.markdown(f"**Updated:** {esc(fmt_date(c['updated_at']))}")
                        if c.get("notes"):
                            st.caption(f"📝 {c['notes'][:300]}")
                        st.markdown("")
                        if st.button("🗑️ Delete Case", key=f"del_{c['id']}", type="secondary", width='stretch'):
                            delete_case(c["id"])
                            st.success("✅ Deleted!")
                            st.rerun()

                with analyses_tab:
                    db = get_db()
                    saved = db.get_case_analyses(c["id"])
                    if saved:
                        st.caption(f"{len(saved)} saved analysis(es) for this case")
                        for sa in saved:
                            task_lbl = TASK_TYPES.get(sa.get("task", ""), {}).get("label", sa.get("task", ""))
                            mode_lbl = RESPONSE_MODES.get(sa.get("mode", ""), {}).get("label", sa.get("mode", ""))
                            st.markdown(f"""<div class="history-item">
                                <strong>{esc(sa.get('query', '')[:120])}</strong><br>
                                <small>{esc(fmt_date(sa.get('timestamp', '')))} · {esc(task_lbl)} · {esc(mode_lbl)}</small>
                            </div>""", unsafe_allow_html=True)

                            sa_view, sa_export, sa_del = st.columns([2, 2, 1])
                            with sa_view:
                                if st.button("👁️ View", key=f"view_sa_{sa['id']}", width='stretch'):
                                    st.markdown(f'<div class="response-box">{esc(sa["response"])}</div>', unsafe_allow_html=True)
                            with sa_export:
                                sa_fname = f"Case_Analysis_{sa['id']}"
                                st.download_button(
                                    "📥 TXT", export_txt(sa["response"], f"Case Analysis — {c.get('title', '')}"),
                                    f"{sa_fname}.txt", "text/plain",
                                    key=f"sa_dl_{sa['id']}", width='stretch',
                                )
                            with sa_del:
                                if st.button("🗑️", key=f"del_sa_{sa['id']}", width='stretch', help="Delete this analysis"):
                                    db.delete_case_analysis(sa["id"])
                                    st.success("Deleted!")
                                    st.rerun()
                    else:
                        st.info("No analyses saved to this case yet. Use 'Save to Case' in the AI Assistant or Research tab.")


# ═══════════════════════════════════════════════════════
# PAGE: CALENDAR
# ═══════════════════════════════════════════════════════
def render_calendar():
    st.markdown("""<div class="page-header">
        <h2>📅 Hearing Calendar</h2>
        <p>Upcoming hearings and deadlines at a glance</p>
    </div>""", unsafe_allow_html=True)

    hearings = get_hearings()
    if not hearings:
        st.info("No upcoming hearings. Add cases with hearing dates in the Case Manager.")
        return

    overdue = [h for h in hearings if days_until(h["date"]) < 0]
    today_h = [h for h in hearings if days_until(h["date"]) == 0]
    week_h = [h for h in hearings if 0 < days_until(h["date"]) <= 7]

    sc1, sc2, sc3, sc4 = st.columns(4)
    with sc1:
        st.metric("Total Hearings", len(hearings))
    with sc2:
        st.metric("⚠️ Overdue", len(overdue))
    with sc3:
        st.metric("📍 Today", len(today_h))
    with sc4:
        st.metric("This Week", len(week_h))

    st.markdown("---")

    for h in hearings:
        d = days_until(h["date"])
        if d < 0:
            badge_class, border_color = "badge-err", "#dc2626"
        elif d <= 3:
            badge_class, border_color = "badge-err", "#dc2626"
        elif d <= 7:
            badge_class, border_color = "badge-warn", "#f59e0b"
        else:
            badge_class, border_color = "badge-ok", "#059669"

        st.markdown(f"""<div class="custom-card" style="border-left: 4px solid {border_color};">
            <h4>{esc(h['title'])}</h4>
            Suit: <strong>{esc(h['suit'])}</strong> · Court: {esc(h['court'])} · Status: {esc(h['status'])}<br>
            📅 <strong>{esc(fmt_date(h['date']))}</strong>
            <span class="badge {badge_class}">{esc(relative_date(h['date']))}</span>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: TEMPLATES (FULL CRUD)
# ═══════════════════════════════════════════════════════
def render_templates():
    st.markdown("""<div class="page-header">
        <h2>📋 Document Templates</h2>
        <p>Built-in and custom Nigerian legal document templates</p>
    </div>""", unsafe_allow_html=True)

    tab_browse, tab_add, tab_manage = st.tabs(["📄 Browse Templates", "➕ Add Custom", "⚙️ Manage Custom"])

    all_templates = get_all_templates()

    with tab_browse:
        cats = sorted(set(t["cat"] for t in all_templates))
        sel_cat = st.selectbox("Filter by Category", ["All"] + cats, key="tmpl_cat_sel")

        templates = all_templates if sel_cat == "All" else [t for t in all_templates if t["cat"] == sel_cat]

        for t in templates:
            is_builtin = t.get("builtin", False)
            badge_html = '<span class="badge badge-ok">Built-in</span>' if is_builtin else '<span class="badge badge-info">Custom</span>'
            st.markdown(f"""<div class="custom-card">
                <h4>{esc(t['name'])}</h4>
                <span class="badge badge-info">{esc(t['cat'])}</span> {badge_html}
            </div>""", unsafe_allow_html=True)

            tc1, tc2, tc3 = st.columns(3)
            with tc1:
                if st.button("👁️ Preview", key=f"prev_t_{t['id']}", width='stretch'):
                    st.code(t["content"], language=None)
            with tc2:
                if st.button("📋 Load to AI", key=f"load_t_{t['id']}", width='stretch'):
                    st.session_state.loaded_template = t["content"]
                    st.success(f"✅ '{t['name']}' loaded! Go to AI Assistant tab.")
            with tc3:
                st.download_button(
                    "📥 Download", t["content"],
                    f"{t['name'].replace(' ', '_')}.txt", "text/plain",
                    key=f"dl_t_{t['id']}", width='stretch',
                )

    with tab_add:
        st.markdown("#### ➕ Create Custom Template")
        with st.form("add_template_form", clear_on_submit=True):
            tmpl_name = st.text_input("Template Name *", key="tmpl_name_inp")
            tmpl_cat = st.text_input("Category *", placeholder="e.g. Corporate, Litigation, Property", key="tmpl_cat_inp")
            tmpl_content = st.text_area("Template Content *", height=300,
                                        placeholder="Type your template here.\nUse [PLACEHOLDER] for variable fields.",
                                        key="tmpl_content_inp")

            if st.form_submit_button("➕ Add Template", type="primary"):
                if tmpl_name.strip() and tmpl_cat.strip() and tmpl_content.strip():
                    new_tmpl = {
                        "id": f"custom_{new_id()}",
                        "name": tmpl_name.strip(),
                        "cat": tmpl_cat.strip(),
                        "content": tmpl_content.strip(),
                        "builtin": False,
                        "created_at": datetime.now().isoformat(),
                    }
                    st.session_state.custom_templates.append(new_tmpl)
                    persist("custom_templates")
                    st.success(f"✅ Template '{tmpl_name}' created!")
                    st.rerun()
                else:
                    st.error("❌ All fields are required.")

    with tab_manage:
        custom = st.session_state.custom_templates
        if not custom:
            st.info("No custom templates yet. Add one in the ➕ Add Custom tab.")
            return

        st.caption(f"{len(custom)} custom template(s)")
        for i, t in enumerate(custom):
            st.markdown(f"""<div class="custom-card">
                <h4>{esc(t['name'])}</h4>
                <span class="badge badge-info">{esc(t['cat'])}</span>
                <span class="badge badge-info">Custom</span>
                <small> · Created: {esc(fmt_date(t.get('created_at', '')))}</small>
            </div>""", unsafe_allow_html=True)

            with st.expander(f"✏️ Edit / Delete: {t['name']}", expanded=False):
                edit_name = st.text_input("Name", value=t["name"], key=f"et_name_{t['id']}")
                edit_cat = st.text_input("Category", value=t["cat"], key=f"et_cat_{t['id']}")
                edit_content = st.text_area("Content", value=t["content"], height=200, key=f"et_content_{t['id']}")

                ec1, ec2 = st.columns(2)
                with ec1:
                    if st.button("💾 Save Changes", key=f"et_save_{t['id']}", width='stretch'):
                        st.session_state.custom_templates[i]["name"] = edit_name.strip()
                        st.session_state.custom_templates[i]["cat"] = edit_cat.strip()
                        st.session_state.custom_templates[i]["content"] = edit_content.strip()
                        st.session_state.custom_templates[i]["updated_at"] = datetime.now().isoformat()
                        persist("custom_templates")
                        st.success("✅ Template updated!")
                        st.rerun()
                with ec2:
                    if st.button("🗑️ Delete Template", key=f"et_del_{t['id']}", type="secondary", width='stretch'):
                        st.session_state.custom_templates.pop(i)
                        persist("custom_templates")
                        st.success("✅ Deleted!")
                        st.rerun()


# ═══════════════════════════════════════════════════════
# PAGE: CLIENTS
# ═══════════════════════════════════════════════════════
def render_clients():
    st.markdown("""<div class="page-header">
        <h2>👥 Client Manager</h2>
        <p>Manage your client database and track engagement</p>
    </div>""", unsafe_allow_html=True)

    tab_list, tab_add = st.tabs(["👥 All Clients", "➕ Add Client"])

    with tab_add:
        with st.form("add_client_form", clear_on_submit=True):
            st.markdown("#### ➕ New Client")
            cc1, cc2 = st.columns(2)
            with cc1:
                name = st.text_input("Client Name *", key="cl_name_inp")
                email = st.text_input("Email", key="cl_email_inp")
                phone = st.text_input("Phone", key="cl_phone_inp")
            with cc2:
                cl_type = st.selectbox("Type", CLIENT_TYPES, key="cl_type_inp")
                address = st.text_area("Address", height=80, key="cl_addr_inp")
            notes = st.text_input("Notes", key="cl_notes_inp")

            if st.form_submit_button("➕ Add Client", type="primary"):
                if name.strip():
                    add_client({
                        "name": name.strip(), "email": email.strip(),
                        "phone": phone.strip(), "type": cl_type,
                        "address": address.strip(), "notes": notes.strip(),
                    })
                    st.success(f"✅ Client '{name}' added!")
                    st.rerun()
                else:
                    st.error("❌ Client name is required.")

    with tab_list:
        clients = st.session_state.clients
        if not clients:
            st.info("No clients yet. Add one in the ➕ Add Client tab.")
            return

        search = st.text_input("🔍 Search clients", key="cl_search_inp", placeholder="Name, email, type…")
        filtered = clients
        if search.strip():
            s = search.strip().lower()
            filtered = [c for c in filtered if s in c.get("name", "").lower() or s in c.get("email", "").lower() or s in c.get("type", "").lower()]

        for cl in filtered:
            cc = client_case_count(cl["id"])
            bill = client_billable(cl["id"])
            st.markdown(f"""<div class="custom-card">
                <h4>{esc(cl.get('name', ''))}</h4>
                <span class="badge badge-info">{esc(cl.get('type', ''))}</span>
                📧 {esc(cl.get('email', '—'))} · 📞 {esc(cl.get('phone', '—'))}<br>
                📁 {cc} case{'s' if cc != 1 else ''} · 💰 {esc(fmt_currency(bill))}
                {f" · 📝 {esc(cl.get('notes', '')[:80])}" if cl.get('notes') else ""}
            </div>""", unsafe_allow_html=True)

            bc1, bc2 = st.columns([1, 4])
            with bc1:
                if st.button("🗑️ Delete", key=f"del_cl_{cl['id']}", width='stretch'):
                    delete_client(cl["id"])
                    st.success("✅ Deleted!")
                    st.rerun()


# ═══════════════════════════════════════════════════════
# PAGE: BILLING (WITH AI COST TRACKER)
# ═══════════════════════════════════════════════════════
def render_billing():
    st.markdown("""<div class="page-header">
        <h2>💰 Billing & Cost Tracker</h2>
        <p>Time entries, invoicing, financial reports, and AI usage costs</p>
    </div>""", unsafe_allow_html=True)

    tab_time, tab_inv, tab_report, tab_costs = st.tabs(
        ["⏱️ Time Entries", "📄 Invoices", "📊 Reports", "🤖 AI Costs"]
    )

    # ── Time Entries ──
    with tab_time:
        with st.form("add_time_form", clear_on_submit=True):
            st.markdown("#### ➕ New Time Entry")
            bt1, bt2 = st.columns(2)
            with bt1:
                cl_names = [c.get("name", "?") for c in st.session_state.clients]
                if not cl_names:
                    st.warning("Add a client first.")
                    cl_sel_b = None
                else:
                    cl_sel_b = st.selectbox("Client *", cl_names, key="bill_cl_inp")
                desc = st.text_input("Description *", key="bill_desc_inp")
            with bt2:
                hours = st.number_input("Hours *", min_value=0.0, step=0.25, key="bill_hrs_inp")
                rate = st.number_input("Rate (₦/hr) *", min_value=0.0, step=1000.0, value=50000.0, key="bill_rate_inp")
                entry_date = st.date_input("Date", key="bill_date_inp")

            if st.form_submit_button("➕ Add Entry", type="primary"):
                if cl_sel_b and desc.strip() and hours > 0:
                    cidx = cl_names.index(cl_sel_b)
                    add_time_entry({
                        "client_id": st.session_state.clients[cidx]["id"],
                        "client_name": cl_sel_b,
                        "description": desc.strip(),
                        "hours": hours, "rate": rate,
                        "date": str(entry_date),
                    })
                    st.success(f"✅ {hours}h @ {fmt_currency(rate)}/hr added!")
                    st.rerun()
                else:
                    st.error("❌ Fill all required fields.")

        entries = st.session_state.time_entries
        if entries:
            st.markdown("#### 📋 Recent Entries")
            for te in reversed(entries[-20:]):
                st.markdown(f"""<div class="custom-card">
                    <strong>{esc(te.get('description', ''))}</strong><br>
                    {esc(te.get('client_name', ''))} ·
                    {te.get('hours', 0)}h @ {esc(fmt_currency(te.get('rate', 0)))}/hr ·
                    <strong>{esc(fmt_currency(te.get('amount', 0)))}</strong> ·
                    {esc(fmt_date(te.get('date', '')))}
                </div>""", unsafe_allow_html=True)
                if st.button("🗑️", key=f"del_te_{te['id']}", help="Delete entry"):
                    delete_time_entry(te["id"])
                    st.rerun()

    # ── Invoices ──
    with tab_inv:
        st.markdown("#### 📄 Generate Invoice")
        if st.session_state.clients:
            cl_names_inv = [c.get("name", "?") for c in st.session_state.clients]
            inv_client = st.selectbox("Client", cl_names_inv, key="inv_cl_sel")
            if st.button("📄 Generate Invoice", type="primary", key="gen_inv_btn", width='stretch'):
                cidx = cl_names_inv.index(inv_client)
                cid = st.session_state.clients[cidx]["id"]
                inv = make_invoice(cid)
                if inv:
                    st.success(f"✅ Invoice {inv['invoice_no']} — {fmt_currency(inv['total'])}")
                    st.rerun()
                else:
                    st.warning("No billable entries for this client.")
        else:
            st.info("Add clients first.")

        if st.session_state.invoices:
            st.markdown("#### 📋 All Invoices")
            for inv in reversed(st.session_state.invoices):
                firm = get_firm_name()
                inv_text = (
                    f"{firm}\n\n"
                    f"INVOICE: {inv['invoice_no']}\n"
                    f"Date: {fmt_date(inv['date'])}\n"
                    f"Client: {inv['client_name']}\n"
                    f"Status: {inv['status']}\n\n"
                    f"{'='*40}\n"
                )
                for e in inv.get("entries", []):
                    inv_text += f"{e.get('description', '')} | {e.get('hours', 0)}h | {fmt_currency(e.get('amount', 0))}\n"
                inv_text += f"{'='*40}\nTOTAL: {fmt_currency(inv['total'])}\n"

                st.markdown(f"""<div class="custom-card">
                    <h4>{esc(inv['invoice_no'])}</h4>
                    {esc(inv['client_name'])} · {esc(fmt_date(inv['date']))} ·
                    <strong>{esc(fmt_currency(inv['total']))}</strong> ·
                    <span class="badge badge-info">{esc(inv['status'])}</span>
                </div>""", unsafe_allow_html=True)

                ic1, ic2, ic3 = st.columns(3)
                with ic1:
                    st.download_button("📥 TXT", export_txt(inv_text, f"Invoice {inv['invoice_no']}"),
                                       f"Invoice_{inv['invoice_no']}.txt", "text/plain",
                                       key=f"inv_txt_{inv['id']}", width='stretch')
                with ic2:
                    safe_pdf_download(inv_text, f"Invoice {inv['invoice_no']}",
                                      f"Invoice_{inv['invoice_no']}", f"inv_pdf_{inv['id']}")
                with ic3:
                    safe_docx_download(inv_text, f"Invoice {inv['invoice_no']}",
                                       f"Invoice_{inv['invoice_no']}", f"inv_docx_{inv['id']}",
                                       doc_type="invoice",
                                       meta={"invoice_no": inv.get("invoice_no",""), "client": inv.get("client_name",""), "matter": inv.get("matter",""), "amount": f"₦{inv.get('amount',0):,.2f}"})

    # ── Billing Reports ──
    with tab_report:
        st.markdown("#### 📊 Billing Summary")
        entries = st.session_state.time_entries
        if entries:
            th = total_hours()
            tb = total_billable()
            avg = tb / th if th else 0

            rc1, rc2, rc3 = st.columns(3)
            with rc1:
                st.metric("Total Hours", f"{th:.1f}")
            with rc2:
                st.metric("Total Billable", fmt_currency(tb))
            with rc3:
                st.metric("Avg Rate/hr", fmt_currency(avg))

            if HAS_PLOTLY:
                df = pd.DataFrame(entries)
                if "client_name" in df.columns and "amount" in df.columns:
                    chart_df = df.groupby("client_name")["amount"].sum().reset_index()
                    chart_df.columns = ["Client", "Amount"]
                    fig = px.bar(chart_df, x="Client", y="Amount",
                                 title="Billable Amount by Client",
                                 color_discrete_sequence=["#059669"])
                    st.plotly_chart(fig, width='stretch')

                if "date" in df.columns and "hours" in df.columns:
                    df["date"] = pd.to_datetime(df["date"], errors="coerce")
                    time_df = df.dropna(subset=["date"]).groupby("date")["hours"].sum().reset_index()
                    if not time_df.empty:
                        fig2 = px.line(time_df, x="date", y="hours",
                                       title="Hours Over Time",
                                       color_discrete_sequence=["#059669"])
                        st.plotly_chart(fig2, width='stretch')
        else:
            st.info("No time entries to report.")

    # ── AI Cost Tracker ──
    with tab_costs:
        st.markdown("#### 🤖 AI Usage & Cost Tracker")
        db = get_db()
        summary = db.get_cost_summary()

        kc1, kc2, kc3 = st.columns(3)
        with kc1:
            st.metric("Today", f"${summary['daily_cost']:.4f}", f"{summary['daily_calls']} calls")
        with kc2:
            st.metric("This Month", f"${summary['monthly_cost']:.4f}", f"{summary['monthly_calls']} calls")
        with kc3:
            st.metric("All Time", f"${summary['total_cost']:.4f}", f"{summary['total_calls']} calls")

        st.markdown("---")

        logs = db.get_cost_logs(100)
        if logs:
            st.markdown("#### 📋 Recent API Calls")

            if HAS_PLOTLY and len(logs) > 1:
                log_df = pd.DataFrame(logs)
                log_df["timestamp"] = pd.to_datetime(log_df["timestamp"], errors="coerce")
                log_df["date"] = log_df["timestamp"].dt.date

                # Daily cost chart
                daily_df = log_df.groupby("date")["estimated_cost"].sum().reset_index()
                daily_df.columns = ["Date", "Cost ($)"]
                if len(daily_df) > 1:
                    fig_cost = px.bar(daily_df, x="Date", y="Cost ($)",
                                      title="Daily AI Cost",
                                      color_discrete_sequence=["#3b82f6"])
                    st.plotly_chart(fig_cost, width='stretch')

                # Calls by task
                if "task" in log_df.columns:
                    task_df = log_df.groupby("task").agg(
                        calls=("id", "count"),
                        total_cost=("estimated_cost", "sum")
                    ).reset_index()
                    task_df.columns = ["Task", "Calls", "Cost ($)"]
                    fig_task = px.pie(task_df, values="Calls", names="Task",
                                     title="API Calls by Task Type")
                    st.plotly_chart(fig_task, width='stretch')

                # Calls by model
                if "model" in log_df.columns:
                    model_df = log_df.groupby("model").agg(
                        calls=("id", "count"),
                        total_cost=("estimated_cost", "sum")
                    ).reset_index()
                    model_df.columns = ["Model", "Calls", "Cost ($)"]
                    st.dataframe(model_df, width='stretch', hide_index=True)

            # Log table
            st.markdown("#### 📜 Call Log")
            for log in logs[:50]:
                task_lbl = TASK_TYPES.get(log.get("task", ""), {}).get("label", log.get("task", ""))
                mode_lbl = RESPONSE_MODES.get(log.get("mode", ""), {}).get("label", log.get("mode", ""))
                st.markdown(f"""<div class="history-item">
                    <small>{esc(fmt_date(log.get('timestamp', '')))} ·
                    {esc(log.get('model', ''))} ·
                    {esc(task_lbl)} · {esc(mode_lbl)} ·
                    In: {log.get('input_chars', 0):,}c · Out: {log.get('output_chars', 0):,}c ·
                    <strong>${log.get('estimated_cost', 0):.5f}</strong></small><br>
                    <small>{esc(log.get('query_preview', '')[:100])}</small>
                </div>""", unsafe_allow_html=True)

            # Export cost logs
            if st.button("📥 Export Cost Logs (CSV)", key="export_cost_csv", width='stretch'):
                cost_df = pd.DataFrame(logs)
                csv_data = cost_df.to_csv(index=False)
                st.download_button(
                    "⬇️ Download CSV", csv_data,
                    f"lexiassist_cost_logs_{datetime.now():%Y%m%d}.csv",
                    "text/csv", key="dl_cost_csv", width='stretch',
                )
        else:
            st.info("No API calls logged yet. Use the AI Assistant to generate your first analysis.")

        st.caption(f"💡 Costs estimated at ${COST_PER_1M_INPUT}/1M input tokens + ${COST_PER_1M_OUTPUT}/1M output tokens (approx Gemini 2.5 Flash pricing).")


# ═══════════════════════════════════════════════════════
# PAGE: TOOLS (EDITABLE REFERENCES)
# ═══════════════════════════════════════════════════════
def render_tools():
    st.markdown("""<div class="page-header">
        <h2>🔧 Legal Reference Tools</h2>
        <p>Limitation periods · Court hierarchy · Legal maxims — view and customise</p>
    </div>""", unsafe_allow_html=True)

    tab_lim, tab_calc, tab_court, tab_maxim = st.tabs(
        ["⏳ Limitation Periods", "🧮 Deadline Calculator", "🏛️ Court Hierarchy", "📜 Legal Maxims"]
    )

    # ── Limitation Periods (editable) ──
    with tab_lim:
        sub_view, sub_add = st.tabs(["📋 View All", "➕ Add Custom"])

        with sub_view:
            st.markdown("#### ⏳ Limitation Periods (Nigeria)")
            all_lim = get_all_limitation_periods()
            df_lim = pd.DataFrame(all_lim)
            if not df_lim.empty:
                df_lim.columns = ["Cause of Action", "Limitation Period", "Authority"]
                st.dataframe(df_lim, width='stretch', hide_index=True)
                st.download_button(
                    "📥 Download CSV", df_lim.to_csv(index=False),
                    "limitation_periods_nigeria.csv", "text/csv", key="dl_lim_csv",
                )

            # Show custom entries with delete option
            custom_lim = st.session_state.custom_limitation_periods
            if custom_lim:
                st.markdown("---")
                st.markdown("##### ✏️ Custom Entries")
                for i, lp in enumerate(custom_lim):
                    lc1, lc2 = st.columns([5, 1])
                    with lc1:
                        st.markdown(f"""<div class="tool-card">
                            <strong>{esc(lp['cause'])}</strong> — {esc(lp['period'])}<br>
                            <small>{esc(lp['authority'])}</small>
                            <span class="badge badge-info">Custom</span>
                        </div>""", unsafe_allow_html=True)
                    with lc2:
                        if st.button("🗑️", key=f"del_lim_{i}", help="Delete this entry"):
                            st.session_state.custom_limitation_periods.pop(i)
                            persist("custom_limitation_periods")
                            st.rerun()

        with sub_add:
            st.markdown("#### ➕ Add Custom Limitation Period")
            with st.form("add_lim_form", clear_on_submit=True):
                lim_cause = st.text_input("Cause of Action *", key="lim_cause_inp")
                lim_period = st.text_input("Limitation Period *", placeholder="e.g. 6 years", key="lim_period_inp")
                lim_auth = st.text_input("Authority *", placeholder="e.g. Limitation Act, s. X", key="lim_auth_inp")
                if st.form_submit_button("➕ Add", type="primary"):
                    if lim_cause.strip() and lim_period.strip() and lim_auth.strip():
                        st.session_state.custom_limitation_periods.append({
                            "cause": lim_cause.strip(),
                            "period": lim_period.strip(),
                            "authority": lim_auth.strip(),
                        })
                        persist("custom_limitation_periods")
                        st.success("✅ Added!")
                        st.rerun()
                    else:
                        st.error("❌ All fields required.")

    
    # ── Smart Deadline Calculator ──
    with tab_calc:
        st.markdown("#### 🧮 AI Limitation Deadline Calculator")
        st.caption("Describe your case facts and the AI will compute your exact limitation deadline and days remaining.")
        calc_facts = st.text_area(
            "Case Facts",
            height=150,
            placeholder="e.g. My client was involved in a road accident on 15 March 2022 in Lagos. The negligent driver works for a government ministry. No action has been filed yet.",
            key="calc_facts_ta",
        )
        calc_btn = st.button(
            "🧮 Calculate Deadline",
            type="primary",
            disabled=not calc_facts.strip(),
            key="calc_deadline_btn",
            width='stretch',
        )
        if calc_btn and calc_facts.strip():
            calc_prompt = f"""
You are a Nigerian limitation period expert. Analyse these facts and compute ALL applicable
limitation periods. Today's date is {date.today().strftime('%d %B %Y')}.

Respond ONLY in this exact JSON format, nothing else:
{{
  "causes_of_action": [
    {{
      "cause": "Negligence/Tort",
      "limitation_period": "3 years",
      "authority": "Limitation Act Cap L16 LFN 2004, s.8(1)(b)",
      "event_date": "2022-03-15",
      "deadline_date": "2025-03-15",
      "days_remaining": 0,
      "status": "EXPIRED/URGENT/WARNING/SAFE",
      "special_notes": "Any special rule e.g. POPA notice requirement"
    }}
  ],
  "most_urgent": "Name of most urgent cause of action",
  "immediate_action": "What lawyer must do right now"
}}

FACTS: {calc_facts}
"""
            with st.spinner("⏱️ Computing limitation deadlines..."):
                raw = generate(calc_prompt, IDENTITY_CORE, "brief", "analysis")
            try:
                clean = raw.strip().replace("```json", "").replace("```", "").strip()
                data = json.loads(clean)
                causes = data.get("causes_of_action", [])
                st.markdown("---")
                for ca in causes:
                    status = ca.get("status", "SAFE")
                    days = int(ca.get("days_remaining", 0))
                    if status == "EXPIRED":
                        card_color = "#fee2e2"
                        badge_class = "badge-err"
                        icon = "🔴"
                        days_text = f"EXPIRED {abs(days)} days ago"
                    elif status == "URGENT":
                        card_color = "#fef3c7"
                        badge_class = "badge-warn"
                        icon = "🟡"
                        days_text = f"{days} days remaining"
                    elif status == "WARNING":
                        card_color = "#fefce8"
                        badge_class = "badge-warn"
                        icon = "🟠"
                        days_text = f"{days} days remaining"
                    else:
                        card_color = "#f0fdf4"
                        badge_class = "badge-ok"
                        icon = "🟢"
                        days_text = f"{days} days remaining"
                    st.markdown(f"""
<div style="background:{card_color};border-radius:0.75rem;padding:1.2rem;
margin-bottom:1rem;border:1px solid #e5e7eb;">
  <div style="display:flex;justify-content:space-between;align-items:center;">
    <h4 style="margin:0;">{icon} {esc(ca.get('cause',''))}</h4>
    <span class="badge {badge_class}">{esc(days_text)}</span>
  </div>
  <div style="margin-top:0.5rem;">
    ⏳ <strong>Limitation Period:</strong> {esc(ca.get('limitation_period',''))}
    &nbsp;|&nbsp;
    📅 <strong>Deadline:</strong> {esc(ca.get('deadline_date',''))}
  </div>
  <div>📖 <strong>Authority:</strong> {esc(ca.get('authority',''))}</div>
  {f"<div>⚠️ <strong>Note:</strong> {esc(ca.get('special_notes',''))}</div>"
    if ca.get('special_notes') else ""}
</div>""", unsafe_allow_html=True)
                st.error(f"🚨 Most Urgent: **{data.get('most_urgent', '')}**")
                st.warning(f"⚡ Immediate Action: {data.get('immediate_action', '')}")
            except Exception:
                st.markdown(raw)
    # ── PRE-ACTION NOTICE CHECKER (merged into same tab) ──
        st.markdown("---")
        st.markdown("#### ⚠️ Pre-Action Notice & Compliance Checker")
        st.caption(
            "Find out exactly what you must do BEFORE filing suit — "
            "notices, time gaps, letters, and statutory requirements. "
            "Missing these kills cases before they start."
        )

        pre_facts = st.text_area(
            "Case Facts for Pre-Action Check",
            height=130,
            key="pre_action_facts_ta",
            placeholder="""e.g. Client wants to sue the Lagos State Government
for wrongful termination of a contract worth ₦50M.
The contract was terminated in January 2024.
No pre-action steps have been taken yet.""",
        )

        pre_btn = st.button(
            "⚠️ Check Pre-Action Requirements",
            type="primary",
            disabled=not pre_facts.strip(),
            key="pre_action_btn",
            width='stretch',
        )

        if pre_btn and pre_facts.strip():
            pre_prompt = f"""
You are a senior Nigerian litigation lawyer. Analyse the facts below and
identify ALL pre-action requirements that must be satisfied before filing
suit in Nigeria. Today's date is {date.today().strftime('%d %B %Y')}.

Respond ONLY in this exact JSON format, nothing else:
{{
  "can_sue_immediately": false,
  "overall_status": "PRE-ACTION REQUIRED / READY TO FILE / INCOMPLETE",
  "summary": "One paragraph explaining the pre-action position",
  "requirements": [
    {{
      "requirement": "Pre-Action Notice to Government",
      "authority": "Public Officers Protection Act, s.2 / Attorney General Notice",
      "is_mandatory": true,
      "deadline_to_comply": "30 days before filing",
      "action_required": "Serve statutory notice on the relevant Ministry",
      "sample_wording": "One sentence sample wording for the notice or letter",
      "consequence_of_omission": "Suit will be statute-barred / struck out",
      "status": "PENDING/DONE/NOT APPLICABLE"
    }}
  ],
  "total_waiting_period": "Total days to wait before filing e.g. 30 days",
  "earliest_filing_date": "Estimated earliest date suit can be filed",
  "immediate_actions": [
    "Action 1 to take right now",
    "Action 2 to take right now"
  ],
  "common_mistakes": [
    "Common mistake lawyers make in this type of case"
  ]
}}

CASE FACTS: {pre_facts}
"""
            with st.spinner("⚠️ Checking pre-action requirements..."):
                pre_raw = generate(
                    pre_prompt, IDENTITY_CORE, "brief", "procedure"
                )
            try:
                pre_clean = (
                    pre_raw.strip()
                    .replace("```json", "")
                    .replace("```", "")
                    .strip()
                )
                pre_data = json.loads(pre_clean)

                # ── Overall status banner ──
                overall = pre_data.get("overall_status", "PRE-ACTION REQUIRED")
                can_sue = pre_data.get("can_sue_immediately", False)

                if can_sue:
                    banner_color = "#f0fdf4"
                    banner_border = "#059669"
                    banner_icon = "✅"
                    banner_text_color = "#059669"
                else:
                    banner_color = "#fef3c7"
                    banner_border = "#f59e0b"
                    banner_icon = "⚠️"
                    banner_text_color = "#d97706"

                st.markdown(f"""
<div style="background:{banner_color};border:2px solid {banner_border};
border-radius:0.75rem;padding:1.2rem;margin:1rem 0;">
  <h4 style="margin:0;color:{banner_text_color};">
    {banner_icon} {esc(overall)}
  </h4>
  <p style="margin:0.6rem 0 0 0;">{esc(pre_data.get('summary',''))}</p>
  <div style="margin-top:0.6rem;">
    ⏳ <strong>Total waiting period:</strong>
    {esc(pre_data.get('total_waiting_period',''))} &nbsp;|&nbsp;
    📅 <strong>Earliest filing date:</strong>
    {esc(pre_data.get('earliest_filing_date',''))}
  </div>
</div>""", unsafe_allow_html=True)

                # ── Requirements ──
                reqs = pre_data.get("requirements", [])
                if reqs:
                    st.markdown(
                        f"##### 📋 {len(reqs)} Pre-Action Requirement(s)"
                    )
                    for req in reqs:
                        is_mandatory = req.get("is_mandatory", False)
                        status = req.get("status", "PENDING")

                        if status == "NOT APPLICABLE":
                            req_bg = "#f8fafc"
                            req_border = "#cbd5e1"
                            status_badge = "badge-info"
                        elif status == "DONE":
                            req_bg = "#f0fdf4"
                            req_border = "#059669"
                            status_badge = "badge-ok"
                        else:
                            req_bg = "#fef3c7"
                            req_border = "#f59e0b"
                            status_badge = "badge-warn"

                        mandatory_html = (
                            '<span class="badge badge-err">MANDATORY</span>'
                            if is_mandatory
                            else '<span class="badge badge-info">Recommended</span>'
                        )

                        st.markdown(f"""
<div style="background:{req_bg};border-left:4px solid {req_border};
border-radius:0.5rem;padding:1rem;margin-bottom:0.8rem;">
  <div style="display:flex;justify-content:space-between;
  align-items:flex-start;margin-bottom:0.4rem;">
    <strong>{esc(req.get('requirement',''))}</strong>
    <div>
      {mandatory_html}
      <span class="badge {status_badge}">{esc(status)}</span>
    </div>
  </div>
  <div>
    📖 <strong>Authority:</strong>
    <code>{esc(req.get('authority',''))}</code>
  </div>
  <div>
    ⏱️ <strong>Deadline:</strong> {esc(req.get('deadline_to_comply',''))}
  </div>
  <div>
    ✅ <strong>Action:</strong> {esc(req.get('action_required',''))}
  </div>
  {f'<div>📝 <strong>Sample wording:</strong> <em>{esc(req.get("sample_wording",""))}</em></div>'
    if req.get('sample_wording') else ''}
  <div style="color:#dc2626;">
    🚫 <strong>If omitted:</strong>
    {esc(req.get('consequence_of_omission',''))}
  </div>
</div>""", unsafe_allow_html=True)

                # ── Immediate actions ──
                immediate = pre_data.get("immediate_actions", [])
                if immediate:
                    st.markdown("##### ⚡ Immediate Actions")
                    for ia in immediate:
                        st.markdown(f"- {esc(ia)}")

                # ── Common mistakes ──
                mistakes = pre_data.get("common_mistakes", [])
                if mistakes:
                    with st.expander(
                        "🚨 Common Mistakes to Avoid", expanded=False
                    ):
                        for m in mistakes:
                            st.markdown(f"- {esc(m)}")

                # ── Export ──
                pre_report = (
                    f"PRE-ACTION COMPLIANCE REPORT\n"
                    f"Date: {datetime.now():%d %B %Y at %H:%M}\n"
                    f"Status: {overall}\n"
                    f"Earliest Filing: "
                    f"{pre_data.get('earliest_filing_date','')}\n\n"
                    f"SUMMARY:\n{pre_data.get('summary','')}\n\n"
                    f"REQUIREMENTS:\n"
                )
                for req in reqs:
                    pre_report += (
                        f"- {req.get('requirement','')} | "
                        f"{req.get('authority','')} | "
                        f"Deadline: {req.get('deadline_to_comply','')}\n"
                        f"  Action: {req.get('action_required','')}\n"
                        f"  If omitted: "
                        f"{req.get('consequence_of_omission','')}\n\n"
                    )
                if immediate:
                    pre_report += "IMMEDIATE ACTIONS:\n"
                    for ia in immediate:
                        pre_report += f"- {ia}\n"

                pre_fname = (
                    f"PreAction_Report_{datetime.now():%Y%m%d_%H%M}"
                )
                pe1, pe2, pe3 = st.columns(3)
                with pe1:
                    st.download_button(
                        "📥 TXT Report",
                        export_txt(
                            pre_report,
                            "Pre-Action Compliance Report",
                        ),
                        f"{pre_fname}.txt",
                        "text/plain",
                        key="pre_dl_txt",
                        width='stretch',
                    )
                with pe2:
                    st.download_button(
                        "📥 HTML Report",
                        export_html(
                            pre_report,
                            "Pre-Action Compliance Report",
                        ),
                        f"{pre_fname}.html",
                        "text/html",
                        key="pre_dl_html",
                        width='stretch',
                    )
                with pe3:
                    safe_pdf_download(
                        pre_report,
                        "Pre-Action Compliance Report",
                        pre_fname,
                        "pre_dl_pdf",
                    )

                st.markdown("""<div class="disclaimer">
                    <strong>⚖️ Disclaimer:</strong>
                    Pre-action requirements vary by state, court, and
                    defendant type. Always verify requirements for the
                    specific jurisdiction and court before filing.
                </div>""", unsafe_allow_html=True)

            except Exception:
                st.markdown(pre_raw)
    
    # ── Court Hierarchy ──
    with tab_court:
        st.markdown("#### 🏛️ Nigerian Court Hierarchy")
        st.caption("From the Supreme Court down to courts of first instance")
        for c in COURT_HIERARCHY:
            indent = "&nbsp;&nbsp;&nbsp;&nbsp;" * (c["level"] - 1)
            level_label = {1: "APEX", 2: "APPELLATE", 3: "SUPERIOR", 4: "LOWER"}.get(c["level"], "")
            st.markdown(f"""<div class="tool-card">
                {indent}{c['icon']} <strong>{esc(c['name'])}</strong>
                <span class="badge badge-info">{level_label}</span><br>
                {indent}&nbsp;&nbsp;&nbsp;&nbsp;<small>{esc(c['desc'])}</small>
            </div>""", unsafe_allow_html=True)

    # ── Legal Maxims (editable) ──
    with tab_maxim:
        sub_maxim_view, sub_maxim_add = st.tabs(["📋 View All", "➕ Add Custom"])

        with sub_maxim_view:
            st.markdown("#### 📜 Legal Maxims")
            search = st.text_input("🔍 Search maxims", key="maxim_search_inp", placeholder="E.g. 'nemo' or 'remedy'")
            all_maxims = get_all_maxims()
            maxims = all_maxims
            if search.strip():
                s = search.strip().lower()
                maxims = [m for m in maxims if s in m["maxim"].lower() or s in m["meaning"].lower()]

            st.caption(f"Showing {len(maxims)} maxim{'s' if len(maxims) != 1 else ''}")
            for m in maxims:
                is_custom = m not in DEFAULT_LEGAL_MAXIMS
                badge_extra = ' <span class="badge badge-info">Custom</span>' if is_custom else ""
                st.markdown(f"""<div class="tool-card">
                    <strong><em>{esc(m['maxim'])}</em></strong>{badge_extra}<br>
                    {esc(m['meaning'])}
                </div>""", unsafe_allow_html=True)

            # Manage custom maxims
            custom_maxims = st.session_state.custom_maxims
            if custom_maxims:
                st.markdown("---")
                st.markdown("##### ✏️ Manage Custom Maxims")
                for i, m in enumerate(custom_maxims):
                    mc1, mc2 = st.columns([5, 1])
                    with mc1:
                        st.caption(f"**{m['maxim']}** — {m['meaning']}")
                    with mc2:
                        if st.button("🗑️", key=f"del_maxim_{i}", help="Delete"):
                            st.session_state.custom_maxims.pop(i)
                            persist("custom_maxims")
                            st.rerun()

        with sub_maxim_add:
            st.markdown("#### ➕ Add Custom Maxim")
            with st.form("add_maxim_form", clear_on_submit=True):
                maxim_latin = st.text_input("Latin Maxim *", key="maxim_latin_inp")
                maxim_meaning = st.text_input("English Meaning *", key="maxim_meaning_inp")
                if st.form_submit_button("➕ Add Maxim", type="primary"):
                    if maxim_latin.strip() and maxim_meaning.strip():
                        st.session_state.custom_maxims.append({
                            "maxim": maxim_latin.strip(),
                            "meaning": maxim_meaning.strip(),
                        })
                        persist("custom_maxims")
                        st.success("✅ Maxim added!")
                        st.rerun()
                    else:
                        st.error("❌ Both fields required.")

# ═══════════════════════════════════════════════════════
# PAGE: CONFLICT OF INTEREST CHECKER
# ═══════════════════════════════════════════════════════
CONFLICT_PROMPT = """
You are a Nigerian legal ethics expert applying the Rules of Professional
Conduct for Legal Practitioners 2007 (RPC).

A lawyer wants to take on a new matter. Check whether any conflict of interest
exists against the existing client and case data provided.

Respond ONLY in this exact JSON format, nothing else:
{{
  "overall_verdict": "CLEAR / POTENTIAL CONFLICT / SERIOUS CONFLICT",
  "confidence": 85,
  "summary": "One paragraph summary of the conflict analysis",
  "conflicts_found": [
    {{
      "conflict_type": "Direct/Indirect/Positional/Former Client",
      "severity": "High/Medium/Low",
      "existing_party": "Name of existing client or case party",
      "new_party": "Name from new matter that conflicts",
      "reason": "Specific reason this is a conflict under RPC",
      "rpc_rule": "Specific RPC rule number e.g. Rule 17(1)"
    }}
  ],
  "recommendations": [
    "Specific recommendation 1",
    "Specific recommendation 2"
  ],
  "disclosure_required": true,
  "can_proceed_with_consent": true,
  "consent_note": "What consent is needed if proceeding, or empty string"
}}

EXISTING CLIENTS:
{existing_clients}

EXISTING CASES AND PARTIES:
{existing_cases}

NEW MATTER DETAILS:
{new_matter}
"""


def render_conflict_checker():
    st.markdown("""<div class="page-header">
        <h2>🔍 Conflict of Interest Checker</h2>
        <p>RPC-compliant conflict scanning across all clients and cases
        before accepting a new matter</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    clients = st.session_state.clients
    cases = st.session_state.cases

    # ── Stats ──
    cs1, cs2, cs3 = st.columns(3)
    with cs1:
        st.metric("Clients on Record", len(clients))
    with cs2:
        st.metric("Cases on Record", len(cases))
    with cs3:
        st.metric("Parties Indexed",
                  len(clients) + sum(
                      1 for c in cases if c.get("title")
                  ))

    if not clients and not cases:
        st.warning(
            "⚠️ No clients or cases on record yet. "
            "Add clients and cases first so the checker has data to scan against."
        )
        return

    st.markdown("---")

    # ── New matter input ──
    st.markdown("### 📋 New Matter Details")
    st.caption(
        "Enter details of the prospective new client or matter. "
        "The checker will scan all existing clients and cases for conflicts."
    )

    nc1, nc2 = st.columns(2)
    with nc1:
        new_client_name = st.text_input(
            "Prospective Client Name *",
            key="conflict_client_name",
            placeholder="e.g. Alhaji Musa Danladi",
        )
        new_opponent = st.text_input(
            "Opposing Party Name(s)",
            key="conflict_opponent",
            placeholder="e.g. Bright Ventures Ltd, Mr Chen",
        )
        new_matter_type = st.text_input(
            "Matter Type",
            key="conflict_matter_type",
            placeholder="e.g. Land dispute, Debt recovery, Employment",
        )
    with nc2:
        new_court = st.text_input(
            "Court / Tribunal",
            key="conflict_court",
            placeholder="e.g. Federal High Court Lagos",
        )
        new_related_parties = st.text_input(
            "Other Related Parties / Companies",
            key="conflict_related",
            placeholder="e.g. Parent company, directors, guarantors",
        )
        new_former_counsel = st.text_input(
            "Previous Counsel (if known)",
            key="conflict_prev_counsel",
            placeholder="e.g. ABC & Co — counsel to opponent",
        )

    new_facts = st.text_area(
        "Brief Description of New Matter",
        height=100,
        key="conflict_facts",
        placeholder="""e.g. Prospective client wants to sue Bright Ventures Ltd
for breach of a supply contract worth ₦25M. The dispute
arose in January 2024 in Lagos.""",
    )

    # ── Manual party list builder ──
    with st.expander("➕ Add Extra Parties to Scan (optional)", expanded=False):
        st.caption(
            "Add any additional names — subsidiaries, aliases, related companies — "
            "that should be checked even if not in your client database."
        )
        extra_parties = st.text_area(
            "Extra names (one per line)",
            height=80,
            key="conflict_extra_parties",
            placeholder="SinoPower Ltd\nEmeka Holdings\nMrs Chidinma Obi",
        )

    check_btn = st.button(
        "🔍 Run Conflict Check",
        type="primary",
        width='stretch',
        key="conflict_check_btn",
        disabled=not new_client_name.strip(),
    )

    if check_btn and new_client_name.strip():
        # Build existing clients string
        clients_str = ""
        for cl in clients:
            clients_str += (
                f"- {cl.get('name','')} | "
                f"Type: {cl.get('type','')} | "
                f"Email: {cl.get('email','')} | "
                f"Notes: {cl.get('notes','')}\n"
            )
        if not clients_str:
            clients_str = "No clients on record."

        # Build existing cases string
        cases_str = ""
        for c in cases:
            client_name_for_case = get_client_name(c.get("client_id", ""))
            cases_str += (
                f"- {c.get('title','')} | "
                f"Suit: {c.get('suit_no','')} | "
                f"Court: {c.get('court','')} | "
                f"Our Client: {client_name_for_case} | "
                f"Status: {c.get('status','')} | "
                f"Notes: {c.get('notes','')}\n"
            )
        if not cases_str:
            cases_str = "No cases on record."

        # Build new matter string
        extra = extra_parties.strip() if extra_parties.strip() else "None"
        new_matter_str = (
            f"Prospective Client: {new_client_name}\n"
            f"Opposing Party: {new_opponent or 'Not specified'}\n"
            f"Matter Type: {new_matter_type or 'Not specified'}\n"
            f"Court: {new_court or 'Not specified'}\n"
            f"Other Related Parties: {new_related_parties or 'None'}\n"
            f"Previous Counsel: {new_former_counsel or 'Unknown'}\n"
            f"Extra parties to scan: {extra}\n"
            f"Facts: {new_facts or 'Not provided'}\n"
        )

        prompt = CONFLICT_PROMPT.format(
            existing_clients=clients_str,
            existing_cases=cases_str,
            new_matter=new_matter_str,
        )

        with st.spinner("🔍 Scanning all clients and cases for conflicts..."):
            raw = generate(prompt, IDENTITY_CORE, "brief", "advisory")

        try:
            clean = raw.strip().replace("```json", "").replace("```", "").strip()
            data = json.loads(clean)
            st.session_state["conflict_result"] = data
            st.session_state["conflict_matter"] = new_client_name
            st.rerun()
        except Exception:
            st.markdown(raw)

    # ── Display result ──
    result = st.session_state.get("conflict_result", {})
    if not result:
        st.markdown("---")
        st.info(
            "Fill in the new matter details above and click "
            "**🔍 Run Conflict Check** to scan your entire client and case database."
        )
        return

    st.markdown("---")
    st.markdown(
        f"### 🔍 Conflict Check — "
        f"{esc(st.session_state.get('conflict_matter',''))}"
    )

    # ── Verdict banner ──
    verdict = result.get("overall_verdict", "CLEAR")
    confidence = int(result.get("confidence", 0))

    if verdict == "CLEAR":
        v_color = "#059669"
        v_bg = "#f0fdf4"
        v_icon = "✅"
        v_border = "#059669"
    elif verdict == "POTENTIAL CONFLICT":
        v_color = "#d97706"
        v_bg = "#fffbeb"
        v_icon = "⚠️"
        v_border = "#f59e0b"
    else:
        v_color = "#dc2626"
        v_bg = "#fef2f2"
        v_icon = "🚫"
        v_border = "#dc2626"

    st.markdown(f"""
<div style="background:{v_bg};border:2px solid {v_border};
border-radius:0.75rem;padding:1.4rem;margin-bottom:1.2rem;">
  <div style="display:flex;justify-content:space-between;align-items:center;">
    <h3 style="margin:0;color:{v_color};">{v_icon} {esc(verdict)}</h3>
    <span style="font-size:1.4rem;font-weight:800;color:{v_color};">
      {confidence}% confidence
    </span>
  </div>
  <p style="margin:0.8rem 0 0 0;">{esc(result.get('summary',''))}</p>
</div>""", unsafe_allow_html=True)

    # ── Conflicts found ──
    conflicts = result.get("conflicts_found", [])
    if conflicts:
        st.markdown(f"#### 🔴 {len(conflicts)} Conflict(s) Found")
        for cf in conflicts:
            sev = cf.get("severity", "Low")
            sev_color = (
                "#dc2626" if sev == "High"
                else ("#d97706" if sev == "Medium" else "#059669")
            )
            st.markdown(f"""
<div style="border-left:4px solid {sev_color};background:#fff;
border-radius:0.5rem;padding:1rem;margin-bottom:0.8rem;
box-shadow:0 1px 4px rgba(0,0,0,0.05);">
  <div style="display:flex;justify-content:space-between;">
    <strong>{esc(cf.get('conflict_type',''))} Conflict</strong>
    <span style="color:{sev_color};font-weight:700;">
      {esc(sev)} Severity
    </span>
  </div>
  <div style="margin-top:0.5rem;">
    🏢 <strong>Existing party:</strong> {esc(cf.get('existing_party',''))}
    &nbsp;↔️&nbsp;
    <strong>New party:</strong> {esc(cf.get('new_party',''))}
  </div>
  <div>📖 <strong>Reason:</strong> {esc(cf.get('reason',''))}</div>
  <div>⚖️ <strong>RPC Rule:</strong>
    <code>{esc(cf.get('rpc_rule',''))}</code>
  </div>
</div>""", unsafe_allow_html=True)
    else:
        st.success("✅ No specific conflicts identified in the database.")

    # ── Recommendations ──
    recs = result.get("recommendations", [])
    if recs:
        st.markdown("#### 💡 Recommendations")
        for rec in recs:
            st.markdown(f"- {esc(rec)}")

    # ── Consent / Disclosure ──
    st.markdown("#### 📋 Compliance Summary")
    comp1, comp2 = st.columns(2)
    with comp1:
        disc = result.get("disclosure_required", False)
        st.markdown(
            f"**Disclosure Required:** "
            f"{'🔴 YES' if disc else '🟢 NO'}"
        )
    with comp2:
        consent = result.get("can_proceed_with_consent", False)
        st.markdown(
            f"**Can Proceed with Consent:** "
            f"{'🟡 YES (with conditions)' if consent else '🔴 NO'}"
        )
    if result.get("consent_note"):
        st.info(f"📋 **Consent Note:** {result['consent_note']}")

    # ── Export conflict report ──
    st.markdown("---")
    report_text = (
        f"CONFLICT OF INTEREST CHECK REPORT\n"
        f"Matter: {st.session_state.get('conflict_matter','')}\n"
        f"Date: {datetime.now():%d %B %Y at %H:%M}\n"
        f"Verdict: {verdict} ({confidence}% confidence)\n\n"
        f"SUMMARY:\n{result.get('summary','')}\n\n"
    )
    if conflicts:
        report_text += "CONFLICTS FOUND:\n"
        for cf in conflicts:
            report_text += (
                f"- {cf.get('conflict_type','')} | "
                f"{cf.get('severity','')} | "
                f"{cf.get('existing_party','')} vs "
                f"{cf.get('new_party','')} | "
                f"{cf.get('rpc_rule','')}\n"
                f"  Reason: {cf.get('reason','')}\n"
            )
    if recs:
        report_text += "\nRECOMMENDATIONS:\n"
        for rec in recs:
            report_text += f"- {rec}\n"

    ec1, ec2, ec3 = st.columns(3)
    fname = (
        f"ConflictCheck_{st.session_state.get('conflict_matter','').replace(' ','_')}"
        f"_{datetime.now():%Y%m%d_%H%M}"
    )
    with ec1:
        st.download_button(
            "📥 TXT Report",
            export_txt(report_text, "Conflict of Interest Report"),
            f"{fname}.txt", "text/plain",
            key="conflict_dl_txt", width='stretch',
        )
    with ec2:
        st.download_button(
            "📥 HTML Report",
            export_html(report_text, "Conflict of Interest Report"),
            f"{fname}.html", "text/html",
            key="conflict_dl_html", width='stretch',
        )
    with ec3:
        safe_pdf_download(
            report_text,
            "Conflict of Interest Report",
            fname, "conflict_dl_pdf",
        )

    # ── Clear ──
    if st.button("🗑️ Clear Results", key="conflict_clear_btn",
                 width='stretch'):
        st.session_state["conflict_result"] = {}
        st.session_state["conflict_matter"] = ""
        st.rerun()

    st.markdown("""<div class="disclaimer">
        <strong>⚖️ Disclaimer:</strong> This conflict check is AI-assisted
        and supplements — but does not replace — manual conflict screening.
        Always apply your professional judgment. Rules of Professional Conduct
        for Legal Practitioners 2007 applies.
    </div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════
# PAGE: SMART PLEADINGS DRAFTER
# ═══════════════════════════════════════════════════════
PLEADING_TYPES = {
    "statement_of_claim": {
        "label": "📄 Statement of Claim",
        "desc": "Originating pleading setting out claimant's facts and reliefs",
    },
    "statement_of_defence": {
        "label": "🛡️ Statement of Defence",
        "desc": "Defendant's response admitting or denying each allegation",
    },
    "reply": {
        "label": "↩️ Reply to Defence",
        "desc": "Claimant's response to new matters raised in defence",
    },
    "counterclaim": {
        "label": "⚔️ Counter-Claim",
        "desc": "Defendant's claim against the claimant",
    },
    "originating_summons": {
        "label": "📋 Originating Summons",
        "desc": "For matters begun by summons — questions of law or construction",
    },
    "motion_on_notice": {
        "label": "📬 Motion on Notice",
        "desc": "Interlocutory application with supporting affidavit",
    },
    "affidavit": {
        "label": "📜 Supporting Affidavit",
        "desc": "Sworn statement of facts in support of a motion or application",
    },
    "written_address": {
        "label": "✍️ Written Address",
        "desc": "Final address or skeleton argument for court",
    },
    "notice_of_appeal": {
        "label": "🔔 Notice of Appeal",
        "desc": "Formal notice of appeal with grounds",
    },
    "writ_of_summons": {
        "label": "📃 Writ of Summons",
        "desc": "Originating process for High Court actions",
    },
}

PLEADING_PROMPT = """
You are a senior Nigerian litigation lawyer drafting court documents.
Draft the {pleading_type} described below in full, professional Nigerian court format.

STRICT RULES:
1. Use the exact suit number, parties, and court provided
2. Use [PLACEHOLDER] only for information not provided
3. Include all mandatory formal requirements for this document type in Nigerian courts
4. Number all paragraphs correctly
5. Include proper heading, title, body, relief/prayer section, and signature block
6. Apply the correct rules of court for the specified court
7. Do NOT add commentary or strategy — draft only

CASE DETAILS:
Case Title: {case_title}
Suit Number: {suit_no}
Court: {court}
Claimant: {claimant}
Defendant: {defendant}
Case Type: {case_type}
Key Facts: {facts}
Specific Instructions: {instructions}

Draft the complete {pleading_type} now:
"""


def render_pleadings():
    st.markdown("""<div class="page-header">
        <h2>📜 Smart Pleadings Drafter</h2>
        <p>Generate court-ready pleadings pulled directly from your case file —
        no manual typing of parties, court, or suit number</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    cases = st.session_state.cases
    if not cases:
        st.info("No cases found. Add a case in the 📁 Cases tab first — "
                "the drafter pulls parties, court, and suit number from your case file automatically.")
        return

    # ── Case selector ──
    st.markdown("### 📁 Select Case")
    st.caption("All case details are pulled automatically from your saved case file.")

    case_names = [
        f"{c.get('title', 'Untitled')} ({c.get('suit_no', '—')})"
        for c in cases
    ]
    pc1, pc2 = st.columns([3, 1])
    with pc1:
        selected_case_name = st.selectbox(
            "Choose case",
            case_names,
            key="pleading_case_sel",
            label_visibility="collapsed",
        )
    selected_idx = case_names.index(selected_case_name)
    selected_case = cases[selected_idx]

    with pc2:
        st.metric("Status", selected_case.get("status", "—"))

    # ── Auto-populated case details ──
    st.markdown("---")
    st.markdown("### 📋 Case Details (Auto-Populated)")
    st.caption("Review and edit any field before generating.")

    pd1, pd2 = st.columns(2)
    with pd1:
        case_title = st.text_input(
            "Case Title",
            value=selected_case.get("title", ""),
            key="pl_case_title",
        )
        suit_no = st.text_input(
            "Suit Number",
            value=selected_case.get("suit_no", ""),
            key="pl_suit_no",
        )
        court = st.text_input(
            "Court",
            value=selected_case.get("court", ""),
            key="pl_court",
        )
    with pd2:
        claimant = st.text_input(
            "Claimant / Applicant",
            value="",
            placeholder="e.g. Chief Emeka Obi",
            key="pl_claimant",
        )
        defendant = st.text_input(
            "Defendant / Respondent",
            value="",
            placeholder="e.g. Lagos State Government",
            key="pl_defendant",
        )
        case_type_pl = st.text_input(
            "Case Type",
            value="",
            placeholder="e.g. Breach of Contract, Land Dispute",
            key="pl_case_type",
        )

    facts = st.text_area(
        "Key Facts",
        value=selected_case.get("notes", ""),
        height=120,
        key="pl_facts",
        placeholder="""e.g. Claimant and Defendant entered into a contract on 1 Jan 2023.
Defendant received goods worth ₦12M and refused payment.
Demand letters sent on 1 March and 1 April 2023. No response.""",
    )

    # ── Pleading type selector ──
    st.markdown("---")
    st.markdown("### 📄 Select Document to Draft")

    pl_keys = list(PLEADING_TYPES.keys())
    pleading_type_key = st.selectbox(
        "Document Type",
        pl_keys,
        format_func=lambda x: f"{PLEADING_TYPES[x]['label']} — {PLEADING_TYPES[x]['desc']}",
        key="pleading_type_sel",
    )
    selected_pleading = PLEADING_TYPES[pleading_type_key]

    # Special instructions
    instructions = st.text_area(
        "Special Instructions (optional)",
        height=80,
        key="pl_instructions",
        placeholder="""e.g. Include a claim for general damages of ₦5M and special damages of ₦12M.
Add an application for accelerated hearing.
This is a counter-claim so defendant becomes counter-claimant.""",
    )

    mode = st.session_state.response_mode
    st.info(f"**Mode:** {RESPONSE_MODES[mode]['label']} — "
            f"Comprehensive mode produces the most complete pleadings.")

    # ── Generate button ──
    generate_btn = st.button(
        f"📜 Draft {selected_pleading['label']}",
        type="primary",
        width='stretch',
        key="pleading_generate_btn",
        disabled=not (case_title.strip() and court.strip()),
    )

    if generate_btn:
        prompt = PLEADING_PROMPT.format(
            pleading_type=selected_pleading["label"],
            case_title=case_title.strip(),
            suit_no=suit_no.strip() or "[SUIT NUMBER TO BE ASSIGNED]",
            court=court.strip(),
            claimant=claimant.strip() or "[CLAIMANT NAME]",
            defendant=defendant.strip() or "[DEFENDANT NAME]",
            case_type=case_type_pl.strip() or "General Civil Matter",
            facts=facts.strip() or "As will be adduced at trial",
            instructions=instructions.strip() or "None",
        )
        system = build_system_prompt("drafting", mode)
        with st.spinner(
            f"📜 Drafting {selected_pleading['label']}..."
        ):
            result = generate(prompt, system, mode, "drafting")

        st.session_state["pleading_result"] = result
        st.session_state["pleading_title"] = selected_pleading["label"]
        st.session_state["pleading_case_id"] = selected_case["id"]
        st.session_state["pleading_case_title"] = case_title
        add_to_history(
            f"[Pleading] {selected_pleading['label']} — {case_title}",
            result, "drafting", mode,
        )
        st.rerun()

    # ── Display result ──
    result = st.session_state.get("pleading_result", "")
    pleading_title = st.session_state.get("pleading_title", "Pleading")
    pleading_case_id = st.session_state.get("pleading_case_id", "")
    pleading_case_title = st.session_state.get("pleading_case_title", "")

    if result:
        st.markdown("---")
        st.markdown(f"### {pleading_title}")
        st.caption(f"Case: {esc(pleading_case_title)}")

        # ── Export row ──
        fname = (
            f"LexiAssist_{pleading_type_key}_{pleading_case_title.replace(' ','_')}"
            f"_{datetime.now():%Y%m%d_%H%M}"
        )
        ex1, ex2, ex3, ex4 = st.columns(4)
        with ex1:
            st.download_button(
                "📥 TXT",
                export_txt(result, pleading_title),
                f"{fname}.txt", "text/plain",
                key="pl_dl_txt", width='stretch',
            )
        with ex2:
            st.download_button(
                "📥 HTML",
                export_html(result, pleading_title),
                f"{fname}.html", "text/html",
                key="pl_dl_html", width='stretch',
            )
        with ex3:
            safe_pdf_download(result, pleading_title, fname, "pl_dl_pdf")
        with ex4:
            safe_docx_download(result, pleading_title, fname, "pl_dl_docx", doc_type="pleading")

        st.markdown(
            f'<div class="response-box">{esc(result)}</div>',
            unsafe_allow_html=True,
        )

        # ── Save to Case ──
        if pleading_case_id:
            sv1, sv2 = st.columns([3, 1])
            with sv1:
                st.caption(f"Save this pleading to: **{esc(pleading_case_title)}**")
            with sv2:
                if st.button(
                    "💾 Save to Case",
                    key="pl_save_case_btn",
                    type="primary",
                    width='stretch',
                ):
                    save_analysis_to_case(
                        pleading_case_id,
                        f"[{pleading_title}]",
                        result, "drafting", mode,
                    )
                    st.success(
                        f"✅ {pleading_title} saved to case: {pleading_case_title}"
                    )

        # ── Clear ──
        if st.button("🗑️ Clear Draft", key="pl_clear_btn", width='stretch'):
            st.session_state["pleading_result"] = ""
            st.session_state["pleading_title"] = ""
            st.rerun()

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> Review all AI-drafted pleadings
            carefully before filing. Verify all facts, parties, and reliefs
            against your instructions. Counsel remains responsible for all
            documents filed in court.
        </div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════
# PAGE: MATTER LIFECYCLE AUTOMATION
# ═══════════════════════════════════════════════════════
CASE_TYPE_OPTIONS = [
    "Breach of Contract",
    "Land / Property Dispute",
    "Criminal Defence",
    "Employment / Wrongful Termination",
    "Fundamental Rights Enforcement",
    "Debt Recovery",
    "Matrimonial / Family Law",
    "Company / Commercial Dispute",
    "Defamation",
    "Personal Injury / Negligence",
    "Election Petition",
    "Judicial Review / Certiorari",
    "Tenancy / Landlord-Tenant",
    "Probate / Estate Administration",
    "Other (describe below)",
]

LIFECYCLE_PROMPT = """
You are a senior Nigerian litigation lawyer. Generate a complete matter lifecycle workflow
for the case described below.

Respond ONLY in this exact JSON format, nothing else:
{{
  "case_type": "Breach of Contract",
  "court_recommendation": "Lagos State High Court",
  "estimated_duration": "12-18 months",
  "total_stages": 6,
  "stages": [
    {{
      "stage_number": 1,
      "stage_name": "Client Intake & Brief",
      "description": "One sentence describing this stage",
      "duration_estimate": "1-3 days",
      "required_documents": [
        "Instructions letter from client",
        "Copies of contract documents",
        "Evidence of breach"
      ],
      "required_actions": [
        "Obtain full instructions from client",
        "Review all contract documents",
        "Conduct conflict of interest check"
      ],
      "deadline_trigger": "Immediately on instruction",
      "warning": "Any critical warning for this stage or empty string"
    }}
  ],
  "limitation_alert": "Limitation period warning if any",
  "pre_action_requirements": [
    "Any pre-action notices required before filing"
  ],
  "top_risks": [
    "Top 3 risks for this matter"
  ],
  "immediate_next_step": "Single most important action to take right now"
}}

CASE DETAILS:
Case Type: {case_type}
Case Title: {case_title}
Court: {court}
Brief Facts: {facts}
"""


def render_lifecycle():
    st.markdown("""<div class="page-header">
        <h2>⚡ Matter Lifecycle Automation</h2>
        <p>Auto-generate complete case workflows · stages · deadlines · documents · actions</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    cases = st.session_state.cases
    if not cases:
        st.info("No cases found. Add a case in the 📁 Cases tab first, then return here to generate its lifecycle.")
        return

    db = get_db()

    # ── Case selector ──
    st.markdown("### 📁 Select a Case")
    lc1, lc2 = st.columns([3, 1])
    with lc1:
        case_names = [
            f"{c.get('title', 'Untitled')} ({c.get('suit_no', '—')})"
            for c in cases
        ]
        selected_case_name = st.selectbox(
            "Choose case",
            case_names,
            key="lifecycle_case_sel",
            label_visibility="collapsed",
        )
    selected_idx = case_names.index(selected_case_name)
    selected_case = cases[selected_idx]
    case_id = selected_case["id"]

    with lc2:
        st.metric("Status", selected_case.get("status", "—"))

    # ── Case type + facts input ──
    st.markdown("---")
    st.markdown("### ⚙️ Configure Lifecycle Generation")
    gc1, gc2 = st.columns(2)
    with gc1:
        case_type = st.selectbox(
            "Case Type *",
            CASE_TYPE_OPTIONS,
            key="lifecycle_case_type",
        )
        if case_type == "Other (describe below)":
            case_type = st.text_input(
                "Describe case type",
                key="lifecycle_custom_type",
                placeholder="e.g. Insurance claim dispute",
            )
    with gc2:
        court_input = st.text_input(
            "Court (if known)",
            value=selected_case.get("court", ""),
            key="lifecycle_court_inp",
            placeholder="e.g. Federal High Court Lagos",
        )

    facts_input = st.text_area(
        "Brief Facts (optional but improves accuracy)",
        height=120,
        key="lifecycle_facts_inp",
        placeholder="""e.g. Client entered into a supply contract with defendant in Jan 2023.
Defendant received goods worth ₦12M but refused to pay after 90 days.
Multiple demand letters sent. No response. Client wants to sue.""",
    )

    generate_btn = st.button(
        "⚡ Generate Matter Lifecycle",
        type="primary",
        width='stretch',
        key="lifecycle_generate_btn",
        disabled=not case_type,
    )

    if generate_btn:
        prompt = LIFECYCLE_PROMPT.format(
            case_type=case_type,
            case_title=selected_case.get("title", ""),
            court=court_input or "To be determined",
            facts=facts_input or "Not provided",
        )
        with st.spinner("⚡ Building matter lifecycle workflow..."):
            raw = generate(prompt, IDENTITY_CORE, "standard", "analysis")
        try:
            clean = raw.strip().replace("```json", "").replace("```", "").strip()
            lifecycle_data = json.loads(clean)
            db.save_lifecycle(case_id, lifecycle_data)
            # Initialise progress — all stages incomplete
            progress = {
                str(i): False
                for i in range(1, lifecycle_data.get("total_stages", 0) + 1)
            }
            existing_progress = db.load_lifecycle_progress(case_id)
            if not existing_progress:
                db.save_lifecycle_progress(case_id, progress)
            st.success("✅ Lifecycle generated and saved to this case!")
            st.rerun()
        except Exception:
            st.markdown(raw)

    # ── Display saved lifecycle ──
    lifecycle = db.load_lifecycle(case_id)
    if not lifecycle:
        st.info("No lifecycle generated for this case yet. Fill in the form above and click Generate.")
        return

    progress = db.load_lifecycle_progress(case_id)
    if not progress:
        progress = {
            str(i): False
            for i in range(1, lifecycle.get("total_stages", 0) + 1)
        }

    st.markdown("---")

    # ── Summary banner ──
    sc1, sc2, sc3, sc4 = st.columns(4)
    completed = sum(1 for v in progress.values() if v)
    total = lifecycle.get("total_stages", 0)
    pct = int((completed / total) * 100) if total else 0
    with sc1:
        st.metric("Case Type", lifecycle.get("case_type", "—"))
    with sc2:
        st.metric("Recommended Court", lifecycle.get("court_recommendation", "—"))
    with sc3:
        st.metric("Est. Duration", lifecycle.get("estimated_duration", "—"))
    with sc4:
        st.metric("Progress", f"{completed}/{total} stages ({pct}%)")

    # Progress bar
    st.markdown(f"""
<div style="background:#e5e7eb;border-radius:999px;height:16px;margin:0.5rem 0 1.5rem 0;">
  <div style="width:{pct}%;background:#059669;height:16px;border-radius:999px;
  transition:width 0.5s;"></div>
</div>""", unsafe_allow_html=True)

    # ── Alerts ──
    if lifecycle.get("limitation_alert"):
        st.error(f"⏳ **Limitation Alert:** {lifecycle['limitation_alert']}")
    if lifecycle.get("pre_action_requirements"):
        with st.expander("⚠️ Pre-Action Requirements (must complete before filing)", expanded=True):
            for req in lifecycle["pre_action_requirements"]:
                st.markdown(f"- {esc(req)}")
    if lifecycle.get("top_risks"):
        with st.expander("🔴 Top Risks for This Matter", expanded=False):
            for risk in lifecycle["top_risks"]:
                st.markdown(f"- {esc(risk)}")

    st.markdown(f"""
<div style="background:#f0fdf4;border-left:4px solid #059669;
padding:1rem;border-radius:0.5rem;margin-bottom:1.5rem;">
  <strong>⚡ Immediate Next Step:</strong> {esc(lifecycle.get('immediate_next_step', ''))}
</div>""", unsafe_allow_html=True)

    # ── Stages ──
    st.markdown("### 📋 Case Stages")
    stages = lifecycle.get("stages", [])

    for stage in stages:
        stage_num = str(stage.get("stage_number", ""))
        is_done = progress.get(stage_num, False)

        if is_done:
            card_color = "#f0fdf4"
            border_color = "#059669"
            status_icon = "✅"
        else:
            card_color = "#ffffff"
            border_color = "#e5e7eb"
            status_icon = "⬜"

        # Check if previous stage done (for sequential enforcement)
        prev_done = True
        if stage.get("stage_number", 1) > 1:
            prev_done = progress.get(str(stage.get("stage_number", 1) - 1), False)

        with st.expander(
            f"{status_icon} Stage {stage_num}: {stage.get('stage_name', '')} "
            f"— {stage.get('duration_estimate', '')}",
            expanded=not is_done,
        ):
            st.markdown(f"""
<div style="background:{card_color};border:1px solid {border_color};
border-radius:0.75rem;padding:1.2rem;">
  <p>{esc(stage.get('description', ''))}</p>
  <p><strong>⏱️ Duration:</strong> {esc(stage.get('duration_estimate', ''))} &nbsp;|&nbsp;
  <strong>📅 Trigger:</strong> {esc(stage.get('deadline_trigger', ''))}</p>
  {f'<div style="background:#fef3c7;border-left:3px solid #f59e0b;padding:0.6rem;border-radius:0.3rem;margin-top:0.5rem;"><strong>⚠️ Warning:</strong> {esc(stage.get("warning",""))}</div>' if stage.get("warning") else ""}
</div>""", unsafe_allow_html=True)

            dc1, dc2 = st.columns(2)
            with dc1:
                st.markdown("**📄 Required Documents:**")
                for doc in stage.get("required_documents", []):
                    st.markdown(f"- {esc(doc)}")
            with dc2:
                st.markdown("**✅ Required Actions:**")
                for action in stage.get("required_actions", []):
                    st.markdown(f"- {esc(action)}")

            st.markdown("")
            btn_col1, btn_col2, btn_col3 = st.columns([2, 2, 1])
            with btn_col1:
                if not is_done:
                    if st.button(
                        f"✅ Mark Stage {stage_num} Complete",
                        key=f"lc_done_{case_id}_{stage_num}",
                        type="primary",
                        width='stretch',
                    ):
                        progress[stage_num] = True
                        db.save_lifecycle_progress(case_id, progress)
                        st.success(f"Stage {stage_num} marked complete!")
                        st.rerun()
                else:
                    if st.button(
                        f"↩️ Reopen Stage {stage_num}",
                        key=f"lc_undo_{case_id}_{stage_num}",
                        width='stretch',
                    ):
                        progress[stage_num] = False
                        db.save_lifecycle_progress(case_id, progress)
                        st.rerun()
            with btn_col2:
                # Generate document for this stage
                if st.button(
                    f"📄 Draft Stage Document",
                    key=f"lc_draft_{case_id}_{stage_num}",
                    width='stretch',
                ):
                    draft_prompt = (
                        f"Case: {selected_case.get('title', '')}\n"
                        f"Court: {selected_case.get('court', '')}\n"
                        f"Suit No: {selected_case.get('suit_no', '')}\n"
                        f"Stage: {stage.get('stage_name', '')}\n"
                        f"Required Documents: {', '.join(stage.get('required_documents', []))}\n\n"
                        f"Draft the most important document needed for this stage. "
                        f"Use [PLACEHOLDER] for missing information."
                    )
                    system = build_system_prompt("drafting", "standard")
                    with st.spinner(f"📄 Drafting {stage.get('stage_name','')} document..."):
                        draft_result = generate(draft_prompt, system, "standard", "drafting")
                    st.markdown(f'<div class="response-box">{esc(draft_result)}</div>',
                                unsafe_allow_html=True)
                    save_analysis_to_case(
                        case_id,
                        f"[Lifecycle Stage {stage_num}] {stage.get('stage_name','')}",
                        draft_result, "drafting", "standard",
                    )
                    fname = f"Stage{stage_num}_{stage.get('stage_name','').replace(' ','_')}"
                    dl1, dl2 = st.columns(2)
                    with dl1:
                        st.download_button(
                            "📥 TXT", export_txt(draft_result, stage.get("stage_name", "")),
                            f"{fname}.txt", "text/plain",
                            key=f"lc_dl_txt_{case_id}_{stage_num}",
                            width='stretch',
                        )
                    with dl2:
                        safe_docx_download(
                            draft_result, stage.get("stage_name", ""),
                            fname, f"lc_dl_docx_{case_id}_{stage_num}",
                        )

    # ── Regenerate ──
    st.markdown("---")
    rg1, rg2 = st.columns(2)
    with rg1:
        if st.button(
            "🔄 Regenerate Lifecycle",
            key="lifecycle_regen_btn",
            width='stretch',
        ):
            db.save_lifecycle(case_id, {})
            db.save_lifecycle_progress(case_id, {})
            st.success("Lifecycle cleared. Scroll up and regenerate.")
            st.rerun()
    with rg2:
        # Export full lifecycle as TXT
        lifecycle_text = f"MATTER LIFECYCLE — {selected_case.get('title','')}\n"
        lifecycle_text += f"Court: {lifecycle.get('court_recommendation','')}\n"
        lifecycle_text += f"Duration: {lifecycle.get('estimated_duration','')}\n\n"
        for s in stages:
            lifecycle_text += f"STAGE {s.get('stage_number','')}: {s.get('stage_name','')}\n"
            lifecycle_text += f"  {s.get('description','')}\n"
            lifecycle_text += f"  Documents: {', '.join(s.get('required_documents',[]))}\n"
            lifecycle_text += f"  Actions: {', '.join(s.get('required_actions',[]))}\n\n"
        st.download_button(
            "📥 Export Full Lifecycle (TXT)",
            export_txt(lifecycle_text, f"Matter Lifecycle — {selected_case.get('title','')}"),
            f"Lifecycle_{selected_case.get('title','').replace(' ','_')}.txt",
            "text/plain",
            key="lifecycle_export_btn",
            width='stretch',
        )
# ═══════════════════════════════════════════════════════
# PAGE: WITNESS PREPARATION ENGINE
# ═══════════════════════════════════════════════════════
def _wp_extract_section(text: str, header_fragment: str) -> str:
    """Extract text between two witness prep section headers."""
    lines = text.split("\n")
    capture = False
    collected = []
    for line in lines:
        if header_fragment.upper() in line.upper() and "═" in line:
            capture = True
            continue
        if capture and "═══" in line and collected:
            break
        if capture:
            collected.append(line)
    return "\n".join(collected).strip()


def render_witness_prep():
    st.markdown("""<div class="page-header">
        <h2>🎯 Witness Preparation Engine</h2>
        <p>Input case facts and witness role → Examination-in-chief · Cross-exam risks ·
        Re-examination · Coaching notes · Multi-witness contradiction check</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    # Ensure session log exists
    if "wp_witness_log" not in st.session_state:
        st.session_state["wp_witness_log"] = []

    # ── Main tabs ──
    tab_prep, tab_log, tab_contra = st.tabs([
        "🎯 Prepare a Witness",
        f"👥 Witness Log ({len(st.session_state['wp_witness_log'])})",
        "🔍 Contradiction Check",
    ])

    # ═══════════════════════════════════════════════════
    # TAB 1 — PREPARE A WITNESS
    # ═══════════════════════════════════════════════════
    with tab_prep:
        wp1, wp2 = st.columns([2, 1])
        with wp1:
            wp_facts = st.text_area(
                "Case Facts *",
                height=210,
                key="wp_facts_ta",
                placeholder="""Describe the key facts of the case as they relate to this witness.

Example: The witness, Mrs Amaka Obi, is a neighbour of the claimant. She was present on
3 January 2024 when the defendant's vehicle collided with the claimant's gate at Ikeja.
She heard the crash, came outside within 2 minutes, saw the defendant exit the vehicle,
and heard him say 'I lost control'. She took three photographs on her phone.
Opponent may argue she was too far away to hear clearly and has a prior land dispute
with the defendant.""",
            )
        with wp2:
            wp_role = st.text_input(
                "Witness Role *",
                key="wp_role_inp",
                placeholder="e.g. Eyewitness, Expert (valuation), Claimant",
            )
            wp_name = st.text_input(
                "Witness Name (optional)",
                key="wp_name_inp",
                placeholder="e.g. Mrs Amaka Obi",
            )
            wp_case_type = st.selectbox(
                "Case Type (optional)",
                ["— Select —"] + CASE_TYPE_OPTIONS,
                key="wp_case_type_sel",
            )
            case_type_val = "" if wp_case_type == "— Select —" else wp_case_type
            mode = st.session_state.response_mode
            st.info(f"Mode: {RESPONSE_MODES[mode]['label']}")
            wp_generate_btn = st.button(
                "🎯 Prepare Witness",
                type="primary",
                width='stretch',
                key="wp_generate_btn",
                disabled=not (wp_facts.strip() and wp_role.strip()),
            )

        if wp_generate_btn and wp_facts.strip() and wp_role.strip():
            prompt = WITNESS_PREP_PROMPT.format(
                case_facts=wp_facts.strip(),
                witness_role=wp_role.strip(),
                case_type=case_type_val or "Not specified",
            )
            with st.spinner("🎯 Preparing witness brief…"):
                raw = generate(prompt, WITNESS_PREP_SYSTEM, mode, "analysis")
            label = wp_name.strip() or wp_role.strip()
            st.session_state["wp_result"] = raw
            st.session_state["wp_role_label"] = label
            st.session_state["wp_facts_saved"] = wp_facts.strip()
            # Add to witness log
            st.session_state["wp_witness_log"].append({
                "id": new_id(),
                "name": wp_name.strip() or f"Witness {len(st.session_state['wp_witness_log'])+1}",
                "role": wp_role.strip(),
                "case_type": case_type_val or "Not specified",
                "facts": wp_facts.strip(),
                "result": raw,
                "timestamp": datetime.now().strftime("%d %b %Y %H:%M"),
            })
            st.rerun()

        # ── Display result ──
        result = st.session_state.get("wp_result", "")
        role_label = st.session_state.get("wp_role_label", "Witness")
        facts_saved = st.session_state.get("wp_facts_saved", "")

        if result:
            st.markdown("---")
            sec1 = _wp_extract_section(result, "EXAMINATION-IN-CHIEF")
            sec2 = _wp_extract_section(result, "CROSS-EXAMINATION")
            sec3 = _wp_extract_section(result, "COACHING NOTES")

            if not (sec1 and sec2 and sec3):
                st.markdown(f'<div class="response-box">{esc(result)}</div>', unsafe_allow_html=True)
            else:
                s1_tab, s2_tab, s3_tab, s4_tab = st.tabs([
                    "📋 Examination-in-Chief",
                    "⚔️ Cross-Examination Risks",
                    "🧭 Coaching Notes",
                    "↩️ Re-Examination",
                ])

                with s1_tab:
                    st.markdown(f"""
<div style="background:#f0fdf4;border-left:4px solid #059669;border-radius:0.75rem;
padding:1.5rem;line-height:1.8;">
  <h4 style="margin:0 0 1rem 0;color:#059669;">📋 Examination-in-Chief — {esc(role_label)}</h4>
  <div style="white-space:pre-wrap;font-size:0.95rem;">{esc(sec1)}</div>
</div>""", unsafe_allow_html=True)

                with s2_tab:
                    st.markdown(f"""
<div style="background:#fef2f2;border-left:4px solid #dc2626;border-radius:0.75rem;
padding:1.5rem;line-height:1.8;">
  <h4 style="margin:0 0 1rem 0;color:#dc2626;">⚔️ Cross-Examination Risks</h4>
  <div style="white-space:pre-wrap;font-size:0.95rem;">{esc(sec2)}</div>
</div>""", unsafe_allow_html=True)

                with s3_tab:
                    st.markdown(f"""
<div style="background:#fffbeb;border-left:4px solid #f59e0b;border-radius:0.75rem;
padding:1.5rem;line-height:1.8;">
  <h4 style="margin:0 0 1rem 0;color:#d97706;">🧭 Coaching Notes for the Witness</h4>
  <div style="white-space:pre-wrap;font-size:0.95rem;">{esc(sec3)}</div>
</div>""", unsafe_allow_html=True)

                with s4_tab:
                    st.markdown("""
<div style="background:#eff6ff;border-left:4px solid #3b82f6;border-radius:0.6rem;
padding:0.9rem 1.2rem;margin-bottom:1rem;">
  <strong style="color:#1d4ed8;">↩️ Re-Examination Questions</strong><br>
  <small style="color:#475569;">Generated from the cross-examination attack points above.
  Re-examination is limited to matters arising from cross-examination (Evidence Act 2011, s.215).</small>
</div>""", unsafe_allow_html=True)

                    reexam_result = st.session_state.get("wp_reexam_result", "")
                    if not reexam_result:
                        if st.button(
                            "↩️ Generate Re-Examination Questions",
                            type="primary",
                            key="wp_reexam_btn",
                            width='stretch',
                        ):
                            reexam_p = REEXAM_PROMPT.format(
                                witness_role=role_label,
                                case_facts=facts_saved,
                                cross_exam_risks=sec2,
                            )
                            with st.spinner("↩️ Generating re-examination questions…"):
                                reexam_result = generate(reexam_p, REEXAM_SYSTEM, "standard", "analysis")
                            st.session_state["wp_reexam_result"] = reexam_result
                            st.rerun()
                    else:
                        st.markdown(f"""
<div style="background:#eff6ff;border-left:4px solid #3b82f6;border-radius:0.75rem;
padding:1.5rem;line-height:1.8;white-space:pre-wrap;font-size:0.95rem;">
{esc(reexam_result)}</div>""", unsafe_allow_html=True)
                        re1, re2 = st.columns(2)
                        with re1:
                            st.download_button(
                                "📥 Download Re-Examination (TXT)",
                                export_txt(reexam_result, f"Re-Examination — {role_label}"),
                                f"ReExam_{role_label.replace(' ','_')}_{datetime.now():%Y%m%d}.txt",
                                "text/plain", key="wp_reexam_dl_txt", width='stretch',
                            )
                        with re2:
                            if st.button("🔄 Regenerate", key="wp_reexam_regen", width='stretch'):
                                st.session_state["wp_reexam_result"] = ""
                                st.rerun()

            # ── Save to Case ──
            st.markdown("---")
            cases = st.session_state.cases
            if cases:
                st.markdown("##### 💾 Save to Case File")
                sv1, sv2 = st.columns([3, 1])
                with sv1:
                    save_case_options = {c["id"]: f"{c.get('title','Untitled')} [{c.get('status','')}]"
                                         for c in cases}
                    save_case_id = st.selectbox(
                        "Select Case",
                        list(save_case_options.keys()),
                        format_func=lambda x: save_case_options[x],
                        key="wp_save_case_sel",
                    )
                with sv2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("💾 Save", type="primary", key="wp_save_case_btn", width='stretch'):
                        save_analysis_to_case(
                            save_case_id,
                            f"[Witness Prep] {role_label}",
                            result, "analysis", mode,
                        )
                        st.success(f"✅ Saved to case: {save_case_options[save_case_id]}")

            # ── Export row ──
            st.markdown("##### 📥 Export")
            fname = f"WitnessPrep_{role_label.replace(' ','_')}_{datetime.now():%Y%m%d_%H%M}"
            ex1, ex2, ex3, ex4 = st.columns(4)
            with ex1:
                st.download_button("📥 TXT", export_txt(result, f"Witness Prep — {role_label}"),
                    f"{fname}.txt", "text/plain", key="wp_dl_txt", width='stretch')
            with ex2:
                st.download_button("📥 HTML", export_html(result, f"Witness Prep — {role_label}"),
                    f"{fname}.html", "text/html", key="wp_dl_html", width='stretch')
            with ex3:
                safe_pdf_download(result, f"Witness Prep — {role_label}", fname, "wp_dl_pdf")
            with ex4:
                safe_docx_download(result, f"Witness Prep — {role_label}", fname, "wp_dl_docx", doc_type="witness", meta={"role": role_label})

            if st.button("🗑️ Clear Current Brief", key="wp_clear_btn", width='stretch'):
                for k in ["wp_result", "wp_role_label", "wp_facts_saved", "wp_reexam_result"]:
                    st.session_state[k] = ""
                st.rerun()

            st.markdown("""<div class="disclaimer">
                <strong>⚖️ Disclaimer:</strong> AI-generated witness preparation materials are a
                drafting aid only. Review all questions against actual witness statements. Do not
                share coaching notes or cross-exam analysis with opposing counsel.
            </div>""", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════
    # TAB 2 — WITNESS SESSION LOG
    # ═══════════════════════════════════════════════════
    with tab_log:
        log = st.session_state["wp_witness_log"]
        if not log:
            st.info("No witnesses prepared yet in this session. Use the 'Prepare a Witness' tab to get started.")
        else:
            st.markdown(f"##### 👥 {len(log)} Witness(es) Prepared This Session")
            st.caption("All witness briefs are held in memory for this session. Use the Contradiction Check tab to compare accounts.")

            for i, entry in enumerate(log):
                with st.expander(
                    f"{'👤'} {esc(entry['name'])} — {esc(entry['role'])} "
                    f"· {esc(entry['timestamp'])}",
                    expanded=False,
                ):
                    log_sec1 = _wp_extract_section(entry["result"], "EXAMINATION-IN-CHIEF")
                    log_sec2 = _wp_extract_section(entry["result"], "CROSS-EXAMINATION")
                    log_sec3 = _wp_extract_section(entry["result"], "COACHING NOTES")

                    if log_sec1 and log_sec2 and log_sec3:
                        lt1, lt2, lt3 = st.tabs(["📋 Exam-in-Chief", "⚔️ Cross-Exam Risks", "🧭 Coaching"])
                        with lt1:
                            st.markdown(f'<div style="white-space:pre-wrap;font-size:0.9rem;'
                                        f'background:#f0fdf4;padding:1rem;border-radius:0.5rem;">'
                                        f'{esc(log_sec1)}</div>', unsafe_allow_html=True)
                        with lt2:
                            st.markdown(f'<div style="white-space:pre-wrap;font-size:0.9rem;'
                                        f'background:#fef2f2;padding:1rem;border-radius:0.5rem;">'
                                        f'{esc(log_sec2)}</div>', unsafe_allow_html=True)
                        with lt3:
                            st.markdown(f'<div style="white-space:pre-wrap;font-size:0.9rem;'
                                        f'background:#fffbeb;padding:1rem;border-radius:0.5rem;">'
                                        f'{esc(log_sec3)}</div>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<div class="response-box" style="font-size:0.88rem;">'
                                    f'{esc(entry["result"])}</div>', unsafe_allow_html=True)

                    # Quick export per witness
                    loge1, loge2, loge3 = st.columns(3)
                    lname = entry["name"].replace(" ", "_")
                    with loge1:
                        st.download_button(
                            "📥 TXT", export_txt(entry["result"], f"Witness Prep — {entry['name']}"),
                            f"WitnessPrep_{lname}.txt", "text/plain",
                            key=f"wp_log_dl_{i}", width='stretch',
                        )
                    with loge2:
                        safe_pdf_download(
                            entry["result"], f"Witness Prep — {entry['name']}",
                            f"WitnessPrep_{lname}", f"wp_log_pdf_{i}",
                        )
                    with loge3:
                        if st.button("🗑️ Remove from Log", key=f"wp_log_del_{i}", width='stretch'):
                            st.session_state["wp_witness_log"].pop(i)
                            st.rerun()

            if st.button("🗑️ Clear Entire Witness Log", key="wp_log_clear_all", width='stretch'):
                st.session_state["wp_witness_log"] = []
                st.rerun()

    # ═══════════════════════════════════════════════════
    # TAB 3 — CONTRADICTION CHECK
    # ═══════════════════════════════════════════════════
    with tab_contra:
        log = st.session_state["wp_witness_log"]
        st.markdown("#### 🔍 Multi-Witness Contradiction Detector")
        st.caption(
            "Select two or more witnesses from your session log. "
            "AI will identify contradictions, gaps, and corroborations between their accounts — "
            "and suggest how to reconcile them before trial."
        )

        if len(log) < 2:
            st.warning(
                "⚠️ You need at least 2 prepared witnesses in your session log to run a contradiction check. "
                "Prepare more witnesses first."
            )
        else:
            # Multi-select from log
            witness_options = {entry["id"]: f"{entry['name']} ({entry['role']})" for entry in log}
            selected_ids = st.multiselect(
                "Select Witnesses to Compare (minimum 2)",
                list(witness_options.keys()),
                format_func=lambda x: witness_options[x],
                default=list(witness_options.keys())[:min(2, len(witness_options))],
                key="wp_contra_sel",
            )

            contra_btn = st.button(
                "🔍 Run Contradiction Check",
                type="primary",
                width='stretch',
                key="wp_contra_btn",
                disabled=len(selected_ids) < 2,
            )

            if contra_btn and len(selected_ids) >= 2:
                selected_entries = [e for e in log if e["id"] in selected_ids]
                summaries = ""
                for idx, entry in enumerate(selected_entries, 1):
                    summaries += f"\n{'='*50}\nWITNESS {idx}: {entry['name']} ({entry['role']})\n"
                    summaries += f"Case Type: {entry['case_type']}\n\n"
                    summaries += f"PREPARED BRIEF:\n{entry['result'][:3000]}\n"

                contra_prompt = CONTRADICTION_PROMPT.format(
                    count=len(selected_entries),
                    witness_summaries=summaries,
                )
                with st.spinner(f"🔍 Analysing {len(selected_entries)} witnesses for contradictions…"):
                    contra_result = generate(contra_prompt, CONTRADICTION_SYSTEM, "standard", "analysis")
                st.session_state["wp_contra_result"] = contra_result
                st.rerun()

            contra_result = st.session_state.get("wp_contra_result", "")
            if contra_result:
                st.markdown("---")
                st.markdown(f'<div class="response-box">{esc(contra_result)}</div>',
                            unsafe_allow_html=True)
                st.markdown("---")
                cd1, cd2, cd3 = st.columns(3)
                with cd1:
                    st.download_button(
                        "📥 Export Contradiction Report (TXT)",
                        export_txt(contra_result, "Witness Contradiction Analysis"),
                        f"ContradictionCheck_{datetime.now():%Y%m%d_%H%M}.txt",
                        "text/plain", key="wp_contra_dl_txt", width='stretch',
                    )
                with cd2:
                    safe_pdf_download(
                        contra_result, "Witness Contradiction Analysis",
                        f"ContradictionCheck_{datetime.now():%Y%m%d_%H%M}", "wp_contra_dl_pdf",
                    )
                with cd3:
                    if st.button("🗑️ Clear Result", key="wp_contra_clear", width='stretch'):
                        st.session_state["wp_contra_result"] = ""
                        st.rerun()

                st.markdown("""<div class="disclaimer">
                    <strong>⚖️ Disclaimer:</strong> Contradiction analysis is AI-assisted.
                    Counsel must independently review all witness statements before trial.
                    Intra-party contradictions must be resolved before witnesses take the box.
                </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: LEGAL NEWS FEED
# ═══════════════════════════════════════════════════════
def render_legal_news():
    st.markdown("""<div class="page-header">
        <h2>📰 Nigerian Legal News Feed</h2>
        <p>AI-powered digest of recent legal developments · Bookmarks · Case Relevance Scan ·
        Deep Dive analysis — searchable by subject area</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    # Ensure bookmarks list exists in session
    if "nf_bookmarks" not in st.session_state:
        st.session_state["nf_bookmarks"] = []

    bookmarks = st.session_state["nf_bookmarks"]

    # ── Top-level tabs ──
    tab_feed, tab_bookmarks, tab_scan = st.tabs([
        "📰 Live Feed",
        f"📌 Reading List ({len(bookmarks)})",
        "🎯 Case Relevance Scan",
    ])

    # ═══════════════════════════════════════════════════
    # TAB 1 — LIVE FEED
    # ═══════════════════════════════════════════════════
    with tab_feed:
        # ── Controls ──
        nf1, nf2, nf3 = st.columns([2, 2, 1])
        with nf1:
            nf_subject = st.selectbox(
                "Subject Area",
                NEWS_FEED_SUBJECTS,
                key="nf_subject_sel",
            )
        with nf2:
            nf_search = st.text_input(
                "🔍 Search within feed",
                key="nf_search_inp",
                placeholder="e.g. land registration, employment, tax",
            )
        with nf3:
            st.markdown("<br>", unsafe_allow_html=True)
            nf_generate_btn = st.button(
                "🔄 Fetch Latest",
                type="primary",
                width='stretch',
                key="nf_generate_btn",
            )

        if nf_generate_btn:
            subject_val = nf_subject if nf_subject != "All Areas" else "all major practice areas of Nigerian law"
            prompt = NEWS_FEED_PROMPT.format(
                subject_area=subject_val,
                today=date.today().strftime("%d %B %Y"),
            )
            with st.spinner(f"📰 Fetching legal developments — {nf_subject}…"):
                raw = generate(prompt, NEWS_FEED_SYSTEM, "brief", "research")
            try:
                clean = raw.strip().replace("```json", "").replace("```", "").strip()
                feed_data = json.loads(clean)
                st.session_state["nf_feed_data"] = feed_data
                st.session_state["nf_subject_loaded"] = nf_subject
                # Clear any stale deep-dive results
                st.session_state["nf_deepdive"] = {}
            except Exception:
                st.session_state["nf_feed_data"] = {"_raw": raw, "items": []}
                st.session_state["nf_subject_loaded"] = nf_subject

        feed_data = st.session_state.get("nf_feed_data", None)
        subject_loaded = st.session_state.get("nf_subject_loaded", "")
        if "nf_deepdive" not in st.session_state:
            st.session_state["nf_deepdive"] = {}

        if feed_data is None:
            st.markdown("""
<div style="background:#f8fafc;border:1.5px dashed #cbd5e1;border-radius:0.85rem;
padding:2.5rem;text-align:center;color:#64748b;">
  <h3 style="margin:0 0 0.5rem 0;">📰 Your Legal Feed is Empty</h3>
  <p style="margin:0;">Select a subject area and click <strong>Fetch Latest</strong>
  to load Nigerian legal developments.</p>
</div>""", unsafe_allow_html=True)

        elif "_raw" in feed_data:
            st.warning("⚠️ Could not parse as structured data. Showing raw output:")
            st.markdown(f'<div class="response-box">{esc(feed_data["_raw"])}</div>',
                        unsafe_allow_html=True)

        else:
            items = feed_data.get("items", [])
            gen_date = feed_data.get("generated_date", date.today().strftime("%d %B %Y"))

            # ── Header ──
            hd1, hd2 = st.columns([3, 1])
            with hd1:
                st.markdown(f"""
<div style="padding:0.6rem 1rem;background:#f1f5f9;border-radius:0.5rem;
display:inline-block;font-size:0.9rem;">
  📅 <strong>Ref date:</strong> {esc(gen_date)} &nbsp;|&nbsp;
  📂 <strong>Subject:</strong> {esc(subject_loaded)} &nbsp;|&nbsp;
  📰 <strong>{len(items)} items</strong> &nbsp;|&nbsp;
  📌 <strong>{len(bookmarks)} bookmarked</strong>
</div>""", unsafe_allow_html=True)
            with hd2:
                if st.button("🗑️ Clear Feed", key="nf_clear_btn", width='stretch'):
                    st.session_state["nf_feed_data"] = None
                    st.session_state["nf_subject_loaded"] = ""
                    st.session_state["nf_deepdive"] = {}
                    st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)

            # ── Filter by search ──
            search_val = nf_search.strip().lower()
            display_items = items
            if search_val:
                display_items = [
                    item for item in items
                    if search_val in item.get("title", "").lower()
                    or search_val in item.get("summary", "").lower()
                    or search_val in item.get("key_takeaway", "").lower()
                    or search_val in item.get("practice_impact", "").lower()
                ]

            if not display_items:
                st.info(f"No items match '{nf_search}'. Try a different term or clear the filter.")
            else:
                for item in display_items:
                    item_id = str(item.get("id", 0))
                    title = item.get("title", "Untitled Development")
                    summary = item.get("summary", "")
                    takeaway = item.get("key_takeaway", "")
                    impact = item.get("practice_impact", "")

                    # Check if bookmarked
                    is_bookmarked = any(b.get("id") == item_id for b in bookmarks)
                    bm_icon = "📌" if is_bookmarked else "🔖"

                    with st.expander(f"{'📌' if is_bookmarked else '📰'} {esc(title)}", expanded=False):
                        st.markdown(f"""
<div style="background:#ffffff;border:1px solid #e2e8f0;border-radius:0.75rem;padding:1.2rem;">
  <p style="margin:0 0 0.9rem 0;font-size:0.95rem;line-height:1.7;color:#1e293b;">{esc(summary)}</p>
  <div style="background:#f0fdf4;border-left:3px solid #059669;padding:0.7rem 1rem;
  border-radius:0.4rem;margin-bottom:0.7rem;">
    <strong style="color:#059669;">🔑 Key Takeaway:</strong>
    <span style="font-size:0.93rem;"> {esc(takeaway)}</span>
  </div>
  <div style="background:#eff6ff;border-left:3px solid #3b82f6;padding:0.7rem 1rem;
  border-radius:0.4rem;">
    <strong style="color:#1d4ed8;">⚖️ Practice Impact:</strong>
    <span style="font-size:0.93rem;"> {esc(impact)}</span>
  </div>
</div>""", unsafe_allow_html=True)

                        # ── Action buttons ──
                        act1, act2, act3 = st.columns(3)

                        with act1:
                            bm_label = "📌 Bookmarked" if is_bookmarked else "🔖 Bookmark"
                            if st.button(bm_label, key=f"nf_bm_{item_id}", width='stretch'):
                                if is_bookmarked:
                                    st.session_state["nf_bookmarks"] = [
                                        b for b in bookmarks if b.get("id") != item_id
                                    ]
                                    st.success("Removed from Reading List.")
                                else:
                                    st.session_state["nf_bookmarks"].append({
                                        "id": item_id,
                                        "title": title,
                                        "summary": summary,
                                        "key_takeaway": takeaway,
                                        "practice_impact": impact,
                                        "subject": subject_loaded,
                                        "saved_at": datetime.now().strftime("%d %b %Y %H:%M"),
                                    })
                                    st.success("✅ Added to Reading List.")
                                st.rerun()

                        with act2:
                            dd_key = f"nf_dd_{item_id}"
                            dd_result = st.session_state["nf_deepdive"].get(item_id, "")
                            if not dd_result:
                                if st.button("🔬 Deep Dive Analysis", key=dd_key, width='stretch'):
                                    dd_prompt = NEWS_DEEPDIVE_PROMPT.format(
                                        title=title, summary=summary,
                                        takeaway=takeaway, impact=impact,
                                    )
                                    with st.spinner(f"🔬 Analysing: {title[:50]}…"):
                                        dd_result = generate(dd_prompt, NEWS_DEEPDIVE_SYSTEM, "standard", "analysis")
                                    st.session_state["nf_deepdive"][item_id] = dd_result
                                    st.rerun()
                            else:
                                if st.button("🔬 Hide Deep Dive", key=dd_key, width='stretch'):
                                    st.session_state["nf_deepdive"].pop(item_id, None)
                                    st.rerun()

                        with act3:
                            st.download_button(
                                "📥 Export Item",
                                export_txt(
                                    f"TITLE: {title}\n\nSUMMARY:\n{summary}\n\n"
                                    f"KEY TAKEAWAY:\n{takeaway}\n\nPRACTICE IMPACT:\n{impact}",
                                    title,
                                ),
                                f"LegalNews_{item_id}_{datetime.now():%Y%m%d}.txt",
                                "text/plain",
                                key=f"nf_dl_{item_id}",
                                width='stretch',
                            )

                        # ── Deep Dive result ──
                        if dd_result:
                            st.markdown(f"""
<div style="margin-top:1rem;background:#f8fafc;border:1px solid #cbd5e1;
border-radius:0.75rem;padding:1.4rem;">
  <h5 style="margin:0 0 0.8rem 0;color:#1e293b;">🔬 Full Legal Analysis</h5>
  <div style="white-space:pre-wrap;font-size:0.92rem;line-height:1.75;">{esc(dd_result)}</div>
</div>""", unsafe_allow_html=True)
                            safe_pdf_download(
                                dd_result, f"Deep Dive — {title}",
                                f"DeepDive_{item_id}_{datetime.now():%Y%m%d}",
                                f"nf_dd_pdf_{item_id}",
                            )

            # ── Export full feed ──
            st.markdown("---")
            if items:
                feed_text = f"NIGERIAN LEGAL NEWS FEED\nSubject: {subject_loaded}\nDate: {gen_date}\n\n"
                for item in items:
                    feed_text += f"{'='*60}\n{item.get('title','')}\n\n"
                    feed_text += f"SUMMARY:\n{item.get('summary','')}\n\n"
                    feed_text += f"KEY TAKEAWAY:\n{item.get('key_takeaway','')}\n\n"
                    feed_text += f"PRACTICE IMPACT:\n{item.get('practice_impact','')}\n\n"

                ef1, ef2 = st.columns(2)
                fname = f"LegalNewsFeed_{subject_loaded.replace(' ','_').replace('/','_')}_{datetime.now():%Y%m%d_%H%M}"
                with ef1:
                    st.download_button(
                        "📥 Export Full Feed (TXT)",
                        export_txt(feed_text, f"Nigerian Legal News Feed — {subject_loaded}"),
                        f"{fname}.txt", "text/plain",
                        key="nf_dl_txt", width='stretch',
                    )
                with ef2:
                    st.download_button(
                        "📥 Export Full Feed (HTML)",
                        export_html(feed_text, f"Nigerian Legal News Feed — {subject_loaded}"),
                        f"{fname}.html", "text/html",
                        key="nf_dl_html", width='stretch',
                    )

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> This feed is AI-generated. All case citations are
            [CITATION TO BE VERIFIED]. Verify all developments against official law reports
            and primary sources before relying on them in practice.
        </div>""", unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════
    # TAB 2 — READING LIST / BOOKMARKS
    # ═══════════════════════════════════════════════════
    with tab_bookmarks:
        bookmarks = st.session_state["nf_bookmarks"]
        if not bookmarks:
            st.info("📌 No items bookmarked yet. Open any feed item and click 🔖 Bookmark to save it here.")
        else:
            st.markdown(f"##### 📌 {len(bookmarks)} Saved Item(s)")

            bm_search = st.text_input("🔍 Search reading list", key="bm_search_inp",
                                       placeholder="Search your bookmarks...")
            bm_search_val = bm_search.strip().lower()
            display_bm = bookmarks
            if bm_search_val:
                display_bm = [b for b in bookmarks
                               if bm_search_val in b.get("title", "").lower()
                               or bm_search_val in b.get("summary", "").lower()]

            for i, bm in enumerate(display_bm):
                with st.expander(f"📌 {esc(bm.get('title',''))}"
                                 f" · {esc(bm.get('subject',''))} · {esc(bm.get('saved_at',''))}",
                                 expanded=False):
                    st.markdown(f"""
<div style="background:#fff;border:1px solid #e2e8f0;border-radius:0.75rem;padding:1.1rem;">
  <p style="margin:0 0 0.8rem 0;font-size:0.93rem;line-height:1.7;">{esc(bm.get('summary',''))}</p>
  <div style="background:#f0fdf4;border-left:3px solid #059669;padding:0.6rem 0.9rem;
  border-radius:0.4rem;margin-bottom:0.6rem;">
    <strong style="color:#059669;">🔑</strong> {esc(bm.get('key_takeaway',''))}
  </div>
  <div style="background:#eff6ff;border-left:3px solid #3b82f6;padding:0.6rem 0.9rem;
  border-radius:0.4rem;">
    <strong style="color:#1d4ed8;">⚖️</strong> {esc(bm.get('practice_impact',''))}
  </div>
</div>""", unsafe_allow_html=True)

                    bm_act1, bm_act2 = st.columns(2)
                    with bm_act1:
                        st.download_button(
                            "📥 Export (TXT)",
                            export_txt(
                                f"TITLE: {bm.get('title','')}\n\n"
                                f"SUMMARY:\n{bm.get('summary','')}\n\n"
                                f"KEY TAKEAWAY:\n{bm.get('key_takeaway','')}\n\n"
                                f"PRACTICE IMPACT:\n{bm.get('practice_impact','')}",
                                bm.get("title", ""),
                            ),
                            f"Bookmark_{bm.get('id','x')}_{datetime.now():%Y%m%d}.txt",
                            "text/plain",
                            key=f"bm_dl_{i}",
                            width='stretch',
                        )
                    with bm_act2:
                        if st.button("🗑️ Remove", key=f"bm_del_{i}", width='stretch'):
                            bm_id = bm.get("id")
                            st.session_state["nf_bookmarks"] = [
                                b for b in st.session_state["nf_bookmarks"] if b.get("id") != bm_id
                            ]
                            st.rerun()

            st.markdown("---")
            if st.button("🗑️ Clear All Bookmarks", key="bm_clear_all", width='stretch'):
                st.session_state["nf_bookmarks"] = []
                st.rerun()

    # ═══════════════════════════════════════════════════
    # TAB 3 — CASE RELEVANCE SCAN
    # ═══════════════════════════════════════════════════
    with tab_scan:
        st.markdown("#### 🎯 Case Relevance Scan")
        st.caption(
            "Paste your case facts below. The AI will scan every item in your current feed "
            "and rank them by relevance to your matter — identifying which developments help, "
            "which hurt, and which raise procedural flags."
        )

        feed_data = st.session_state.get("nf_feed_data", None)
        feed_items = feed_data.get("items", []) if (feed_data and "_raw" not in feed_data) else []

        if not feed_items:
            st.warning("⚠️ Load a news feed first (use the 'Live Feed' tab → Fetch Latest). "
                       "The scanner needs items to check against.")
        else:
            st.info(f"📰 {len(feed_items)} item(s) loaded from feed: **{st.session_state.get('nf_subject_loaded', '')}**")

            scan_facts = st.text_area(
                "Your Case Facts *",
                height=200,
                key="nf_scan_facts_ta",
                placeholder="""Describe your current matter. Example:

Client is a tenant in Lagos who was issued a Notice to Quit in January 2024.
The tenancy is a yearly tenancy at ₦800,000 per annum. Landlord claims breach of
tenancy covenants (subletting). Client denies subletting and has receipts of all rent
paid. Matter is before the Lagos State Rent Tribunal.""",
            )

            scan_btn = st.button(
                "🎯 Scan Feed for Relevance",
                type="primary",
                width='stretch',
                key="nf_scan_btn",
                disabled=not scan_facts.strip(),
            )

            if scan_btn and scan_facts.strip():
                news_text = ""
                for item in feed_items:
                    news_text += (
                        f"\n[Item {item.get('id',0)}] TITLE: {item.get('title','')}\n"
                        f"SUMMARY: {item.get('summary','')}\n"
                        f"TAKEAWAY: {item.get('key_takeaway','')}\n"
                        f"PRACTICE IMPACT: {item.get('practice_impact','')}\n"
                    )

                scan_prompt = NEWS_RELEVANCE_PROMPT.format(
                    case_facts=scan_facts.strip(),
                    news_items=news_text,
                )
                with st.spinner(f"🎯 Scanning {len(feed_items)} items against your case facts…"):
                    raw_scan = generate(scan_prompt, NEWS_RELEVANCE_SYSTEM, "brief", "analysis")

                try:
                    clean_scan = raw_scan.strip().replace("```json", "").replace("```", "").strip()
                    scan_data = json.loads(clean_scan)
                    st.session_state["nf_scan_result"] = scan_data
                except Exception:
                    st.session_state["nf_scan_result"] = {"_raw": raw_scan}
                st.rerun()

            scan_result = st.session_state.get("nf_scan_result", None)

            if scan_result:
                st.markdown("---")

                if "_raw" in scan_result:
                    st.markdown(f'<div class="response-box">{esc(scan_result["_raw"])}</div>',
                                unsafe_allow_html=True)
                else:
                    # Summary banner
                    summary_text = scan_result.get("scan_summary", "")
                    if summary_text:
                        st.markdown(f"""
<div style="background:#f0fdf4;border:2px solid #059669;border-radius:0.75rem;
padding:1rem 1.4rem;margin-bottom:1.2rem;">
  <strong style="color:#059669;">🎯 Scan Summary:</strong>
  <span style="font-size:0.95rem;"> {esc(summary_text)}</span>
</div>""", unsafe_allow_html=True)

                    scan_items = scan_result.get("items", [])
                    # Sort by score descending
                    scan_items = sorted(scan_items, key=lambda x: x.get("relevance_score", 0), reverse=True)

                    for si in scan_items:
                        score = si.get("relevance_score", 0)
                        label = si.get("relevance_label", "")
                        fav = si.get("favourable_or_unfavourable", "NEUTRAL")
                        how = si.get("how_it_affects_case", "")
                        si_title = si.get("title", "")

                        if score >= 7:
                            score_color = "#059669"; bg = "#f0fdf4"; border = "#059669"
                        elif score >= 5:
                            score_color = "#d97706"; bg = "#fffbeb"; border = "#f59e0b"
                        elif score >= 1:
                            score_color = "#64748b"; bg = "#f8fafc"; border = "#cbd5e1"
                        else:
                            score_color = "#94a3b8"; bg = "#f8fafc"; border = "#e2e8f0"

                        fav_icons = {
                            "FAVOURABLE": "🟢 Favourable",
                            "UNFAVOURABLE": "🔴 Unfavourable",
                            "NEUTRAL": "⚪ Neutral",
                            "PROCEDURAL": "🔵 Procedural",
                        }
                        fav_label = fav_icons.get(fav, fav)

                        st.markdown(f"""
<div style="background:{bg};border:1px solid {border};border-radius:0.75rem;
padding:1rem 1.2rem;margin-bottom:0.7rem;">
  <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
    <strong style="font-size:0.95rem;">{esc(si_title)}</strong>
    <div style="display:flex;gap:0.5rem;align-items:center;">
      <span style="background:{score_color};color:white;font-weight:700;font-size:0.8rem;
      padding:0.2rem 0.6rem;border-radius:1rem;">{score}/10</span>
      <span style="font-size:0.8rem;color:{score_color};font-weight:600;">{esc(label)}</span>
      <span style="font-size:0.8rem;">{esc(fav_label)}</span>
    </div>
  </div>
  {f'<p style="margin:0;font-size:0.9rem;color:#374151;line-height:1.6;">{esc(how)}</p>' if how else ''}
</div>""", unsafe_allow_html=True)

                    # Export scan report
                    scan_report = f"CASE RELEVANCE SCAN REPORT\nDate: {datetime.now():%d %B %Y at %H:%M}\n\n"
                    scan_report += f"CASE FACTS:\n{st.session_state.get('nf_scan_facts_ta','')}\n\n"
                    scan_report += f"SCAN SUMMARY:\n{summary_text}\n\n"
                    scan_report += "RANKED ITEMS:\n"
                    for si in scan_items:
                        scan_report += (
                            f"\n[Score {si.get('relevance_score',0)}/10 | "
                            f"{si.get('relevance_label','')} | "
                            f"{si.get('favourable_or_unfavourable','')}]\n"
                            f"{si.get('title','')}\n"
                            f"{si.get('how_it_affects_case','')}\n"
                        )

                    sc1, sc2 = st.columns(2)
                    with sc1:
                        st.download_button(
                            "📥 Export Scan Report (TXT)",
                            export_txt(scan_report, "Case Relevance Scan Report"),
                            f"RelevanceScan_{datetime.now():%Y%m%d_%H%M}.txt",
                            "text/plain", key="nf_scan_dl_txt", width='stretch',
                        )
                    with sc2:
                        if st.button("🗑️ Clear Scan", key="nf_scan_clear", width='stretch'):
                            st.session_state["nf_scan_result"] = None
                            st.rerun()

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> Relevance scores are AI-generated assessments.
            Independent legal judgment is required before relying on any matched development.
            Verify all citations against primary sources.
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: NOTES → LEGAL BRIEF CONVERTER
# ═══════════════════════════════════════════════════════
def render_notes_converter():
    st.markdown("""<div class="page-header">
        <h2>📝 Notes → Legal Brief Converter</h2>
        <p>Paste raw client meeting notes — get a structured legal brief,
        retainer letter, letter of demand, or formal advice letter</p>
    </div>""", unsafe_allow_html=True)
    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return
    output_types = {
        "brief":    "📋 Legal Brief (Internal Memo)",
        "retainer": "🤝 Client Retainer Letter",
        "demand":   "📩 Letter of Demand",
        "advice":   "📄 Formal Legal Advice Letter",
    }
    nc1, nc2 = st.columns([2, 1])
    with nc1:
        notes_input = st.text_area(
            "Raw Meeting Notes",
            height=280,
            placeholder="""Paste your raw, unstructured meeting notes here. Example:

Met with Mrs Adaobi today. Her husband died intestate in March.
3 kids. House in Lekki worth maybe 50M. Brother in law is claiming
the house saying it was given to him verbally. She has receipts from
when they bought it in 2011. No will. She wants to know if she can
stop him from selling. Court? How long? Cost?""",
            key="notes_input_ta",
        )
    with nc2:
        output_type = st.selectbox(
            "Convert To",
            list(output_types.keys()),
            format_func=lambda x: output_types[x],
            key="notes_output_type",
        )
        client_name = st.text_input(
            "Client Name",
            placeholder="Mrs Adaobi Okafor",
            key="notes_client_name",
        )
        matter_ref = st.text_input(
            "Matter Reference",
            placeholder="MO/2024/001",
            key="notes_matter_ref",
        )
        mode = st.session_state.response_mode
        st.info(f"Mode: {RESPONSE_MODES[mode]['label']}")
        convert_btn = st.button(
            "✨ Convert Notes",
            type="primary",
            width='stretch',
            disabled=not notes_input.strip(),
            key="notes_convert_btn",
        )
    if convert_btn and notes_input.strip():
        type_prompts = {
            "brief": f"""Convert these raw client meeting notes into a structured
internal legal brief using Nigerian law.
Format strictly as:
CLIENT DETAILS / FACTS AS UNDERSTOOD / ISSUES IDENTIFIED /
APPLICABLE LAW & AUTHORITIES / PRELIMINARY ADVICE /
RECOMMENDED ACTION / RISKS & EXPOSURES
Client: {client_name or '[CLIENT]'} | Ref: {matter_ref or '[REF]'}
Be thorough. Cite Nigerian statutes and cases where relevant.""",
            "retainer": f"""Convert these raw meeting notes into a formal Client
Retainer Letter on Nigerian law firm letterhead format.
Include: scope of engagement, fees structure (use [AMOUNT] placeholders),
our obligations, client obligations, confidentiality clause,
governing law, termination clause, and full signature block.
Client: {client_name or '[CLIENT]'} | Ref: {matter_ref or '[REF]'}
Use standard Nigerian solicitor letter format throughout.""",
            "demand": f"""Convert these raw meeting notes into a formal Letter of
Demand in standard Nigerian solicitor format.
Include: full heading with OUR REF and DATE, RE: line, facts paragraph,
legal position with applicable law, specific demand with exact amount
if mentioned, deadline (7/14/21 days as appropriate), and clear
consequences of non-compliance.
Client: {client_name or '[CLIENT]'} | Ref: {matter_ref or '[REF]'}""",
            "advice": f"""Convert these raw meeting notes into a formal Legal Advice
Letter addressed to the client.
Format: Introduction / Facts as Understood / Legal Position /
Our Advice / Recommended Next Steps / Costs Estimate / Disclaimer
Write in plain English the client can understand.
Explain all legal terms used. No unnecessary Latin.
Client: {client_name or '[CLIENT]'} | Ref: {matter_ref or '[REF]'}""",
        }
        full_prompt = (
            type_prompts[output_type]
            + f"\n\nRAW MEETING NOTES:\n{notes_input.strip()}"
        )
        system = build_system_prompt("drafting", mode)
        with st.spinner(f"✨ Converting notes to {output_types[output_type]}..."):
            result = generate(full_prompt, system, mode, "drafting")
        st.markdown("---")
        st.markdown(f"### {output_types[output_type]}")
        fname = f"LexiAssist_{output_type}_{(client_name or 'client').replace(' ','_')}_{datetime.now():%Y%m%d_%H%M}"
        ex1, ex2, ex3, ex4 = st.columns(4)
        with ex1:
            st.download_button(
                "📥 TXT",
                export_txt(result, output_types[output_type]),
                f"{fname}.txt", "text/plain",
                key="notes_dl_txt", width='stretch',
            )
        with ex2:
            st.download_button(
                "📥 HTML",
                export_html(result, output_types[output_type]),
                f"{fname}.html", "text/html",
                key="notes_dl_html", width='stretch',
            )
        with ex3:
            safe_pdf_download(result, output_types[output_type], fname, "notes_dl_pdf")
        with ex4:
            safe_docx_download(result, output_types[output_type], fname, "notes_dl_docx")
        st.markdown(
            f'<div class="response-box">{esc(result)}</div>',
            unsafe_allow_html=True,
        )
        add_to_history(
            f"[Notes→{output_type.title()}] {notes_input[:80]}",
            result, "drafting", mode,
        )
        cases = st.session_state.cases
        if cases:
            st.markdown("### 💾 Save to Case")
            sc1, sc2 = st.columns([3, 1])
            with sc1:
                case_names = [
                    f"{c.get('title','Untitled')} ({c.get('suit_no','—')})"
                    for c in cases
                ]
                sel = st.selectbox(
                    "Select case:", case_names,
                    key="notes_save_case_sel",
                    label_visibility="collapsed",
                )
            with sc2:
                if st.button("💾 Save", key="notes_save_case_btn",
                             type="primary", width='stretch'):
                    idx = case_names.index(sel)
                    save_analysis_to_case(
                        cases[idx]["id"],
                        f"[Notes→{output_type}] {notes_input[:100]}",
                        result, "drafting", mode,
                    )
                    st.success(f"✅ Saved to: {cases[idx].get('title','')}")
        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> Review all AI-generated documents
            before sending to clients or filing. Verify all legal positions
            and citations independently.
        </div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════
# PAGE: PROFILE
# ═══════════════════════════════════════════════════════
def render_profile():
    st.markdown("""<div class="page-header">
        <h2>👤 User Profile</h2>
        <p>Firm branding, contact details, and security settings</p>
    </div>""", unsafe_allow_html=True)

    profile = st.session_state.profile

    tab_info, tab_security, tab_notif, tab_data = st.tabs(["🏢 Firm Details", "🔐 Security", "🔔 Notifications", "💾 Data Management"])

    # ── Firm Details ──
    with tab_info:
        st.markdown("#### 🏢 Firm / Lawyer Profile")
        st.caption("This information appears on exported documents (PDF, DOCX, HTML, TXT).")

        with st.form("profile_form"):
            p1, p2 = st.columns(2)
            with p1:
                firm_name = st.text_input("Firm Name", value=profile.get("firm_name", ""), key="prof_firm_inp",
                                          placeholder="e.g. Adekunle & Associates")
                lawyer_name = st.text_input("Lawyer Name", value=profile.get("lawyer_name", ""), key="prof_lawyer_inp")
                email = st.text_input("Email", value=profile.get("email", ""), key="prof_email_inp")
            with p2:
                phone = st.text_input("Phone", value=profile.get("phone", ""), key="prof_phone_inp")
                address = st.text_area("Address", value=profile.get("address", ""), height=100, key="prof_addr_inp")

            if st.form_submit_button("💾 Save Profile", type="primary"):
                st.session_state.profile["firm_name"] = firm_name.strip()
                st.session_state.profile["lawyer_name"] = lawyer_name.strip()
                st.session_state.profile["email"] = email.strip()
                st.session_state.profile["phone"] = phone.strip()
                st.session_state.profile["address"] = address.strip()
                persist_profile()
                st.success("✅ Profile saved! Firm name will appear on all exports.")
                st.rerun()

        # Preview
        if profile.get("firm_name"):
            st.markdown("---")
            st.markdown("#### 📄 Export Header Preview")
            st.markdown(f"""<div class="custom-card">
                <h4>{esc(profile.get('firm_name', ''))}</h4>
                {esc(profile.get('lawyer_name', ''))}<br>
                📧 {esc(profile.get('email', ''))} · 📞 {esc(profile.get('phone', ''))}<br>
                📍 {esc(profile.get('address', ''))}
            </div>""", unsafe_allow_html=True)

    # ── Notifications ──
    with tab_notif:
        st.markdown("#### 🔔 Hearing Reminder Emails")
        st.caption("Receive automatic email alerts 7 days and 1 day before each hearing.")
        st.info("💡 Requires a Gmail account with an App Password. Get one at: Google Account → Security → 2-Step Verification → App Passwords")
        with st.form("notif_form"):
            notif_email = st.text_input(
                "Your Email Address (recipient)",
                value=st.session_state.profile.get("notif_email", ""),
                placeholder="yourname@gmail.com",
                key="notif_email_inp",
            )
            notif_smtp_user = st.text_input(
                "Gmail Address (sender)",
                value=st.session_state.profile.get("notif_smtp_user", ""),
                placeholder="sender@gmail.com",
                key="notif_smtp_inp",
            )
            notif_smtp_pass = st.text_input(
                "Gmail App Password",
                type="password",
                key="notif_smtp_pass_inp",
                help="16-character app password from Google Account → Security → App Passwords",
            )
            if st.form_submit_button("💾 Save Notification Settings", type="primary"):
                st.session_state.profile["notif_email"] = notif_email.strip()
                st.session_state.profile["notif_smtp_user"] = notif_smtp_user.strip()
                st.session_state.profile["notif_smtp_pass"] = notif_smtp_pass.strip()
                persist_profile()
                st.success("✅ Notification settings saved!")
        st.markdown("---")
        st.markdown("##### 📬 Send Reminders Now")
        hearings = get_hearings()
        upcoming = [h for h in hearings if 0 <= days_until(h["date"]) <= 7]
        has_email_config = (
            st.session_state.profile.get("notif_email") and
            st.session_state.profile.get("notif_smtp_user") and
            st.session_state.profile.get("notif_smtp_pass")
        )
        if upcoming and has_email_config:
            st.markdown(f"**{len(upcoming)} hearing(s)** within the next 7 days:")
            for h in upcoming:
                d = days_until(h["date"])
                badge = "badge-err" if d <= 1 else ("badge-warn" if d <= 3 else "badge-ok")
                st.markdown(f"""<div class="history-item">
                    <strong>{esc(h['title'])}</strong> ·
                    {esc(h['court'])} ·
                    📅 {esc(fmt_date(h['date']))}
                    <span class="badge {badge}">{esc(relative_date(h['date']))}</span>
                </div>""", unsafe_allow_html=True)
            if st.button(
                "📬 Send Reminder Emails for All Upcoming Hearings",
                key="send_reminders_btn",
                type="primary",
                width='stretch',
            ):
                sent, failed = 0, 0
                firm = get_firm_name()
                for h in upcoming:
                    try:
                        msg = MIMEMultipart("alternative")
                        msg["Subject"] = f"⚖️ Hearing Reminder: {h['title']} — {fmt_date(h['date'])}"
                        msg["From"] = st.session_state.profile["notif_smtp_user"]
                        msg["To"] = st.session_state.profile["notif_email"]
                        body = f"""
<html>
<body style="font-family:Georgia,serif;max-width:600px;margin:auto;padding:20px;color:#1e293b;">
  <h2 style="color:#059669;border-bottom:2px solid #059669;padding-bottom:10px;">
    ⚖️ LexiAssist Hearing Reminder
  </h2>
  <div style="background:#f0fdf4;border-left:4px solid #059669;
  padding:15px;border-radius:8px;margin:20px 0;">
    <h3 style="margin:0 0 10px 0;">{esc(h['title'])}</h3>
    <p style="margin:5px 0;"><strong>Suit Number:</strong> {esc(h['suit'])}</p>
    <p style="margin:5px 0;"><strong>Court:</strong> {esc(h['court'])}</p>
    <p style="margin:5px 0;"><strong>Hearing Date:</strong> {esc(fmt_date(h['date']))}</p>
    <p style="margin:5px 0;"><strong>Days Remaining:</strong>
      <span style="color:#dc2626;font-weight:bold;">{days_until(h['date'])} day(s)</span>
    </p>
  </div>
  <p style="background:#fef3c7;padding:10px;border-radius:6px;">
    ⚠️ Please ensure all court processes, briefs, and appearances
    are prepared well in advance.
  </p>
  <p style="color:#6b7280;font-size:12px;margin-top:30px;
  border-top:1px solid #e5e7eb;padding-top:10px;">
    Sent by <strong>{esc(firm)}</strong> via LexiAssist v8.0 ·
    {datetime.now().strftime('%d %B %Y at %H:%M')}
  </p>
</body>
</html>"""
                        msg.attach(MIMEText(body, "html"))
                        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
                            server.login(
                                st.session_state.profile["notif_smtp_user"],
                                st.session_state.profile["notif_smtp_pass"],
                            )
                            server.sendmail(
                                st.session_state.profile["notif_smtp_user"],
                                st.session_state.profile["notif_email"],
                                msg.as_string(),
                            )
                        sent += 1
                    except Exception as e:
                        failed += 1
                        logger.warning(f"Email send failed: {e}")
                if sent:
                    st.success(f"✅ {sent} reminder email(s) sent to {st.session_state.profile['notif_email']}")
                if failed:
                    st.error(f"❌ {failed} email(s) failed. Check your Gmail App Password and make sure 2-Step Verification is enabled.")
        elif not has_email_config:
            st.info("⚙️ Configure your email settings in the form above to enable reminders.")
        else:
            st.info("✅ No hearings within the next 7 days. You are clear.")
    # ── Security / Account ──
    with tab_security:
        st.markdown("#### 🔐 Account Security")
        st.caption(f"Logged in as: **@{esc(st.session_state.get('current_username',''))}** · "
                   f"Role: **{esc(st.session_state.get('current_user_role','').title())}**")

        st.markdown("##### 🔑 Change Password")
        with st.form("change_pw_form"):
            current_pw = st.text_input("Current Password", type="password", key="cur_pw_inp")
            new_pw = st.text_input("New Password", type="password", key="new_pw_inp")
            confirm_pw = st.text_input("Confirm New Password", type="password", key="confirm_pw_inp")
            if st.form_submit_button("🔐 Update Password", type="primary"):
                if hash_password(current_pw) != profile.get("password_hash", ""):
                    st.error("❌ Current password is incorrect.")
                elif not new_pw:
                    st.error("❌ New password cannot be empty.")
                elif len(new_pw) < 6:
                    st.error("❌ Password must be at least 6 characters.")
                elif new_pw != confirm_pw:
                    st.error("❌ New passwords do not match.")
                else:
                    st.session_state.profile["password_hash"] = hash_password(new_pw)
                    persist_profile()
                    st.success("✅ Password updated successfully!")
                    st.rerun()

        st.markdown("---")
        st.markdown("##### 📋 Account Information")
        uid = st.session_state.get("current_user_id", "")
        if uid:
            user_rec = get_db().get_user_by_id(uid)
            if user_rec:
                ai1, ai2 = st.columns(2)
                with ai1:
                    st.metric("Username", f"@{user_rec.get('username','')}")
                    st.metric("Role", user_rec.get("role", "").title())
                with ai2:
                    st.metric("Joined", fmt_date(user_rec.get("created_at", "")))
                    st.metric("Last Login", fmt_date(user_rec.get("last_login", "")))

        st.markdown("---")
        st.markdown("##### 📱 Active Sessions")
        st.caption("These are all devices and browsers where you are currently signed in.")
        current_token = st.session_state.get("_session_token", "")
        if uid:
            sessions = get_db().get_user_sessions(uid)
            if not sessions:
                st.info("No active sessions found.")
            else:
                for i, sess in enumerate(sessions):
                    is_current = (sess["token"] == current_token)
                    badge = '<span style="background:#059669;color:white;font-size:0.72rem;' \
                            'padding:0.15rem 0.5rem;border-radius:1rem;font-weight:600;">This device</span>' \
                            if is_current else ""
                    st.markdown(f"""
<div style="background:{'#f0fdf4' if is_current else '#f8fafc'};
border:1px solid {'#059669' if is_current else '#e2e8f0'};
border-radius:0.6rem;padding:0.8rem 1rem;margin-bottom:0.5rem;
display:flex;justify-content:space-between;align-items:center;">
  <div>
    🖥️ <strong>Session {i+1}</strong> {badge}<br>
    <small style="color:#64748b;">
      Created: {esc(fmt_date(sess.get('created_at','')))} ·
      Last used: {esc(fmt_date(sess.get('last_used','')))} ·
      Expires: {esc(fmt_date(sess.get('expires_at','')))}
    </small>
  </div>
</div>""", unsafe_allow_html=True)
                    if not is_current:
                        if st.button(f"🚫 Revoke Session {i+1}", key=f"revoke_sess_{i}",
                                     width='stretch'):
                            get_db().revoke_session_token(sess["token"])
                            st.success("✅ Session revoked.")
                            st.rerun()

            st.markdown("")
            if st.button("🚫 Sign Out All Other Devices", key="revoke_all_others",
                         width='stretch'):
                db2 = get_db()
                all_sess = db2.get_user_sessions(uid)
                for sess in all_sess:
                    if sess["token"] != current_token:
                        db2.revoke_session_token(sess["token"])
                st.success("✅ All other sessions revoked.")
                st.rerun()

        st.markdown("---")
        st.markdown("##### 🚪 Sign Out")
        st.caption("Signs you out of this device. Your data is saved.")
        if st.button("🚪 Sign Out Now", key="profile_logout_btn",
                     width='stretch', type="primary"):
            do_logout()

    # ── Data Management ──
    with tab_data:
        st.markdown("#### 💾 Full Backup & Restore")

        # Backup
        st.markdown("##### 📥 Export Full Backup")
        st.caption("Downloads all cases, clients, billing, chat history, templates, references, profile, and cost logs as a single JSON file.")
        if st.button("📦 Generate Full Backup", key="profile_backup_btn", width='stretch', type="primary"):
            export_data = {
                "export_date": datetime.now().isoformat(),
                "version": "8.0",
                "cases": st.session_state.cases,
                "clients": st.session_state.clients,
                "time_entries": st.session_state.time_entries,
                "invoices": st.session_state.invoices,
                "chat_history": st.session_state.chat_history,
                "custom_templates": st.session_state.custom_templates,
                "custom_limitation_periods": st.session_state.custom_limitation_periods,
                "custom_maxims": st.session_state.custom_maxims,
                "profile": {k: v for k, v in st.session_state.profile.items() if k != "password_hash"},
                "cost_logs": get_db().get_cost_logs(500),
            }
            st.download_button(
                "⬇️ Download Full Backup",
                json.dumps(export_data, indent=2, default=str),
                f"lexiassist_full_backup_{datetime.now():%Y%m%d_%H%M}.json",
                "application/json", key="profile_dl_backup",
                width='stretch',
            )

        st.markdown("---")

        # Restore
        st.markdown("##### 📤 Restore from Backup")
        st.caption("Upload a previously exported JSON backup to restore all data.")
        restore_file = st.file_uploader("Upload backup JSON", type=["json"], key="profile_restore_upload")
        if restore_file:
            try:
                raw = json.loads(restore_file.getvalue().decode("utf-8", errors="ignore"))
                if isinstance(raw, dict):
                    st.markdown(f"""<div class="custom-card">
                        <h4>📦 Backup Details</h4>
                        Version: {esc(str(raw.get('version', '?')))} ·
                        Date: {esc(fmt_date(raw.get('export_date', '')))} ·
                        Cases: {len(raw.get('cases', []))} ·
                        Clients: {len(raw.get('clients', []))} ·
                        History: {len(raw.get('chat_history', []))}
                    </div>""", unsafe_allow_html=True)

                    if st.button("⚠️ Restore This Backup (Overwrites Current Data)", type="primary",
                                 key="confirm_restore_btn", width='stretch'):
                        for k in ["cases", "clients", "time_entries", "invoices", "chat_history",
                                   "custom_templates", "custom_limitation_periods", "custom_maxims"]:
                            if k in raw:
                                st.session_state[k] = raw[k]
                                persist(k)
                        if "profile" in raw and isinstance(raw["profile"], dict):
                            for pk, pv in raw["profile"].items():
                                if pk != "password_hash":
                                    st.session_state.profile[pk] = pv
                            persist_profile()
                        st.success("✅ Backup restored successfully!")
                        st.rerun()
                else:
                    st.error("❌ Invalid backup file format.")
            except Exception as e:
                st.error(f"❌ Error reading backup: {e}")

        st.markdown("---")

        # Data stats
        st.markdown("##### 📊 Current Data Summary")
        ds1, ds2, ds3, ds4 = st.columns(4)
        with ds1:
            st.metric("Cases", len(st.session_state.cases))
            st.metric("Clients", len(st.session_state.clients))
        with ds2:
            st.metric("Time Entries", len(st.session_state.time_entries))
            st.metric("Invoices", len(st.session_state.invoices))
        with ds3:
            st.metric("AI Sessions", len(st.session_state.chat_history))
            st.metric("Custom Templates", len(st.session_state.custom_templates))
        with ds4:
            cost_s = get_db().get_cost_summary()
            st.metric("API Calls Logged", cost_s["total_calls"])
            st.metric("Custom Maxims", len(st.session_state.custom_maxims))

        st.markdown("---")

        # Danger zone
        st.markdown("##### ⚠️ Danger Zone")
        st.caption("These actions cannot be undone. Export a backup first!")
        dz1, dz2 = st.columns(2)
        with dz1:
            if st.button("🗑️ Clear All Chat History", key="clear_all_history", width='stretch'):
                st.session_state.chat_history = []
                persist("chat_history")
                st.success("✅ Chat history cleared.")
                st.rerun()
        with dz2:
            if st.button("🗑️ Reset All Data", key="reset_all_data", type="secondary", width='stretch'):
                for k in ["cases", "clients", "time_entries", "invoices", "chat_history",
                           "custom_templates", "custom_limitation_periods", "custom_maxims"]:
                    st.session_state[k] = []
                    persist(k)
                st.session_state.last_response = ""
                st.session_state.original_query = ""
                st.session_state.research_results = ""
                st.success("✅ All data reset. Profile and password preserved.")
                st.rerun()





# ═══════════════════════════════════════════════════════
# PAGE: LEGAL FEE & STAMP DUTY CALCULATOR
# ═══════════════════════════════════════════════════════
def render_fee_calculator():
    st.markdown("""<div class="page-header">
        <h2>⚖️ Legal Fee & Stamp Duty Calculator</h2>
        <p>Nigerian scale fees · Stamp duty on instruments · Court filing fees · Professional fee note</p>
    </div>""", unsafe_allow_html=True)

    tab_land, tab_stamp, tab_court, tab_feenote = st.tabs([
        "🏠 Solicitor's Scale Fees",
        "📄 Stamp Duty",
        "🏛️ Court Filing Fees",
        "🧾 Generate Fee Note",
    ])

    # ═══════════════════════════════════════
    # TAB 1 — SOLICITOR'S SCALE FEES (LAND)
    # ═══════════════════════════════════════
    with tab_land:
        st.markdown("#### 🏠 Land Matters Remuneration Scale")
        st.caption(
            "Based on the Legal Practitioners (Remuneration for Legal Documentation "
            "and Other Land Matters) Order. Applies to: Deeds of Assignment, Conveyances, "
            "Governor's Consent, Mortgages, and related documentation."
        )
        lf1, lf2 = st.columns([2, 1])
        with lf1:
            land_value = st.number_input(
                "Transaction / Property Value (₦)",
                min_value=0.0, value=5_000_000.0, step=100_000.0,
                format="%.2f", key="land_val_inp",
            )
            st.caption(f"Entered: **{fmt_currency(land_value)}**")
        with lf2:
            include_vat = st.checkbox("Add 7.5% VAT on fees", value=True, key="land_vat_chk")
            show_breakdown = st.checkbox("Show band-by-band breakdown", value=True, key="land_bband")

        if st.button("🔢 Calculate Fees", type="primary", key="land_calc_btn", width='stretch'):
            st.session_state["lf_value"] = land_value
            st.session_state["lf_vat"] = include_vat

        lf_value = st.session_state.get("lf_value", 0.0)
        if lf_value > 0:
            base_fee, breakdown = compute_land_fee(lf_value)
            vat = base_fee * 0.075 if include_vat else 0.0
            total = base_fee + vat

            st.markdown("---")
            # Result cards
            r1, r2, r3 = st.columns(3)
            with r1:
                st.markdown(f"""<div class="stat-card">
                    <div class="stat-value">{fmt_currency(base_fee)}</div>
                    <div class="stat-label">Base Solicitor's Fee</div>
                </div>""", unsafe_allow_html=True)
            with r2:
                st.markdown(f"""<div class="stat-card">
                    <div class="stat-value">{fmt_currency(vat)}</div>
                    <div class="stat-label">VAT (7.5%)</div>
                </div>""", unsafe_allow_html=True)
            with r3:
                st.markdown(f"""<div class="stat-card">
                    <div class="stat-value">{fmt_currency(total)}</div>
                    <div class="stat-label">Total Chargeable</div>
                </div>""", unsafe_allow_html=True)

            effective_rate = (base_fee / lf_value * 100) if lf_value > 0 else 0
            st.info(f"💡 Effective rate: **{effective_rate:.3f}%** on {fmt_currency(lf_value)}")

            if show_breakdown:
                st.markdown("##### 📊 Band-by-Band Breakdown")
                import pandas as pd
                df = pd.DataFrame([
                    {
                        "Band": row["band"],
                        "Taxable Amount": fmt_currency(row["taxable"]),
                        "Rate": row["rate"],
                        "Fee": fmt_currency(row["fee"]),
                    }
                    for row in breakdown
                ])
                if include_vat:
                    df.loc[len(df)] = {"Band": "VAT (7.5%)", "Taxable Amount": "", "Rate": "7.5%", "Fee": fmt_currency(vat)}
                df.loc[len(df)] = {"Band": "TOTAL", "Taxable Amount": "", "Rate": "", "Fee": fmt_currency(total)}
                st.dataframe(df, width='stretch', hide_index=True)

            # Store for fee note tab
            st.session_state["fn_land_fee"] = base_fee
            st.session_state["fn_land_vat"] = vat
            st.session_state["fn_land_total"] = total
            st.session_state["fn_land_value"] = lf_value

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Note:</strong> Scale fees under the Land Matters Remuneration Order represent
            the minimum chargeable by a legal practitioner. Fees may be higher by agreement.
            Minimum fee: ₦10,000. Always issue a formal Fee Agreement Letter.
        </div>""", unsafe_allow_html=True)

    # ═══════════════════════════════════════
    # TAB 2 — STAMP DUTY
    # ═══════════════════════════════════════
    with tab_stamp:
        st.markdown("#### 📄 Stamp Duty Calculator")
        st.caption(
            "Stamp Duties Act Cap S8 LFN 2004, as amended by the Finance Acts 2019–2023. "
            "Stamp duty is payable before or within 30 days of execution of the instrument."
        )

        sd1, sd2 = st.columns([2, 1])
        with sd1:
            instrument_key = st.selectbox(
                "Instrument Type",
                list(STAMP_DUTY_INSTRUMENTS.keys()),
                format_func=lambda k: STAMP_DUTY_INSTRUMENTS[k]["label"],
                key="sd_instrument_sel",
            )
        with sd2:
            st.markdown("<br>", unsafe_allow_html=True)
            inst = STAMP_DUTY_INSTRUMENTS[instrument_key]
            st.info(f"💡 **Rate:** {inst['note']}")

        basis = inst["basis"]
        sd_value = 0.0
        sd_years = 1.0
        sd_annual = 0.0

        if basis == "flat":
            st.metric("Stamp Duty Payable", fmt_currency(inst.get("flat", 0)))
            st.session_state["fn_stamp_duty"] = float(inst.get("flat", 0))
        else:
            v1, v2 = st.columns(2)
            with v1:
                if basis in ("property_value", "consideration", "contract_value", "secured_amount"):
                    sd_value = st.number_input("Transaction / Property Value (₦)",
                        min_value=0.0, value=10_000_000.0, step=500_000.0,
                        format="%.2f", key="sd_value_inp")
                elif basis == "loan_amount":
                    sd_value = st.number_input("Loan / Secured Amount (₦)",
                        min_value=0.0, value=5_000_000.0, step=250_000.0,
                        format="%.2f", key="sd_loan_inp")
                elif basis == "guaranteed_sum":
                    sd_value = st.number_input("Guaranteed / Indemnified Sum (₦)",
                        min_value=0.0, value=5_000_000.0, step=250_000.0,
                        format="%.2f", key="sd_guar_inp")
                elif "annual_rent" in basis:
                    sd_annual = st.number_input("Annual Rent (₦)",
                        min_value=0.0, value=1_200_000.0, step=100_000.0,
                        format="%.2f", key="sd_rent_inp")
            with v2:
                if basis == "annual_rent_x_years":
                    sd_years = st.number_input("Number of Years",
                        min_value=0.5, max_value=6.9, value=2.0, step=0.5, key="sd_years_inp")
                    st.caption(f"Annual rent: {fmt_currency(sd_annual)} × {sd_years} years")

            duty = compute_stamp_duty(
                instrument_key,
                value=sd_value if sd_value > 0 else sd_annual,
                years=sd_years,
                annual_rent=sd_annual,
            )

            if st.button("🔢 Calculate Stamp Duty", type="primary",
                         key="sd_calc_btn", width='stretch'):
                st.session_state["sd_result"] = duty

            sd_result = st.session_state.get("sd_result", None)
            if sd_result is not None:
                st.markdown("---")
                sc1, sc2, sc3 = st.columns(3)
                with sc1:
                    st.metric("Stamp Duty", fmt_currency(sd_result))
                with sc2:
                    effective = (sd_result / (sd_value or sd_annual * sd_years or 1)) * 100
                    st.metric("Effective Rate", f"{effective:.3f}%")
                with sc3:
                    st.metric("Penalty (if late > 30 days)", fmt_currency(sd_result * 0.1 + 50))
                st.markdown(f"""
<div style="background:#fffbeb;border-left:3px solid #f59e0b;padding:0.8rem 1rem;
border-radius:0.4rem;margin-top:0.5rem;font-size:0.9rem;">
  ⚠️ <strong>Reminder:</strong> Stamp duty must be paid within 30 days of execution.
  Late stamping attracts a 10% penalty plus ₦50 administrative charge.
  Unstamped instruments are inadmissible in evidence (Stamp Duties Act, s.22).
</div>""", unsafe_allow_html=True)
                st.session_state["fn_stamp_duty"] = sd_result

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Note:</strong> Rates reflect the Stamp Duties Act and Finance Act amendments.
            Stamp duty on electronic transactions and receipts (₦50 on transactions above ₦10,000)
            may apply separately. Confirm with FIRS or relevant State tax authority.
        </div>""", unsafe_allow_html=True)

    # ═══════════════════════════════════════
    # TAB 3 — COURT FILING FEES
    # ═══════════════════════════════════════
    with tab_court:
        st.markdown("#### 🏛️ Court Filing Fee Estimator")
        st.caption(
            "Indicative fees based on the relevant Rules of Court. "
            "Actual fees vary by state, court division, and current court orders. "
            "Always verify with the specific court registry before filing."
        )

        cf1, cf2 = st.columns(2)
        with cf1:
            court_key = st.selectbox(
                "Select Court",
                list(COURT_FILING_FEES.keys()),
                format_func=lambda k: COURT_FILING_FEES[k]["label"],
                key="cf_court_sel",
            )
        with cf2:
            claim_val = st.number_input(
                "Claim Value (₦)",
                min_value=0.0, value=10_000_000.0, step=500_000.0,
                format="%.2f", key="cf_claim_inp",
            )
            st.caption(f"Claim: **{fmt_currency(claim_val)}**")

        if st.button("🔢 Get Filing Fees", type="primary", key="cf_calc_btn", width='stretch'):
            st.session_state["cf_result"] = (court_key, claim_val)

        cf_result = st.session_state.get("cf_result", None)
        if cf_result:
            ck, cv = cf_result
            court = COURT_FILING_FEES[ck]
            filing_fee, court_note = get_court_filing_fee(ck, cv)
            appeal_fee = court.get("appeal_fee", 0)

            st.markdown("---")
            ff1, ff2, ff3 = st.columns(3)
            with ff1:
                st.metric("Originating Process Fee", fmt_currency(filing_fee))
            with ff2:
                st.metric("Estimated Appeal Fee", fmt_currency(appeal_fee))
            with ff3:
                st.metric("Filing + Service (est.)",
                          fmt_currency(filing_fee + filing_fee * 0.3))

            st.info(f"ℹ️ **{court['label']}:** {court_note}")

            # All bands table
            st.markdown("##### 📊 Full Fee Schedule — " + court["label"])
            band_rows = []
            for band in court["bands"]:
                cap = band["claim_max"]
                band_rows.append({
                    "Claim Range": band["label"],
                    "Filing Fee": fmt_currency(band["fee"]),
                })
            import pandas as pd
            st.dataframe(pd.DataFrame(band_rows), width='stretch', hide_index=True)

            # Other cost estimates
            st.markdown("##### 💰 Estimated Total Costs to File")
            items = [
                ("Court filing fee", filing_fee),
                ("Sheriff / Process server fee (est.)", filing_fee * 0.2),
                ("Certified true copies (est.)", 2_000),
                ("Solicitor's filing charges (est.)", 5_000),
            ]
            total_costs = sum(v for _, v in items)
            for label, val in items:
                st.markdown(f"- {label}: **{fmt_currency(val)}**")
            st.markdown(f"**Estimated total to file: {fmt_currency(total_costs)}**")
            st.session_state["fn_court_fee"] = filing_fee

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> Filing fees are indicative only. Court registries
            periodically revise fees. Always obtain an official fee schedule from the registry
            before filing or advising clients on litigation costs.
        </div>""", unsafe_allow_html=True)

    # ═══════════════════════════════════════
    # TAB 4 — GENERATE PROFESSIONAL FEE NOTE
    # ═══════════════════════════════════════
    with tab_feenote:
        st.markdown("#### 🧾 Professional Fee Note Generator")
        st.caption(
            "Generate a formal, professionally formatted Fee Note / Bill of Costs "
            "ready to issue to your client. Uses values computed in the other tabs."
        )
        fn1, fn2 = st.columns(2)
        with fn1:
            fn_client = st.text_input("Client Name *", key="fn_client_inp",
                                      placeholder="Chief Emeka Obi / ABC Ltd")
            fn_matter = st.text_input("Matter Description *", key="fn_matter_inp",
                                      placeholder="Purchase of property at No. 5 Bourdillon, Ikoyi")
            fn_ref = st.text_input("Our Reference", key="fn_ref_inp",
                                   placeholder="EO/2025/001")
            fn_date = st.text_input("Date", value=datetime.now().strftime("%d %B %Y"),
                                    key="fn_date_inp")
        with fn2:
            fn_land = st.number_input("Professional Fees (₦)", min_value=0.0,
                value=float(st.session_state.get("fn_land_fee", 0) or 0),
                step=1_000.0, format="%.2f", key="fn_land_inp")
            fn_stamp = st.number_input("Stamp Duty Paid (₦)", min_value=0.0,
                value=float(st.session_state.get("fn_stamp_duty", 0) or 0),
                step=500.0, format="%.2f", key="fn_stamp_inp")
            fn_court_fee_val = st.number_input("Court / Registry Fees (₦)", min_value=0.0,
                value=float(st.session_state.get("fn_court_fee", 0) or 0),
                step=500.0, format="%.2f", key="fn_court_inp")
            fn_disbursements = st.number_input("Other Disbursements (₦)", min_value=0.0,
                value=0.0, step=500.0, format="%.2f", key="fn_disb_inp")
            fn_vat = st.checkbox("Add 7.5% VAT on professional fees", value=True, key="fn_vat_chk")

        fn_notes = st.text_area("Additional Notes / Description of Services",
                                height=80, key="fn_notes_inp",
                                placeholder="E.g. Includes perfection of title, CAC searches, preparation of Deed of Assignment, and obtaining Governor's Consent.")

        gen_btn = st.button("🧾 Generate Fee Note", type="primary",
                            key="fn_gen_btn", width='stretch',
                            disabled=not (fn_client.strip() and fn_matter.strip()))

        if gen_btn:
            vat_amount = fn_land * 0.075 if fn_vat else 0.0
            subtotal = fn_land + fn_stamp + fn_court_fee_val + fn_disbursements
            total_due = subtotal + vat_amount
            firm = get_firm_name()
            lawyer = st.session_state.profile.get("lawyer_name", "")
            firm_address = st.session_state.profile.get("address", "")
            firm_email = st.session_state.profile.get("email", "")
            firm_phone = st.session_state.profile.get("phone", "")

            fee_note_text = f"""
{'='*65}
{firm.upper() if firm and firm != 'LexiAssist' else 'LAW FIRM'}
{'SOLICITORS & ADVOCATES' if firm and firm != 'LexiAssist' else ''}
{firm_address}
Tel: {firm_phone}  |  Email: {firm_email}
{'='*65}

PROFESSIONAL FEE NOTE / BILL OF COSTS

Date:         {fn_date}
Our Ref:      {fn_ref or '[REF]'}
To:           {fn_client}

RE: {fn_matter}

{'─'*65}

DESCRIPTION OF SERVICES:
{fn_notes or 'Professional legal services rendered in connection with the above matter.'}

{'─'*65}

FEES AND DISBURSEMENTS:
{'─'*65}"""

            items = []
            if fn_land > 0:
                items.append(("Professional / Solicitor's Fees", fn_land))
            if fn_stamp > 0:
                items.append(("Stamp Duty on Instrument", fn_stamp))
            if fn_court_fee_val > 0:
                items.append(("Court / Registry Filing Fees", fn_court_fee_val))
            if fn_disbursements > 0:
                items.append(("Other Disbursements", fn_disbursements))

            for desc, val in items:
                fee_note_text += f"\n  {desc:<45} {fmt_currency(val):>15}"

            fee_note_text += f"""
{'─'*65}
  Sub-Total                                          {fmt_currency(subtotal):>15}"""
            if vat_amount > 0:
                fee_note_text += f"""
  VAT @ 7.5% (on professional fees)                 {fmt_currency(vat_amount):>15}"""
            fee_note_text += f"""
{'─'*65}
  TOTAL AMOUNT DUE                                   {fmt_currency(total_due):>15}
{'='*65}

PAYMENT:
Kindly remit the sum of {fmt_currency(total_due)} to:
  Bank:           [BANK NAME]
  Account Name:   {firm or '[FIRM NAME]'}
  Account No:     [ACCOUNT NUMBER]

Payment is due within 14 days of this fee note.
Kindly quote our reference {fn_ref or '[REF]'} on all payments.

{'─'*65}
{lawyer or '[AUTHORISED SIGNATORY]'}
For: {firm or '[FIRM NAME]'}

⚠️ All fees are subject to the Rules of Professional Conduct for
Legal Practitioners 2007 and are subject to review if the matter
becomes more complex than currently anticipated.
{'='*65}
"""
            st.session_state["fn_generated"] = fee_note_text
            st.session_state["fn_total_due"] = total_due

        fn_generated = st.session_state.get("fn_generated", "")
        if fn_generated:
            st.markdown("---")
            st.markdown(f'<div class="response-box" style="font-family:monospace;font-size:0.88rem;">'
                        f'{esc(fn_generated)}</div>', unsafe_allow_html=True)

            total_due = st.session_state.get("fn_total_due", 0)
            st.success(f"✅ Total Due: **{fmt_currency(total_due)}**")

            fn_fname = f"FeeNote_{fn_client.replace(' ','_')}_{datetime.now():%Y%m%d}"
            fne1, fne2, fne3, fne4 = st.columns(4)
            with fne1:
                st.download_button("📥 TXT", export_txt(fn_generated, "Professional Fee Note"),
                    f"{fn_fname}.txt", "text/plain", key="fn_dl_txt", width='stretch')
            with fne2:
                st.download_button("📥 HTML", export_html(fn_generated, "Professional Fee Note"),
                    f"{fn_fname}.html", "text/html", key="fn_dl_html", width='stretch')
            with fne3:
                safe_pdf_download(fn_generated, "Professional Fee Note", fn_fname, "fn_dl_pdf")
            with fne4:
                safe_docx_download(fn_generated, "Professional Fee Note", fn_fname, "fn_dl_docx", doc_type="fee_note")


# ═══════════════════════════════════════════════════════
# PAGE: SETTLEMENT & ADR ADVISOR
# ═══════════════════════════════════════════════════════
def render_settlement_advisor():
    st.markdown("""<div class="page-header">
        <h2>🤝 Settlement & ADR Advisor</h2>
        <p>AI-powered negotiation strategy · Settlement value analysis ·
        Without-prejudice offer drafting · ADR route recommendation</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    st.markdown("""
<div style="background:#eff6ff;border-left:3px solid #3b82f6;padding:0.8rem 1.1rem;
border-radius:0.5rem;margin-bottom:1rem;font-size:0.9rem;">
  💡 <strong>How to use:</strong> Fill in the matter details below. The more specific
  your inputs (especially claim amount and case facts), the more actionable the output.
  The AI will take a firm position on the optimal settlement figure and strategy.
</div>""", unsafe_allow_html=True)

    sa1, sa2 = st.columns([3, 2])
    with sa1:
        sa_facts = st.text_area(
            "Case Facts *",
            height=200,
            key="sa_facts_ta",
            placeholder="""E.g. Client (ABC Ltd) entered into a construction contract with XYZ Builders Ltd
in January 2023 for ₦85M. XYZ abandoned the site in August 2023 after collecting
₦60M (70%). Completion was 35%. ABC incurred ₦25M additional costs to complete
with another contractor. ABC also suffered 6 months' revenue loss estimated at ₦15M.
XYZ claims ABC refused to pay the last instalment of ₦10M. ABC disputes this.""",
        )
        sa_case_type = st.selectbox(
            "Case Type",
            [
                "Breach of Contract", "Debt Recovery", "Property Dispute",
                "Employment / Wrongful Termination", "Company Dispute",
                "Personal Injury / Negligence", "Defamation", "Matrimonial",
                "Construction / Engineering", "Banking & Finance",
                "Intellectual Property", "Other",
            ],
            key="sa_case_type_sel",
        )

    with sa2:
        sa_instructing = st.text_input("Instructing Party", placeholder="ABC Ltd (Claimant)",
                                       key="sa_instruct_inp")
        sa_opposing = st.text_input("Opposing Party", placeholder="XYZ Builders Ltd (Defendant)",
                                    key="sa_oppose_inp")
        sa_amount = st.number_input("Total Claim / Dispute Value (₦)",
            min_value=0.0, value=100_000_000.0, step=1_000_000.0,
            format="%.2f", key="sa_amount_inp")
        st.caption(f"Claim: **{fmt_currency(sa_amount)}**")
        sa_court = st.selectbox("Current Stage",
            ["Pre-litigation", "Letter of Demand sent", "Writ filed / Suit pending",
             "Pleadings stage", "Pre-trial / Mediation ordered", "Trial ongoing",
             "Judgment obtained", "Appeal pending", "Arbitration commenced"],
            key="sa_court_sel")
        sa_strength = st.select_slider(
            "Your Case Strength",
            options=["Very Weak", "Weak", "Moderate", "Strong", "Very Strong"],
            value="Moderate", key="sa_strength_sl",
        )
        sa_urgency = st.selectbox("Time Pressure",
            ["None", "Client needs cash urgently", "Court deadline approaching",
             "Business disruption ongoing", "Preservation risk (assets at risk)",
             "Limitation period approaching"],
            key="sa_urgency_sel")

    mode = st.session_state.response_mode
    st.info(f"Mode: {RESPONSE_MODES[mode]['label']}")
    sa_btn = st.button(
        "🤝 Generate Settlement Strategy",
        type="primary", width='stretch', key="sa_gen_btn",
        disabled=not (sa_facts.strip() and sa_instructing.strip()),
    )

    if sa_btn and sa_facts.strip() and sa_instructing.strip():
        prompt = SETTLEMENT_PROMPT.format(
            instructing_party=sa_instructing.strip(),
            opposing_party=sa_opposing.strip() or "Opposing party",
            case_type=sa_case_type,
            claim_amount=f"{sa_amount:,.2f}",
            court_stage=sa_court,
            strength=sa_strength,
            urgency=sa_urgency,
            case_facts=sa_facts.strip(),
        )
        with st.spinner("🤝 Analysing settlement position and generating strategy…"):
            raw = generate(prompt, SETTLEMENT_SYSTEM, mode, "advisory")
        st.session_state["sa_result"] = raw
        st.session_state["sa_matter_label"] = sa_instructing.strip()
        add_to_history(f"[Settlement] {sa_instructing.strip()} vs {sa_opposing.strip()}", raw, "advisory", mode)
        st.rerun()

    result = st.session_state.get("sa_result", "")
    matter_label = st.session_state.get("sa_matter_label", "Settlement")

    if result:
        st.markdown("---")

        # Parse sections
        def _get_section(text, header):
            lines = text.split("\n")
            capture, collected = False, []
            for line in lines:
                if header.upper() in line.upper() and "═" in line:
                    capture = True; continue
                if capture and "═══" in line and collected:
                    break
                if capture:
                    collected.append(line)
            return "\n".join(collected).strip()

        sec1 = _get_section(result, "SETTLEMENT VALUE")
        sec2 = _get_section(result, "NEGOTIATION STRATEGY")
        sec3 = _get_section(result, "ADR ROUTE")
        sec4 = _get_section(result, "WITHOUT PREJUDICE")
        sec5 = _get_section(result, "RISK IF NO SETTLEMENT")

        if sec1 and sec2:
            t1, t2, t3, t4, t5 = st.tabs([
                "💰 Settlement Value",
                "♟️ Negotiation Strategy",
                "🏛️ ADR Route",
                "✉️ Without Prejudice Offer",
                "⚠️ Litigation Risk",
            ])
            tab_configs = [
                (t1, sec1, "#f0fdf4", "#059669"),
                (t2, sec2, "#eff6ff", "#3b82f6"),
                (t3, sec3, "#f5f3ff", "#7c3aed"),
                (t4, sec4, "#fffbeb", "#f59e0b"),
                (t5, sec5, "#fef2f2", "#dc2626"),
            ]
            for tab, content, bg, border in tab_configs:
                with tab:
                    st.markdown(f"""
<div style="background:{bg};border-left:4px solid {border};border-radius:0.75rem;
padding:1.5rem;line-height:1.8;white-space:pre-wrap;font-size:0.95rem;">
{esc(content)}</div>""", unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="response-box">{esc(result)}</div>', unsafe_allow_html=True)

        # Save to case
        cases = st.session_state.cases
        if cases:
            st.markdown("---")
            sv1, sv2 = st.columns([3, 1])
            with sv1:
                sc_id = st.selectbox("Save to Case", [c["id"] for c in cases],
                    format_func=lambda x: next((c["title"] for c in cases if c["id"] == x), x),
                    key="sa_save_case_sel")
            with sv2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("💾 Save", key="sa_save_btn", type="primary", width='stretch'):
                    save_analysis_to_case(sc_id, f"[Settlement] {matter_label}", result, "advisory", mode)
                    st.success("✅ Saved to case.")

        # Export
        fname = f"Settlement_{matter_label.replace(' ','_')}_{datetime.now():%Y%m%d_%H%M}"
        e1, e2, e3, e4 = st.columns(4)
        with e1:
            st.download_button("📥 TXT", export_txt(result, f"Settlement Strategy — {matter_label}"),
                f"{fname}.txt", "text/plain", key="sa_dl_txt", width='stretch')
        with e2:
            st.download_button("📥 HTML", export_html(result, f"Settlement Strategy — {matter_label}"),
                f"{fname}.html", "text/html", key="sa_dl_html", width='stretch')
        with e3:
            safe_pdf_download(result, f"Settlement Strategy — {matter_label}", fname, "sa_dl_pdf")
        with e4:
            safe_docx_download(result, f"Settlement Strategy — {matter_label}", fname, "sa_dl_docx", doc_type="settlement", meta={"matter": matter_label})

        if st.button("🗑️ Clear", key="sa_clear_btn", width='stretch'):
            st.session_state["sa_result"] = ""
            st.rerun()

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> Settlement strategy is AI-assisted advisory.
            All without-prejudice communications must be reviewed by counsel before transmission.
            Counsel remains professionally responsible for all advice and negotiations.
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: DUE DILIGENCE ENGINE
# ═══════════════════════════════════════════════════════
def render_due_diligence():
    st.markdown("""<div class="page-header">
        <h2>🔎 Due Diligence Engine</h2>
        <p>AI-generated Nigerian-law due diligence checklists · CAC searches ·
        Title verification · Regulatory approvals · Transaction risk flags</p>
    </div>""", unsafe_allow_html=True)

    if not st.session_state.api_configured:
        st.warning("⚠️ Connect your API key first.")
        return

    st.markdown("""
<div style="background:#f0fdf4;border-left:3px solid #059669;padding:0.8rem 1.1rem;
border-radius:0.5rem;margin-bottom:1rem;font-size:0.9rem;">
  💡 <strong>How it works:</strong> Select your transaction type, describe the deal,
  and the AI generates a comprehensive Nigerian-law due diligence checklist tailored
  to your specific matter — with search requirements, risk flags, and a critical path.
</div>""", unsafe_allow_html=True)

    dd1, dd2 = st.columns([3, 2])
    with dd1:
        dd_description = st.text_area(
            "Transaction Description *",
            height=180,
            key="dd_desc_ta",
            placeholder="""Describe the transaction briefly. E.g.:

Client (Sunrise Properties Ltd) intends to acquire a 3-plot commercial property
in Victoria Island, Lagos from Apex Holdings Ltd for ₦450M. The vendor claims
a C of O has been registered in their name. The property currently has a tenant.
Client wants to develop a 12-storey mixed-use building on the site. The vendor
is a company with 3 directors. No prior relationship with vendor.""",
        )
        dd_concerns = st.text_area(
            "Special Concerns / Red Flags (optional)",
            height=80, key="dd_concerns_ta",
            placeholder="E.g. Vendor is a company with recent change of directors. Prior occupant dispute. Area prone to government acquisition.",
        )
    with dd2:
        dd_type = st.selectbox(
            "Transaction Type *",
            list(DD_TRANSACTION_TYPES.keys()),
            format_func=lambda k: DD_TRANSACTION_TYPES[k],
            key="dd_type_sel",
        )
        dd_value = st.number_input(
            "Transaction Value (₦)",
            min_value=0.0, value=100_000_000.0, step=5_000_000.0,
            format="%.2f", key="dd_value_inp",
        )
        st.caption(f"Value: **{fmt_currency(dd_value)}**")
        dd_jurisdiction = st.selectbox(
            "Jurisdiction",
            ["Lagos State", "Abuja (FCT)", "Rivers State", "Kano State",
             "Ogun State", "Oyo State", "Delta State", "Anambra State",
             "Cross River State", "Edo State", "Other — specify in description"],
            key="dd_jur_sel",
        )
        dd_parties = st.text_input(
            "Parties",
            placeholder="Buyer: Sunrise Properties Ltd | Seller: Apex Holdings Ltd",
            key="dd_parties_inp",
        )
        mode = st.session_state.response_mode
        st.info(f"Mode: {RESPONSE_MODES[mode]['label']}")

    dd_btn = st.button(
        f"🔎 Generate Due Diligence Checklist",
        type="primary", width='stretch', key="dd_gen_btn",
        disabled=not dd_description.strip(),
    )

    if dd_btn and dd_description.strip():
        prompt = DD_PROMPT.format(
            transaction_type=DD_TRANSACTION_TYPES[dd_type],
            transaction_value=f"{dd_value:,.2f}",
            jurisdiction=dd_jurisdiction,
            parties=dd_parties.strip() or "As described",
            description=dd_description.strip(),
            special_concerns=dd_concerns.strip() or "None stated",
        )
        with st.spinner(f"🔎 Generating {DD_TRANSACTION_TYPES[dd_type]} due diligence checklist…"):
            raw = generate(prompt, DD_SYSTEM, mode, "advisory")
        st.session_state["dd_result"] = raw
        st.session_state["dd_label"] = f"{DD_TRANSACTION_TYPES[dd_type]} — {dd_jurisdiction}"
        add_to_history(
            f"[Due Diligence] {DD_TRANSACTION_TYPES[dd_type]} — {dd_value:,.0f}",
            raw, "advisory", mode,
        )
        st.rerun()

    result = st.session_state.get("dd_result", "")
    dd_label = st.session_state.get("dd_label", "Due Diligence")

    if result:
        st.markdown("---")
        st.markdown(f"### 🔎 {esc(dd_label)}")

        # Render with colored block
        st.markdown(f"""
<div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:0.85rem;
padding:1.8rem;line-height:1.85;white-space:pre-wrap;font-size:0.93rem;">
{esc(result)}</div>""", unsafe_allow_html=True)

        # Save to case
        cases = st.session_state.cases
        if cases:
            st.markdown("---")
            dv1, dv2 = st.columns([3, 1])
            with dv1:
                dd_case_id = st.selectbox("Save to Case File", [c["id"] for c in cases],
                    format_func=lambda x: next((c["title"] for c in cases if c["id"] == x), x),
                    key="dd_save_case_sel")
            with dv2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("💾 Save to Case", key="dd_save_btn",
                             type="primary", width='stretch'):
                    save_analysis_to_case(dd_case_id, f"[DD Checklist] {dd_label}",
                                          result, "advisory", mode)
                    st.success("✅ Saved to case file.")

        # Export
        fname = f"DueDiligence_{dd_type}_{datetime.now():%Y%m%d_%H%M}"
        de1, de2, de3, de4 = st.columns(4)
        with de1:
            st.download_button("📥 TXT", export_txt(result, f"Due Diligence — {dd_label}"),
                f"{fname}.txt", "text/plain", key="dd_dl_txt", width='stretch')
        with de2:
            st.download_button("📥 HTML", export_html(result, f"Due Diligence — {dd_label}"),
                f"{fname}.html", "text/html", key="dd_dl_html", width='stretch')
        with de3:
            safe_pdf_download(result, f"Due Diligence — {dd_label}", fname, "dd_dl_pdf")
        with de4:
            safe_docx_download(result, f"Due Diligence — {dd_label}", fname, "dd_dl_docx", doc_type="due_diligence", meta={"subject": dd_label})

        if st.button("🗑️ Clear", key="dd_clear_btn", width='stretch'):
            st.session_state["dd_result"] = ""
            st.rerun()

        st.markdown("""<div class="disclaimer">
            <strong>⚖️ Disclaimer:</strong> This checklist is AI-generated and must be reviewed
            by counsel with knowledge of the specific transaction. It does not replace physical
            inspection, official searches, or independent legal advice. All searches must be
            conducted at the relevant registries before advising clients to proceed.
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# PAGE: ADMIN — USER MANAGEMENT
# ═══════════════════════════════════════════════════════
def render_user_management():
    if st.session_state.get("current_user_role") != "admin":
        st.error("🚫 Admin access required.")
        return

    st.markdown("""<div class="page-header">
        <h2>🛡️ User Management</h2>
        <p>Create accounts, manage roles, reset passwords, and remove users</p>
    </div>""", unsafe_allow_html=True)

    db = get_db()
    um_list, um_create, um_stats = st.tabs(["👥 All Users", "➕ Create User", "📊 Usage Stats"])

    # ── All Users ──
    with um_list:
        users = db.list_users()
        current_uid = st.session_state.get("current_user_id", "")
        st.markdown(f"##### 👥 {len(users)} Registered User(s)")

        for user in users:
            uid = user["user_id"]
            is_self = (uid == current_uid)
            role_color = "#059669" if user["role"] == "admin" else "#3b82f6"
            role_label = "🛡️ Admin" if user["role"] == "admin" else "👤 User"

            with st.expander(
                f"{role_label} · @{user['username']} — {user.get('lawyer_name','') or user.get('firm_name','') or ''}",
                expanded=False,
            ):
                u1, u2, u3 = st.columns(3)
                with u1:
                    st.markdown(f"**Username:** @{esc(user['username'])}")
                    st.markdown(f"**Role:** {role_label}")
                    st.markdown(f"**Email:** {esc(user.get('email','') or '—')}")
                with u2:
                    st.markdown(f"**Full Name:** {esc(user.get('lawyer_name','') or '—')}")
                    st.markdown(f"**Firm:** {esc(user.get('firm_name','') or '—')}")
                with u3:
                    st.markdown(f"**Joined:** {esc(fmt_date(user.get('created_at','')))}")
                    st.markdown(f"**Last Login:** {esc(fmt_date(user.get('last_login','')))}")

                st.markdown("---")
                act1, act2, act3 = st.columns(3)

                # Change role
                with act1:
                    if not is_self:
                        new_role = "user" if user["role"] == "admin" else "admin"
                        role_btn_label = f"⬇️ Demote to User" if user["role"] == "admin" else "⬆️ Promote to Admin"
                        if st.button(role_btn_label, key=f"um_role_{uid}", width='stretch'):
                            db.update_user(uid, {"role": new_role})
                            st.success(f"✅ @{user['username']} is now {new_role}.")
                            st.rerun()
                    else:
                        st.caption("(Your own account)")

                # Reset password
                with act2:
                    with st.popover(f"🔑 Reset Password", width='stretch'):
                        with st.form(f"reset_pw_{uid}"):
                            new_temp_pw = st.text_input("New Password", type="password", key=f"tmp_pw_{uid}")
                            if st.form_submit_button("✅ Set Password"):
                                if len(new_temp_pw) < 6:
                                    st.error("Min 6 characters.")
                                else:
                                    db.update_user(uid, {"password_hash": hash_password(new_temp_pw)})
                                    st.success(f"✅ Password reset for @{user['username']}.")

                # Delete user
                with act3:
                    if not is_self:
                        with st.popover(f"🗑️ Delete User", width='stretch'):
                            st.warning(f"Delete @{user['username']}? ALL their data will be permanently erased.")
                            if st.button(f"⚠️ Confirm Delete @{user['username']}",
                                         key=f"um_del_confirm_{uid}", type="primary"):
                                db.delete_user(uid)
                                st.success(f"✅ @{user['username']} deleted.")
                                st.rerun()
                    else:
                        st.caption("Cannot delete yourself.")

    # ── Create User ──
    with um_create:
        st.markdown("##### ➕ Create a New User Account")
        st.caption("Create accounts for colleagues at your firm. They can log in immediately.")
        render_register_form("admin_new_user", admin_mode=True)

    # ── Usage Stats ──
    with um_stats:
        st.markdown("##### 📊 Platform Usage by User")
        users = db.list_users()
        if not users:
            st.info("No users yet.")
        else:
            for user in users:
                uid = user["user_id"]
                cur = db._execute(
                    "SELECT COUNT(*), COALESCE(SUM(estimated_cost),0) FROM cost_logs WHERE user_id = %s",
                    (uid,)
                )
                row = cur.fetchone()
                calls, cost = (row[0], row[1]) if row else (0, 0)
                cur2 = db._execute(
                    "SELECT COUNT(*) FROM kv_store WHERE key LIKE %s", (f"u:{uid}:cases",)
                )
                # Get case count from namespaced kv
                st.markdown(f"""
<div class="custom-card">
  <div style="display:flex;justify-content:space-between;">
    <strong>@{esc(user['username'])}</strong>
    <span class="badge {'badge-ok' if user['role'] == 'admin' else 'badge-info'}">
      {'Admin' if user['role'] == 'admin' else 'User'}
    </span>
  </div>
  <small>🤖 {calls} AI calls · 💰 ${cost:.4f} estimated cost · 
  🕐 Last login: {esc(fmt_date(user.get('last_login','')))}
  </small>
</div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════
def _get_genai_client(key: str):
    return genai.Client(api_key=key)


def add_client(data: dict):
    data["id"] = new_id()
    data["created_at"] = datetime.now().isoformat()
    st.session_state.clients.append(data)
    persist("clients")


def add_time_entry(data: dict):
    data["id"] = new_id()
    data["created_at"] = datetime.now().isoformat()
    data["amount"] = data.get("hours", 0) * data.get("rate", 0)
    st.session_state.time_entries.append(data)
    persist("time_entries")


def add_to_history(query: str, response: str, task: str, mode: str):
    entry = {
        "id": new_id(),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "query": query,
        "response": response,
        "task": task,
        "mode": mode,
        "word_count": len(response.split()),
    }
    st.session_state.chat_history.append(entry)
    persist("chat_history")
    return entry


def auto_connect():
    if st.session_state.api_configured:
        return
    k = _resolve_api_key()
    if k:
        try:
            # Validate key is non-empty; actual connection tested lazily in generate()
            st.session_state.api_key = k
            st.session_state.api_configured = True
            m = safe_secret("GEMINI_MODEL") or os.getenv("GEMINI_MODEL", "")
            if m and m in SUPPORTED_MODELS:
                st.session_state.gemini_model = m
        except Exception as e:
            logger.warning(f"Auto-connect failed: {e}")


def build_system_prompt(task: str, mode: str) -> str:
    base = PROMPTS_BY_MODE.get(mode, PROMPTS_BY_MODE["standard"])
    modifier = TASK_MODIFIERS.get(task, TASK_MODIFIERS["general"])
    return base + modifier


def client_billable(cid: str) -> float:
    return sum(e.get("amount", 0) for e in st.session_state.time_entries if e.get("client_id") == cid)


def client_case_count(cid: str) -> int:
    return sum(1 for c in st.session_state.cases if c.get("client_id") == cid)


def days_until(d) -> int:
    if not d:
        return 9999
    try:
        if isinstance(d, str):
            d = datetime.fromisoformat(d).date()
        if isinstance(d, datetime):
            d = d.date()
        return (d - date.today()).days
    except Exception:
        return 9999


def delete_case(cid: str):
    st.session_state.cases = [c for c in st.session_state.cases if c["id"] != cid]
    persist("cases")
    get_db().delete_case_analyses_for_case(cid)


def delete_client(cid: str):
    st.session_state.clients = [c for c in st.session_state.clients if c["id"] != cid]
    persist("clients")


def delete_time_entry(eid: str):
    st.session_state.time_entries = [e for e in st.session_state.time_entries if e["id"] != eid]
    persist("time_entries")


def estimate_cost(input_text: str, output_text: str) -> float:
    """Estimate API cost from text lengths."""
    input_tokens = len(input_text) / 4
    output_tokens = len(output_text) / 4
    cost = (input_tokens / 1_000_000) * COST_PER_1M_INPUT + (output_tokens / 1_000_000) * COST_PER_1M_OUTPUT
    return round(cost, 6)


def fmt_currency(amount) -> str:
    try:
        return f"₦{float(amount):,.2f}"
    except Exception:
        return "₦0.00"


def fmt_date(d) -> str:
    if not d:
        return "—"
    try:
        if isinstance(d, str):
            d = datetime.fromisoformat(d)
        return d.strftime("%d %b %Y")
    except Exception:
        return str(d)


def generate(prompt: str, system: str, mode: str, task: str = "general") -> str:
    """Core generation with retry, cost logging, and proper token limits."""
    k = _resolve_api_key()
    if not k:
        return "⚠️ No API key configured. Please set up your key."
    mode_cfg = RESPONSE_MODES.get(mode, RESPONSE_MODES["standard"])
    gen_config = _genai_types.GenerateContentConfig(
        system_instruction=system,
        temperature=mode_cfg["temp"],
        top_p=0.92,
        top_k=40,
        max_output_tokens=mode_cfg["tokens"],
    )
    client = _get_genai_client(k)

    for attempt in range(3):
        try:
            resp = client.models.generate_content(
                model=st.session_state.gemini_model,
                contents=prompt,
                config=gen_config,
            )
            if resp and resp.text:
                # Log cost
                cost = estimate_cost(prompt + system, resp.text)
                db = get_db()
                db.add_cost_log({
                    "id": new_id(),
                    "timestamp": datetime.now().isoformat(),
                    "model": st.session_state.gemini_model,
                    "task": task,
                    "mode": mode,
                    "input_chars": len(prompt) + len(system),
                    "output_chars": len(resp.text),
                    "estimated_cost": cost,
                    "query_preview": prompt[:120],
                })
                return resp.text
            return "⚠️ Empty response from AI. Try rephrasing your query."
        except Exception as e:
            if attempt == 2:
                return f"⚠️ Generation error after 3 attempts: {str(e)[:200]}"
            time.sleep(2 * (attempt + 1))
    return "⚠️ Generation failed. Please try again."


def get_active_cases() -> list:
    return [c for c in st.session_state.cases if c.get("status") == "Active"]


def get_all_limitation_periods() -> list:
    custom = st.session_state.get("custom_limitation_periods", [])
    return DEFAULT_LIMITATION_PERIODS + custom


def get_all_maxims() -> list:
    custom = st.session_state.get("custom_maxims", [])
    return DEFAULT_LEGAL_MAXIMS + custom


def get_all_templates() -> list:
    """Combine built-in and custom templates."""
    custom = st.session_state.get("custom_templates", [])
    return DEFAULT_TEMPLATES + custom


def get_client_name(cid: str) -> str:
    for c in st.session_state.clients:
        if c["id"] == cid:
            return c.get("name", "—")
    return "—"


def get_firm_name() -> str:
    """Get firm name for branding on exports."""
    profile = st.session_state.get("profile", {})
    return profile.get("firm_name", "") or "LexiAssist"


def get_hearings() -> list:
    h = []
    for c in st.session_state.cases:
        if c.get("next_hearing") and c.get("status") in ("Active", "Pending"):
            h.append({
                "id": c["id"], "title": c.get("title", ""),
                "date": c["next_hearing"], "court": c.get("court", ""),
                "suit": c.get("suit_no", ""), "status": c.get("status", ""),
            })
    h.sort(key=lambda x: x.get("date", "z"))
    return h


def make_invoice(client_id: str):
    entries = [e for e in st.session_state.time_entries if e.get("client_id") == client_id]
    if not entries:
        return None
    inv = {
        "id": new_id(),
        "invoice_no": f"INV-{datetime.now():%Y%m%d}-{new_id()[:4].upper()}",
        "client_id": client_id,
        "client_name": get_client_name(client_id),
        "entries": entries,
        "total": sum(e.get("amount", 0) for e in entries),
        "date": datetime.now().isoformat(),
        "status": "Draft",
    }
    st.session_state.invoices.append(inv)
    persist("invoices")
    return inv


def manual_connect(key: str) -> bool:
    try:
        client = _get_genai_client(key)
        client.models.generate_content(
            model=st.session_state.gemini_model,
            contents="Test",
            config=_genai_types.GenerateContentConfig(max_output_tokens=10),
        )
        st.session_state.api_key = key
        st.session_state.api_configured = True
        return True
    except Exception as e:
        err = str(e)
        if "403" in err:
            st.error("❌ Invalid API key.")
        elif "429" in err:
            st.error("⚠️ Rate limit — try again shortly.")
        else:
            st.error(f"❌ Connection failed: {err[:120]}")
        return False


def new_id() -> str:
    return uuid.uuid4().hex[:8]


def relative_date(d) -> str:
    n = days_until(d)
    if n == 9999:
        return "—"
    if n < 0:
        return f"{abs(n)}d overdue"
    if n == 0:
        return "TODAY"
    if n == 1:
        return "Tomorrow"
    if n <= 7:
        return f"{n} days"
    return f"{n} days away"


def run_comparison(entry_a: dict, entry_b: dict) -> str:
    prompt = (
        f"ANALYSIS A (from {entry_a.get('timestamp', '')}):\n"
        f"Query: {entry_a.get('query', '')}\n"
        f"Response:\n{entry_a.get('response', '')}\n\n"
        f"{'='*40}\n\n"
        f"ANALYSIS B (from {entry_b.get('timestamp', '')}):\n"
        f"Query: {entry_b.get('query', '')}\n"
        f"Response:\n{entry_b.get('response', '')}"
    )
    return generate(prompt, COMPARISON_PROMPT, "standard", "analysis")


def run_critique(query: str, analysis: str) -> str:
    prompt = f"ORIGINAL QUERY:\n{query}\n\nANALYSIS TO REVIEW:\n{analysis}"
    return generate(prompt, CRITIQUE_PROMPT, "brief", "analysis")


def run_followup(original: str, previous: str, followup: str, mode: str) -> str:
    prompt = f"ORIGINAL QUERY:\n{original}\n\nPREVIOUS ANALYSIS:\n{previous}\n\nFOLLOW-UP QUESTION:\n{followup}"
    return generate(prompt, FOLLOWUP_PROMPT, mode, "general")


def run_issue_spot(query: str) -> str:
    return generate(query, ISSUE_SPOT_PROMPT, "brief", "analysis")


def run_research(query: str, mode: str) -> str:
    system = build_system_prompt("research", mode)
    return generate(query, system, mode, "research")


def safe_secret(key: str, default: str = "") -> str:
    try:
        return st.secrets[key]
    except Exception:
        return default


def save_analysis_to_case(case_id: str, query: str, response: str, task: str, mode: str):
    """Attach an AI analysis to a specific case."""
    db = get_db()
    db.add_case_analysis(case_id, {
        "id": new_id(),
        "query": query,
        "response": response,
        "task": task,
        "mode": mode,
        "timestamp": datetime.now().isoformat(),
    })



def total_billable() -> float:
    return sum(e.get("amount", 0) for e in st.session_state.time_entries)


def total_hours() -> float:
    return sum(e.get("hours", 0) for e in st.session_state.time_entries)


def update_case(cid: str, updates: dict):
    for c in st.session_state.cases:
        if c["id"] == cid:
            c.update(updates)
            c["updated_at"] = datetime.now().isoformat()
    persist("cases")

def init_session_state():
    """Set non-user-specific session defaults. Called every render cycle."""
    simple_defaults = {
        "api_key": "",
        "api_configured": False,
        "gemini_model": DEFAULT_MODEL,
        "theme": "⚖️ Corporate",
        "font_size_scale": 1.0,
        "high_contrast": False,
        "reduce_motion": False,
        "response_mode": "standard",
        "authenticated": False,
        "current_user_id": "",
        "current_username": "",
        "current_user_role": "",
        "user_data_loaded": False,
        "last_response": "",
        "original_query": "",
        "last_task": "general",
        "last_mode": "standard",
        "research_results": "",
        "loaded_template": "",
        "imported_doc": None,
        "selected_history_idx": None,
        "compare_selections": [],
    }
    for k, v in simple_defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v



def esc(text: str) -> str:
    if not text:
        return ""
    return html_mod.escape(str(text))


def extract_file_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        if not HAS_PDF_READ:
            raise ValueError("PDF support not available (install pdfplumber)")
        with pdfplumber.open(BytesIO(data)) as pdf:
            pages = []
            for p in pdf.pages:
                txt = p.extract_text()
                if txt:
                    pages.append(txt)
            return "\n\n".join(pages)
    elif name.endswith((".docx", ".doc")):
        if not HAS_DOCX:
            raise ValueError("DOCX support not available (install python-docx)")
        doc = DocxDocument(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif name.endswith(".txt") or name.endswith(".rtf"):
        return data.decode("utf-8", errors="ignore")
    elif name.endswith((".xlsx", ".xls")):
        if not HAS_XLSX:
            raise ValueError("Excel support not available (install openpyxl)")
        df = pd.read_excel(BytesIO(data))
        return df.to_string(index=False)
    elif name.endswith(".csv"):
        df = pd.read_csv(BytesIO(data))
        return df.to_string(index=False)
    elif name.endswith(".json"):
        obj = json.loads(data.decode("utf-8", errors="ignore"))
        return json.dumps(obj, indent=2)
    else:
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            raise ValueError(f"Unsupported file type: {name}")


def run_ai_query(query: str, task: str, mode: str, context: str = "") -> str:
    system = build_system_prompt(task, mode)
    full_prompt = query
    if context:
        full_prompt = f"DOCUMENT CONTEXT:\n{context[:8000]}\n\nQUERY:\n{query}"
    return generate(full_prompt, system, mode, task)


def add_case(data: dict):
    data["id"] = new_id()
    data["created_at"] = datetime.now().isoformat()
    st.session_state.cases.append(data)
    persist("cases")

def _resolve_api_key() -> str:
    for src in [
        lambda: safe_secret("GEMINI_API_KEY"),
        lambda: os.getenv("GEMINI_API_KEY", ""),
        lambda: st.session_state.get("api_key", ""),
    ]:
        k = src()
        if k and k.strip() and len(k.strip()) >= 10:
            return k.strip()
    return ""

def main():
    init_session_state()
    auto_connect()
    st.markdown(get_theme_css(
        st.session_state.theme,
        font_size_scale=st.session_state.get("font_size_scale", 1.0),
        high_contrast=st.session_state.get("high_contrast", False),
        reduce_motion=st.session_state.get("reduce_motion", False),
    ), unsafe_allow_html=True)

    # ── API setup gate ──
    if not st.session_state.api_configured:
        render_setup_screen()
        return

    db = get_db()
    db.ensure_connected()  # heal stale/aborted connections before any DB work

    # ── Auto-login from persistent session token (survives page refreshes) ──
    if not st.session_state.authenticated:
        try:
            token = st.query_params.get("t", "")
        except Exception:
            token = ""
        if token and not st.session_state.authenticated:
            do_auto_login_from_token(token)

    # ── Auth gate ──
    if not st.session_state.authenticated:
        if not db.has_any_users():
            render_create_admin_screen()
        else:
            render_login_screen()
        return

    # ── Load user data exactly once per session ──
    if not st.session_state.user_data_loaded:
        load_user_data()
        st.session_state.user_data_loaded = True

    render_sidebar()

    is_admin = (st.session_state.current_user_role == "admin")

    # ── TOP NAVIGATION TABS ──
    tab_labels = [
        "🏠 Home",
        "🧠 AI Assistant",
        "📚 Research",
        "📁 Cases",
        "⚡ Lifecycle",
        "📜 Pleadings",
        "🔍 Conflict Check",
        "📅 Calendar",
        "📋 Templates",
        "👥 Clients",
        "💰 Billing",
        "🔧 Tools",
        "📝 Notes → Brief",
        "🎯 Witness Prep",
        "📰 Legal News",
        "⚖️ Fee Calculator",
        "🤝 Settlement",
        "🔎 Due Diligence",
        "👤 Profile",
    ]
    if is_admin:
        tab_labels.append("🛡️ Admin")

    tabs = st.tabs(tab_labels)

    with tabs[0]:  render_home()
    with tabs[1]:  render_ai()
    with tabs[2]:  render_research()
    with tabs[3]:  render_cases()
    with tabs[4]:  render_lifecycle()
    with tabs[5]:  render_pleadings()
    with tabs[6]:  render_conflict_checker()
    with tabs[7]:  render_calendar()
    with tabs[8]:  render_templates()
    with tabs[9]:  render_clients()
    with tabs[10]: render_billing()
    with tabs[11]: render_tools()
    with tabs[12]: render_notes_converter()
    with tabs[13]: render_witness_prep()
    with tabs[14]: render_legal_news()
    with tabs[15]: render_fee_calculator()
    with tabs[16]: render_settlement_advisor()
    with tabs[17]: render_due_diligence()
    with tabs[18]: render_profile()
    if is_admin:
        with tabs[19]: render_user_management()

    # Footer
    st.markdown("---")
    firm = get_firm_name()
    firm_text = f"{esc(firm)} · " if firm and firm != "LexiAssist" else ""
    uname = st.session_state.get("current_username", "")
    user_text = f" · Signed in as @{esc(uname)}" if uname else ""
    st.caption(
        f"⚖️ {firm_text}LexiAssist v8.0 © 2026 · Elite AI Legal Engine for Nigerian Lawyers"
        f"{user_text} · ⚠️ AI-generated information — not legal advice — verify all citations independently"
    )

    # ── Keep-Alive Ping (st.iframe replaces deprecated components.v1.html) ────
    st.iframe(
        src="about:blank",
        height=0,
        scrolling=False,
    )
    # ────────────────────────────────────────────────────────────────────────────


if __name__ == "__main__":
    main()

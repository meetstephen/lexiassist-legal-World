"""
core.py — LexiAssist v7.0 Backend Engine
Smart legal AI with instruction adherence, response-length control, and persistence.
"""
from __future__ import annotations

import html, io, json, logging, os, re, time, uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import google.generativeai as genai
import pandas as pd

try:
    import pdfplumber
    HAS_PDF_READ = True
except ImportError:
    HAS_PDF_READ = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_PDF_WRITE = True
except ImportError:
    HAS_PDF_WRITE = False

try:
    import openpyxl  # noqa: F401
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

logger = logging.getLogger("LexiAssist")

# ═══════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════
SUPPORTED_MODELS = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]
DEFAULT_MODEL = "gemini-2.5-flash"
MODEL_MIGRATION = {
    "gemini-2.0-flash": "gemini-2.5-flash",
    "gemini-2.0-flash-001": "gemini-2.5-flash",
    "gemini-2.0-flash-lite": "gemini-2.5-flash-lite",
    "gemini-2.0-flash-lite-001": "gemini-2.5-flash-lite",
}
CASE_STATUSES = ["Active", "Pending", "Completed", "Archived"]
CLIENT_TYPES = ["Individual", "Corporate", "Government"]
HISTORY_FILE = Path(".lexiassist_history.json")
MAX_UPLOAD_MB = 10

TASK_TYPES = {
    "general":        {"label": "General Query",            "desc": "Any legal question",               "icon": "💬"},
    "analysis":       {"label": "Legal Analysis",           "desc": "Issue spotting & reasoning",        "icon": "🔍"},
    "drafting":       {"label": "Document Drafting",        "desc": "Contracts, pleadings, affidavits",  "icon": "📄"},
    "research":       {"label": "Legal Research",           "desc": "Case law, statutes, authorities",   "icon": "📚"},
    "procedure":      {"label": "Procedural Guidance",      "desc": "Court filing, rules, deadlines",    "icon": "📋"},
    "interpretation": {"label": "Statutory Interpretation", "desc": "Analyze legislation",               "icon": "⚖️"},
    "advisory":       {"label": "Client Advisory",          "desc": "Strategic advice & risk",           "icon": "🎯"},
}

RESPONSE_MODES = {
    "brief":         {"label": "⚡ Brief",         "desc": "Direct answer · 2-5 sentences",           "tokens": 600},
    "standard":      {"label": "📝 Standard",      "desc": "Key issues + guidance · 3-8 paragraphs",  "tokens": 2500},
    "comprehensive": {"label": "🔬 Comprehensive", "desc": "Full multi-pass analysis · CREAC",         "tokens": 10000},
}

# ═══════════════════════════════════════════════════════
# PROMPTS — Short, mode-aware, instruction-adherent
# ═══════════════════════════════════════════════════════
_CORE = """You are LexiAssist — an expert Nigerian legal AI assistant.
Jurisdiction: Nigeria (Constitution 1999 as amended, Federal/State Acts, Nigerian case law).

STRICT RULES:
1. NEVER fabricate case names, citations, or section numbers. If uncertain, state the principle and mark [Citation to be verified].
2. Answer ONLY what is asked. Do NOT add unrequested sections or go beyond the user's scope.
3. Match your response length EXACTLY to the response mode specified below.
4. Be precise, practical, and professionally rigorous.
5. If you don't know, say so. If it depends, say on what — concretely."""

SYSTEM_PROMPTS = {
    "brief": _CORE + """

RESPONSE MODE: BRIEF — Maximum 2-5 sentences.
- Direct legal answer. No headers, no preamble, no bullet lists.
- If missing facts matter, say "it depends on [X]" in ONE sentence.
- If there is a critical deadline risk, note it in ONE sentence.
- Do NOT elaborate beyond what is asked. Brevity is paramount.""",

    "standard": _CORE + """

RESPONSE MODE: STANDARD — 3-8 focused paragraphs.
- Answer the specific question asked. Do NOT write beyond scope.
- Cover: legal position, key authority/statute, main exception, practical next step.
- Use headers ONLY if there are 3+ genuinely distinct sub-issues.
- Briefly note critical missing facts and how they would change the answer.
- Do NOT write a treatise. Focus on what matters most to the user's question.""",

    "comprehensive": _CORE + """

RESPONSE MODE: COMPREHENSIVE — Full structured analysis.
For each issue identified:
- Apply CREAC: Conclusion → Rule (statute + case law) → Explanation → Application → Conclusion
- For each principle: GENERAL RULE → EXCEPTIONS → MINORITY VIEW
- Distinguish: strict law vs equity vs practical enforceability in Nigeria
- DEVIL'S ADVOCATE: state the strongest counter-argument honestly
- Note missing facts and how each would change the analysis
- STRATEGIC: probability assessment, risk, recommended actions with deadlines
- Mark citation confidence: [HIGH CONFIDENCE] / [MODERATE CONFIDENCE] / [VERIFY]
- Flag deadline risks with 🚨 DEADLINE ALERT
- Identify HIDDEN ISSUES the questioner didn't ask about but a Senior Counsel would spot. Mark: ⚠️ HIDDEN ISSUE""",
}

TASK_MODS = {
    "general": "",
    "analysis": "\nTask focus: Legal analysis. Spot issues including hidden ones. Apply structured reasoning to the question asked.",
    "drafting": "\nTask focus: Draft the requested document to professional Nigerian standard. Use [PLACEHOLDER] for missing information. Note required formalities for validity.",
    "research": "\nTask focus: Legal research memo. Cite relevant statutes and case law. Note conflicting authorities and minority views. Mark uncertain citations.",
    "procedure": "\nTask focus: Procedural guidance. Specify correct court, applicable rules, filing requirements, service, deadlines. Flag limitation risks with 🚨.",
    "interpretation": "\nTask focus: Statutory interpretation. Apply literal, golden, and mischief/purposive rules. Note how Nigerian courts have interpreted the provision.",
    "advisory": "\nTask focus: Client advisory. Present viable options with risk/benefit for each. Recommend a strategy. List immediate action items with deadlines.",
}

ISSUE_SPOT_SYSTEM = _CORE + """

YOUR SOLE TASK: Quick issue identification (under 250 words, structured).
- OBVIOUS ISSUES: List each (ISSUE 1, 2…) with area of law and governing principle.
- HIDDEN ISSUES: Issues a junior would miss (HIDDEN A, B…) — limitation traps, locus standi, procedural prerequisites, regulatory dimensions, equitable claims.
- TOP 3-5 MISSING FACTS that would materially change analysis.
- COMPLEXITY RATING: Straightforward / Moderate / Complex / Highly Complex
Do NOT provide full analysis. Issue decomposition ONLY."""

CRITIQUE_SYSTEM = _CORE + """

YOUR SOLE TASK: Brief quality check of the analysis below (under 150 words).
Score 1-5 each: Issue Completeness, Legal Accuracy, Practical Value.
List 1-3 specific gaps with why they matter.
OVERALL GRADE: A / B / C / D. One sentence overall assessment."""

FOLLOWUP_SYSTEM = _CORE + """

Continuing a legal conversation. You have original query, previous analysis, and follow-up question.
- Do NOT repeat previous analysis — focus ONLY on what the follow-up asks.
- If follow-up reveals new issue, address it fresh.
- If asking to go deeper, drill further with additional authority.
- Match response length to the specified mode."""

# ═══════════════════════════════════════════════════════
# REFERENCE DATA
# ═══════════════════════════════════════════════════════
LIMITATION_PERIODS = [
    {"cause": "Simple Contract", "period": "6 years", "authority": "Limitation Act, s. 8(1)(a)"},
    {"cause": "Tort / Negligence", "period": "6 years", "authority": "Limitation Act, s. 8(1)(a)"},
    {"cause": "Personal Injury", "period": "3 years", "authority": "Limitation Act, s. 8(1)(b)"},
    {"cause": "Defamation", "period": "3 years", "authority": "Limitation Act, s. 8(1)(b)"},
    {"cause": "Recovery of Land", "period": "12 years", "authority": "Limitation Act, s. 16"},
    {"cause": "Mortgage", "period": "12 years", "authority": "Limitation Act, s. 18"},
    {"cause": "Recovery of Rent", "period": "6 years", "authority": "Limitation Act, s. 19"},
    {"cause": "Enforcement of Judgment", "period": "12 years", "authority": "Limitation Act, s. 8(1)(d)"},
    {"cause": "Maritime Claims", "period": "2 years", "authority": "Admiralty Jurisdiction Act, s. 10"},
    {"cause": "Labour Disputes", "period": "12 months", "authority": "NIC Act, s. 7(1)(e)"},
    {"cause": "Fundamental Rights", "period": "12 months", "authority": "FREP Rules, Order II r. 1"},
    {"cause": "Tax Assessment Appeal", "period": "30 days", "authority": "FIRS (Est.) Act, s. 59"},
    {"cause": "Public Officer Liability", "period": "3 months notice / 12 months", "authority": "Public Officers Protection Act, s. 2"},
    {"cause": "Insurance Claims", "period": "12 months after disclaimer", "authority": "Insurance Act 2003, s. 72"},
    {"cause": "Winding-Up Petition", "period": "21 days from demand", "authority": "CAMA 2020, s. 572"},
    {"cause": "Election Petition", "period": "21 days after declaration", "authority": "Electoral Act 2022, s. 133(1)"},
    {"cause": "Judicial Review", "period": "3 months", "authority": "FHC Rules, Order 34 r. 3"},
]

COURT_HIERARCHY = [
    {"level": 1, "name": "Supreme Court of Nigeria", "desc": "Final appellate — 7 or 5 Justices", "icon": "🏛️"},
    {"level": 2, "name": "Court of Appeal", "desc": "Intermediate appellate — 16 Divisions", "icon": "⚖️"},
    {"level": 3, "name": "Federal High Court", "desc": "Federal causes: admiralty, revenue, IP, banking", "icon": "🏢"},
    {"level": 3, "name": "State High Courts", "desc": "General civil & criminal per state", "icon": "🏢"},
    {"level": 3, "name": "National Industrial Court", "desc": "Labour & employment disputes", "icon": "🏢"},
    {"level": 3, "name": "Sharia Court of Appeal", "desc": "Islamic personal law appeals", "icon": "🏢"},
    {"level": 3, "name": "Customary Court of Appeal", "desc": "Customary law appeals", "icon": "🏢"},
    {"level": 4, "name": "Magistrate / District Courts", "desc": "Summary jurisdiction", "icon": "📋"},
    {"level": 4, "name": "Area / Customary Courts", "desc": "Customary law first instance", "icon": "📋"},
    {"level": 4, "name": "Sharia Courts", "desc": "Islamic personal law first instance", "icon": "📋"},
    {"level": 5, "name": "Tribunals & Panels", "desc": "Election, Tax Appeal, Code of Conduct", "icon": "📌"},
]

LEGAL_MAXIMS = [
    {"maxim": "Audi alteram partem", "meaning": "Hear the other side — pillar of natural justice"},
    {"maxim": "Nemo judex in causa sua", "meaning": "No one should judge their own cause"},
    {"maxim": "Actus non facit reum nisi mens sit rea", "meaning": "No guilt without guilty mind"},
    {"maxim": "Res judicata", "meaning": "A matter decided — cannot be re-litigated"},
    {"maxim": "Stare decisis", "meaning": "Stand by what has been decided"},
    {"maxim": "Ubi jus ibi remedium", "meaning": "Where there is a right, there is a remedy"},
    {"maxim": "Volenti non fit injuria", "meaning": "No injury to one who consents"},
    {"maxim": "Pacta sunt servanda", "meaning": "Agreements must be honoured"},
    {"maxim": "Nemo dat quod non habet", "meaning": "No one gives what they don't have"},
    {"maxim": "Ignorantia legis neminem excusat", "meaning": "Ignorance of law excuses no one"},
    {"maxim": "Qui facit per alium facit per se", "meaning": "He who acts through another acts himself"},
    {"maxim": "Ex turpi causa non oritur actio", "meaning": "No action from an immoral cause"},
    {"maxim": "Expressio unius est exclusio alterius", "meaning": "Express mention of one excludes others"},
    {"maxim": "Ejusdem generis", "meaning": "General words limited by specific preceding words"},
    {"maxim": "Locus standi", "meaning": "Right or capacity to bring an action"},
    {"maxim": "He who comes to equity must come with clean hands", "meaning": "Equitable relief denied to unconscionable conduct"},
    {"maxim": "Equity regards as done that which ought to be done", "meaning": "Equity treats intended transactions as completed"},
    {"maxim": "Delegatus non potest delegare", "meaning": "A delegate cannot further delegate"},
    {"maxim": "Generalia specialibus non derogant", "meaning": "General provisions don't override specific ones"},
    {"maxim": "Equity follows the law", "meaning": "Equity does not override legal rules except to prevent unconscionability"},
]

# ═══════════════════════════════════════════════════════
# TEMPLATES
# ═══════════════════════════════════════════════════════
TEMPLATES = [
    {"id": "1", "name": "Employment Contract", "cat": "Corporate",
     "content": "EMPLOYMENT CONTRACT\n\nThis Employment Contract is made on [DATE] between:\n\n1. [EMPLOYER NAME] (\"the Employer\")\n   Address: [EMPLOYER ADDRESS] | RC: [NUMBER]\n\n2. [EMPLOYEE NAME] (\"the Employee\")\n   Address: [EMPLOYEE ADDRESS]\n\nTERMS:\n\n1. POSITION: [JOB TITLE]\n2. COMMENCEMENT: [START DATE]\n3. PROBATION: [PERIOD] months\n4. SALARY: N[AMOUNT] monthly\n5. HOURS: [HOURS]/week, Mon-Fri\n6. LEAVE: [NUMBER] days annual\n7. TERMINATION: [NOTICE PERIOD] written notice\n8. CONFIDENTIALITY: Employee maintains confidentiality\n9. GOVERNING LAW: Labour Act of Nigeria\n\nSIGNED:\n_______________ _______________\nEmployer        Employee"},
    {"id": "2", "name": "Tenancy Agreement", "cat": "Property",
     "content": "TENANCY AGREEMENT\n\nMade on [DATE] BETWEEN:\n[LANDLORD] of [ADDRESS] (\"Landlord\")\nAND\n[TENANT] of [ADDRESS] (\"Tenant\")\n\n1. PREMISES: [PROPERTY ADDRESS]\n2. TERM: [DURATION] from [START DATE]\n3. RENT: N[AMOUNT] per [PERIOD]\n4. DEPOSIT: N[AMOUNT] refundable\n5. USE: [Residential/Commercial] only\n6. MAINTENANCE: Tenant keeps premises in good condition\n7. ALTERATIONS: None without Landlord's consent\n8. TERMINATION: [NOTICE PERIOD] written notice\n9. LAW: Lagos Tenancy Law (or applicable state law)\n\nSIGNED:\n_______________ _______________\nLandlord        Tenant"},
    {"id": "3", "name": "Power of Attorney", "cat": "Litigation",
     "content": "GENERAL POWER OF ATTORNEY\n\nI, [GRANTOR NAME], of [ADDRESS], appoint [ATTORNEY NAME] of [ADDRESS] as my Attorney to:\n\n1. Demand, sue for, recover and collect all monies due\n2. Sign and execute contracts and documents\n3. Appear before any court or tribunal\n4. Operate bank accounts\n5. Manage properties and collect rents\n6. Execute and register deeds\n\nThis Power remains in force until revoked in writing.\n\nDated: [DATE]\n_______________\n[GRANTOR NAME]\nWITNESS: _______________"},
    {"id": "4", "name": "Written Address", "cat": "Litigation",
     "content": "IN THE [COURT NAME]\nSUIT NO: [NUMBER]\n\n[PLAINTIFF] v. [DEFENDANT]\n\nWRITTEN ADDRESS OF THE [PLAINTIFF/DEFENDANT]\n\n1.0 INTRODUCTION\n2.0 FACTS\n3.0 ISSUES\n4.0 ARGUMENTS\n5.0 CONCLUSION\n\nDated: [DATE]\n_______________\n[COUNSEL]\nFor: [LAW FIRM]"},
    {"id": "5", "name": "Affidavit", "cat": "Litigation",
     "content": "IN THE [COURT NAME]\nSUIT NO: [NUMBER]\n\nAFFIDAVIT IN SUPPORT OF [MOTION]\n\nI, [DEPONENT], make oath:\n1. I am the [Party].\n2. [Fact 1]\n3. [Fact 2]\n4. This Affidavit is made in good faith.\n\n_______________\nDEPONENT\nSworn at [Location] this [DATE]\nBefore: _______________ COMMISSIONER FOR OATHS"},
    {"id": "6", "name": "Legal Opinion", "cat": "Corporate",
     "content": "LEGAL OPINION — PRIVATE & CONFIDENTIAL\n\nTO: [CLIENT] | FROM: [LAW FIRM] | DATE: [DATE]\nRE: [SUBJECT]\n\n1.0 INTRODUCTION\n2.0 FACTS\n3.0 ISSUES\n4.0 LEGAL FRAMEWORK\n5.0 ANALYSIS\n6.0 CONCLUSION\n7.0 CAVEATS\n\n_______________\n[PARTNER]\nFor: [LAW FIRM]"},
    {"id": "7", "name": "Demand Letter", "cat": "Litigation",
     "content": "[LETTERHEAD]\n[DATE]\n\n[RECIPIENT]\n\nRE: DEMAND FOR N[AMOUNT]\nOUR CLIENT: [CLIENT NAME]\n\nWe are Solicitors to [CLIENT].\n\nFacts: [Background]\n\nPay within 7 DAYS or we institute proceedings.\n\nGovern yourself accordingly.\n\n_______________\n[COUNSEL]\nFor: [LAW FIRM]"},
    {"id": "8", "name": "Board Resolution", "cat": "Corporate",
     "content": "BOARD RESOLUTION — [COMPANY] (RC: [NUMBER])\n[VENUE] — [DATE]\n\nPRESENT: [Directors]\nIN ATTENDANCE: [Company Secretary]\n\nRESOLVED:\n1. [Resolution]\n2. Any Director authorized to execute documents.\n3. Company Secretary to file returns with CAC.\n\nCERTIFIED TRUE COPY\n_______________\nCompany Secretary"},
]
# ═══════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════
def gen_id() -> str:
    return uuid.uuid4().hex[:8]

def fmt_currency(a: float) -> str:
    return f"₦{a:,.2f}"

def esc(t: str) -> str:
    return html.escape(str(t))

def fmt_date(s: str) -> str:
    try:
        return datetime.fromisoformat(s).strftime("%B %d, %Y")
    except Exception:
        return str(s)

def days_until(s: str) -> int:
    try:
        return (datetime.fromisoformat(s).date() - datetime.now().date()).days
    except Exception:
        return 999

def relative_date(s: str) -> str:
    d = days_until(s)
    if d == 0: return "Today"
    if d == 1: return "Tomorrow"
    if d == -1: return "Yesterday"
    if 0 < d <= 7: return f"In {d} days"
    if -7 <= d < 0: return f"{abs(d)} days ago"
    return fmt_date(s)

def normalize_model(name: str) -> str:
    c = (name or "").strip()
    m = MODEL_MIGRATION.get(c, c)
    return m if m in SUPPORTED_MODELS else DEFAULT_MODEL


# ═══════════════════════════════════════════════════════
# FILE EXTRACTION
# ═══════════════════════════════════════════════════════
def extract_file_text(uploaded_file) -> str:
    """Extract text from uploaded file. Raises on unsupported type."""
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    # Size check
    mb = len(data) / (1024 * 1024)
    if mb > MAX_UPLOAD_MB:
        raise ValueError(f"File too large ({mb:.1f}MB). Max is {MAX_UPLOAD_MB}MB.")

    if name.endswith(".pdf"):
        if not HAS_PDF_READ:
            raise RuntimeError("PDF support not installed. Run: pip install pdfplumber")
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
            return "\n".join(pages)
    elif name.endswith(".docx"):
        if not HAS_DOCX:
            raise RuntimeError("DOCX support not installed. Run: pip install python-docx")
        doc = DocxDoc(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    elif name.endswith(".csv"):
        return pd.read_csv(io.BytesIO(data)).to_string(index=False)
    elif name.endswith(".xlsx"):
        if not HAS_XLSX:
            raise RuntimeError("Excel support not installed. Run: pip install openpyxl")
        return pd.read_excel(io.BytesIO(data)).to_string(index=False)
    else:
        raise ValueError(f"Unsupported file type: {name.split('.')[-1]}")


# ═══════════════════════════════════════════════════════
# EXPORT FUNCTIONS
# ═══════════════════════════════════════════════════════
def export_txt(text: str, title: str = "LexiAssist Analysis") -> str:
    """Export analysis as plain text string."""
    stamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    header = f"{'='*60}\n{title}\nGenerated: {stamp}\n{'='*60}\n\n"
    footer = f"\n\n{'='*60}\nDisclaimer: AI-generated legal information. Not legal advice.\nVerify all citations independently. Apply professional judgment.\n{'='*60}"
    return header + text + footer


def export_html(text: str, title: str = "LexiAssist Analysis") -> str:
    """Export analysis as styled HTML."""
    stamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    escaped = esc(text)
    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>{esc(title)}</title>
<style>
body{{font-family:Georgia,serif;line-height:1.8;max-width:850px;margin:40px auto;padding:20px;color:#1e293b}}
h1{{color:#059669;border-bottom:3px solid #059669;padding-bottom:12px;font-size:1.6rem}}
.meta{{color:#64748b;font-size:.85rem;margin-bottom:2rem}}
.content{{white-space:pre-wrap;font-size:15px}}
.disclaimer{{background:#fef3c7;border-left:4px solid #f59e0b;padding:16px;margin-top:32px;border-radius:0 8px 8px 0;font-size:.9rem}}
</style></head><body>
<h1>⚖️ {esc(title)}</h1>
<div class="meta">Generated: {stamp}</div>
<div class="content">{escaped}</div>
<div class="disclaimer"><strong>Disclaimer:</strong> AI-generated legal information for professional reference only. Not legal advice. Verify all citations independently.</div>
</body></html>"""


def export_pdf(text: str, title: str = "LexiAssist Analysis") -> bytes:
    """Export analysis as PDF bytes. Returns empty bytes if fpdf2 not installed."""
    if not HAS_PDF_WRITE:
        return b""
    stamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    safe_title = title.encode("latin-1", "replace").decode("latin-1")
    pdf.cell(0, 12, safe_title, ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, f"Generated: {stamp}", ln=True)
    pdf.ln(6)
    pdf.set_draw_color(5, 150, 105)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)
    pdf.set_text_color(30, 41, 59)
    pdf.set_font("Helvetica", "", 10)
    safe_text = text.replace("\u2014", "--").replace("\u2013", "-")
    safe_text = safe_text.replace("\u2018", "'").replace("\u2019", "'")
    safe_text = safe_text.replace("\u201c", '"').replace("\u201d", '"')
    safe_text = safe_text.replace("\u20a6", "N")  # Naira sign
    safe_text = safe_text.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 5.5, safe_text)
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 4.5, "Disclaimer: AI-generated legal information. Not legal advice. Verify citations. Apply professional judgment.")
    return pdf.output()


def export_docx(text: str, title: str = "LexiAssist Analysis") -> bytes:
    """Export analysis as DOCX bytes. Returns empty bytes if python-docx not installed."""
    if not HAS_DOCX:
        return b""
    stamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    doc = DocxDoc()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_heading(title, level=1)
    meta = doc.add_paragraph(f"Generated: {stamp}")
    meta.style.font.size = Pt(9)
    doc.add_paragraph("─" * 60)
    for para in text.split("\n"):
        if para.strip():
            doc.add_paragraph(para)
        else:
            doc.add_paragraph("")
    doc.add_paragraph("─" * 60)
    disclaimer = doc.add_paragraph(
        "Disclaimer: AI-generated legal information for professional reference only. "
        "Not legal advice. Verify all citations independently."
    )
    disclaimer.runs[0].italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════
# API LAYER
# ═══════════════════════════════════════════════════════
def get_api_key(secrets_fn=None, session_key: str = "") -> str:
    """Get API key from secrets, env, or session."""
    sources = [
        lambda: (secrets_fn() if secrets_fn else ""),
        lambda: os.getenv("GEMINI_API_KEY", ""),
        lambda: session_key,
    ]
    for fn in sources:
        try:
            k = fn()
            if k and len(k.strip()) >= 10:
                return k.strip()
        except Exception:
            pass
    return ""


def configure_api(key: str):
    """Configure genai with key."""
    genai.configure(api_key=key, transport="rest")


def test_connection(key: str, model_name: str) -> tuple[bool, str]:
    """Test API connection. Returns (success, message)."""
    try:
        configure_api(key)
        m = genai.GenerativeModel(normalize_model(model_name))
        m.generate_content("Say OK", generation_config={"max_output_tokens": 8})
        return True, "Connected"
    except Exception as e:
        s = str(e)
        if "403" in s:
            return False, "API key invalid or unauthorized."
        if "429" in s:
            return False, "Rate limit hit. Wait and retry."
        return False, f"Connection error: {s}"


def generate(prompt: str, system: str, model_name: str, gen_cfg: dict) -> str:
    """Generate response from Gemini. Retries up to 3 times."""
    model_name = normalize_model(model_name)
    try:
        model = genai.GenerativeModel(model_name, system_instruction=system)
    except TypeError:
        model = genai.GenerativeModel(model_name)
        prompt = f"{system}\n\n---\n\n{prompt}"

    for attempt in range(3):
        try:
            resp = model.generate_content(prompt, generation_config=gen_cfg)
            return resp.text
        except Exception as e:
            if attempt == 2:
                return f"Error: {e}"
            time.sleep(1.5 * (attempt + 1))
    return "Error: generation failed after retries."


# ═══════════════════════════════════════════════════════
# SMART PIPELINE
# ═══════════════════════════════════════════════════════
def build_system_prompt(mode: str, task: str) -> str:
    """Build system prompt from mode and task type."""
    base = SYSTEM_PROMPTS.get(mode, SYSTEM_PROMPTS["standard"])
    mod = TASK_MODS.get(task, "")
    return base + mod


def run_pipeline(
    query: str,
    task: str,
    mode: str,
    model_name: str,
    doc_context: str = "",
    conv_context: str = "",
) -> dict[str, str]:
    """Run the smart pipeline. Returns dict with keys: issue_spot, ambiguity, main, critique, grade."""
    result = {"issue_spot": "", "ambiguity": "", "main": "", "critique": "", "grade": ""}
    tokens = RESPONSE_MODES.get(mode, RESPONSE_MODES["standard"])["tokens"]
    system = build_system_prompt(mode, task)
    now = datetime.now().strftime("%d %B %Y")

    # Build prompt parts
    parts = [f"DATE: {now}"]
    if doc_context:
        parts.append(f"DOCUMENT CONTEXT:\n{doc_context[:3000]}")
    if conv_context:
        parts.append(f"PRIOR CONTEXT:\n{conv_context[:1500]}")

    gen_main = {"temperature": 0.2, "top_p": 0.88, "top_k": 35, "max_output_tokens": tokens}
    gen_fast = {"temperature": 0.15, "top_p": 0.85, "top_k": 25, "max_output_tokens": 800}
    gen_crit = {"temperature": 0.15, "top_p": 0.85, "top_k": 25, "max_output_tokens": 500}

    if mode == "comprehensive":
        # Pass 1: Issue Spotting
        result["issue_spot"] = generate(
            f"LEGAL SCENARIO:\n\n{query}", ISSUE_SPOT_SYSTEM, model_name, gen_fast
        )
        # Pass 2: Build main prompt with issue context
        parts.append(f"PRE-ANALYSIS — ISSUES IDENTIFIED:\n{result['issue_spot']}")
        parts.append(f"QUERY:\n{query}")
        parts.append("Address EVERY issue identified above (including hidden issues). Apply your full reasoning framework.")
        prompt = "\n\n".join(parts)
        result["main"] = generate(prompt, system, model_name, gen_main)

        # Pass 3: Self-Critique
        if not result["main"].startswith(("Error", "⚠️")):
            critique_prompt = (
                f"QUERY:\n{query}\n\nISSUES:\n{result['issue_spot']}\n\nANALYSIS TO CRITIQUE:\n{result['main']}"
            )
            result["critique"] = generate(critique_prompt, CRITIQUE_SYSTEM, model_name, gen_crit)
            result["grade"] = extract_grade(result["critique"])
    else:
        # Brief / Standard: single call
        parts.append(f"QUERY:\n{query}")
        prompt = "\n\n".join(parts)
        result["main"] = generate(prompt, system, model_name, gen_main)

    return result


def run_followup(
    original_query: str,
    original_response: str,
    followup: str,
    mode: str,
    task: str,
    model_name: str,
) -> str:
    """Run a follow-up query with context."""
    tokens = RESPONSE_MODES.get(mode, RESPONSE_MODES["standard"])["tokens"]
    system = FOLLOWUP_SYSTEM + f"\n\nRESPONSE MODE: {mode.upper()} — max {tokens} tokens." + TASK_MODS.get(task, "")
    prompt = (
        f"ORIGINAL QUERY:\n{original_query[:1000]}\n\n"
        f"PREVIOUS ANALYSIS:\n{original_response[:3000]}\n\n"
        f"FOLLOW-UP QUESTION:\n{followup}\n\n"
        f"Address the follow-up. Do NOT repeat prior analysis."
    )
    gen_cfg = {"temperature": 0.2, "top_p": 0.88, "top_k": 35, "max_output_tokens": tokens}
    return generate(prompt, system, model_name, gen_cfg)


def run_research(query: str, mode: str, model_name: str) -> str:
    """Run a legal research query."""
    tokens = RESPONSE_MODES.get(mode, RESPONSE_MODES["standard"])["tokens"]
    system = build_system_prompt(mode, "research")
    gen_cfg = {"temperature": 0.2, "top_p": 0.88, "top_k": 35, "max_output_tokens": tokens}
    prompt = (
        f"RESEARCH QUESTION:\n{query}\n\n"
        f"Provide a focused legal research response. Cite statutes and cases. "
        f"Mark uncertain citations: [Citation to be verified]."
    )
    return generate(prompt, system, model_name, gen_cfg)


def extract_grade(text: str) -> str:
    """Extract quality grade from critique text."""
    m = re.search(r"OVERALL GRADE:\s*([A-D])", text, re.IGNORECASE)
    return m.group(1).upper() if m else "B"


# ═══════════════════════════════════════════════════════
# HISTORY PERSISTENCE
# ═══════════════════════════════════════════════════════
def save_history(history: list[dict]) -> bool:
    """Save conversation history to file. Returns True on success."""
    try:
        # Keep last 100 entries
        trimmed = history[-100:] if len(history) > 100 else history
        HISTORY_FILE.write_text(json.dumps(trimmed, indent=2, default=str))
        return True
    except Exception as e:
        logger.warning(f"Failed to save history: {e}")
        return False


def load_history() -> list[dict]:
    """Load conversation history from file."""
    try:
        if HISTORY_FILE.exists():
            data = json.loads(HISTORY_FILE.read_text())
            if isinstance(data, list):
                return data
    except Exception as e:
        logger.warning(f"Failed to load history: {e}")
    return []


def clear_history() -> bool:
    """Delete history file."""
    try:
        if HISTORY_FILE.exists():
            HISTORY_FILE.unlink()
        return True
    except Exception as e:
        logger.warning(f"Failed to clear history: {e}")
        return False


# ═══════════════════════════════════════════════════════
# DATA CRUD — Cases, Clients, Time, Invoices
# ═══════════════════════════════════════════════════════
def add_case(cases: list, data: dict) -> list:
    data["id"] = gen_id()
    data["created_at"] = datetime.now().isoformat()
    cases.append(data)
    return cases


def update_case(cases: list, case_id: str, updates: dict) -> list:
    for c in cases:
        if c["id"] == case_id:
            c.update(updates)
            c["updated_at"] = datetime.now().isoformat()
    return cases


def delete_case(cases: list, case_id: str) -> list:
    return [c for c in cases if c["id"] != case_id]


def add_client(clients: list, data: dict) -> list:
    data["id"] = gen_id()
    data["created_at"] = datetime.now().isoformat()
    clients.append(data)
    return clients


def delete_client(clients: list, client_id: str) -> list:
    return [c for c in clients if c["id"] != client_id]


def client_name(clients: list, client_id: str) -> str:
    for c in clients:
        if c["id"] == client_id:
            return c["name"]
    return "—"


def add_time_entry(entries: list, data: dict) -> list:
    data["id"] = gen_id()
    data["created_at"] = datetime.now().isoformat()
    data["amount"] = data["hours"] * data["rate"]
    entries.append(data)
    return entries


def delete_time_entry(entries: list, entry_id: str) -> list:
    return [e for e in entries if e["id"] != entry_id]


def make_invoice(invoices: list, entries: list, clients: list, client_id: str) -> dict | None:
    client_entries = [e for e in entries if e.get("client_id") == client_id]
    if not client_entries:
        return None
    inv = {
        "id": gen_id(),
        "invoice_no": f"INV-{datetime.now():%Y%m%d}-{gen_id()[:4].upper()}",
        "client_id": client_id,
        "client_name": client_name(clients, client_id),
        "entries": client_entries,
        "total": sum(e["amount"] for e in client_entries),
        "date": datetime.now().isoformat(),
        "status": "Draft",
    }
    invoices.append(inv)
    return inv


def total_billable(entries: list) -> float:
    return sum(e.get("amount", 0) for e in entries)


def total_hours(entries: list) -> float:
    return sum(e.get("hours", 0) for e in entries)


def client_billable(entries: list, client_id: str) -> float:
    return sum(e.get("amount", 0) for e in entries if e.get("client_id") == client_id)


def client_case_count(cases: list, client_id: str) -> int:
    return sum(1 for c in cases if c.get("client_id") == client_id)


def get_hearings(cases: list, limit: int = 10) -> list[dict]:
    """Get upcoming hearings sorted by date. Returns a list (not tuple)."""
    hearings = []
    for c in cases:
        if c.get("next_hearing") and c.get("status") == "Active":
            hearings.append({
                "id": c["id"],
                "title": c["title"],
                "date": c["next_hearing"],
                "court": c.get("court", ""),
                "suit": c.get("suit_no", ""),
            })
    hearings.sort(key=lambda x: x["date"])
    return hearings[:limit]


def export_all_data(cases, clients, entries, invoices) -> str:
    """Export all app data as JSON string."""
    return json.dumps({
        "cases": cases,
        "clients": clients,
        "time_entries": entries,
        "invoices": invoices,
    }, indent=2, default=str)